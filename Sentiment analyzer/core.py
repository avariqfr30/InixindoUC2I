import concurrent.futures
import hashlib
import io
import json
import logging
import os
from pathlib import Path
import re
import threading
import textwrap
from datetime import datetime, timedelta
from urllib.parse import urlparse

import chromadb
import diskcache as dc
import markdown
import matplotlib
import matplotlib.patches as patches
import matplotlib.pyplot as plt
import pandas as pd
import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from chromadb.config import Settings
from chromadb.utils import embedding_functions
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from sqlalchemy import create_engine

from pydantic import BaseModel, Field
from ollama import Client

from config import (
    APP_MODE,
    ADOPTION_READINESS_PILLARS,
    CUSTOMER_JOURNEY_STAGES,
    CX_SENTIMENT_STRUCTURE,
    CSV_PATH,
    DATA_DIR,
    DEFAULT_COLOR,
    DEFAULT_SCORE_ENGINE,
    EMBED_MODEL,
    ENABLE_VECTOR_INDEX,
    EXTERNAL_DATA_MODE,
    INTERNAL_DATA_MODE,
    OLLAMA_HOST,
    OSINT_BASE_QUERIES,
    OSINT_CACHE_PATH,
    OSINT_CACHE_TTL_SECONDS,
    OSINT_MAX_SIGNALS,
    OSINT_RECENCY,
    OSINT_RESULTS_PER_QUERY,
    OSINT_SEARCH_LANGUAGE,
    OSINT_SEARCH_REGION,
    SCORE_ENGINE_PROFILES,
    SERPER_API_KEY,
    SENTIMENT_OPTIONS,
    WRITER_FIRM_NAME,
)
from internal_api import InternalApiClient

matplotlib.use("Agg")

logger = logging.getLogger(__name__)

# ==========================================
# PYDANTIC SCHEMAS & FAST CACHING
# ==========================================
class InsightSchema(BaseModel):
    insight: str = Field(description="The extracted insight in Indonesian. 'NOT_FOUND' if missing.")

# Initialize ultra-fast disk caching for OSINT
osint_cache_dir = Path(OSINT_CACHE_PATH).parent / '.osint_cache' if hasattr(OSINT_CACHE_PATH, "parent") else os.path.dirname(OSINT_CACHE_PATH) + '/.osint_cache'
osint_cache = dc.Cache(str(osint_cache_dir))
# ==========================================

CANONICAL_INTERNAL_COLUMNS = (
    "Record ID",
    "Sumber Feedback",
    "Kanal Feedback",
    "Tanggal Feedback",
    "Tipe Stakeholder",
    "Layanan",
    "Lokasi",
    "Tipe Instruktur",
    "Rentang Waktu",
    "Rating",
    "Komentar",
    "Customer Journey Hint",
)

COLUMN_ALIASES = {
    "Record ID": ("record_id", "id", "feedback_id", "ticket_id", "case_id"),
    "Sumber Feedback": ("sumber feedback", "source", "feedback_source", "origin", "source_name"),
    "Kanal Feedback": ("kanal feedback", "channel", "feedback_channel", "touchpoint", "platform", "kanal"),
    "Tanggal Feedback": ("tanggal feedback", "feedback_date", "created_at", "submitted_at", "date", "tanggal"),
    "Tipe Stakeholder": ("tipe stakeholder", "stakeholder_type", "stakeholder", "customer_segment", "customer_type", "segment", "segmen"),
    "Layanan": ("layanan", "service", "service_name", "product", "offering", "service_type"),
    "Lokasi": ("lokasi", "location", "training_location", "city", "kota", "venue_location"),
    "Tipe Instruktur": ("tipe instruktur", "instructor_type", "trainer_type", "coach_type", "internal_ol", "internal_or_ol", "trainer_origin"),
    "Rentang Waktu": ("rentang waktu", "timeframe", "periode", "period", "reporting_period"),
    "Rating": ("rating", "score", "csat", "sentiment_score", "nilai"),
    "Komentar": ("komentar", "comment", "feedback", "feedback_text", "review", "notes", "complaint_text", "customer_comment"),
    "Customer Journey Hint": ("customer_journey_hint", "journey_hint", "journey_stage", "customer_journey_stage", "touchpoint_stage"),
}

DATE_COLUMN_ALIASES = (
    "tanggal feedback", "tanggal", "date", "created_at", "submitted_at", "feedback_date",
)

def append_field(paragraph, instruction):
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    field_instruction = OxmlElement("w:instrText")
    field_instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    field_instruction.text = instruction
    field_separator = OxmlElement("w:fldChar")
    field_separator.set(qn("w:fldCharType"), "separate")
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, field_instruction, field_separator, field_end])

class InternalDataProvider:
    source_name = "internal"

    @staticmethod
    def _normalize_token(value):
        return re.sub(r"[^a-z0-9]+", "_", str(value).strip().lower()).strip("_")

    @classmethod
    def _column_variants(cls, column_name):
        raw_name = str(column_name).strip()
        normalized = cls._normalize_token(raw_name)
        variants = {normalized}

        for part in re.split(r"[.\[\]/\\]+", raw_name):
            token = cls._normalize_token(part)
            if token:
                variants.add(token)

        normalized_parts = [part for part in normalized.split("_") if part]
        if normalized_parts:
            variants.add(normalized_parts[-1])
            if len(normalized_parts) >= 2:
                variants.add("_".join(normalized_parts[-2:]))

        return variants

    @classmethod
    def _find_matching_column(cls, columns, aliases):
        alias_tokens = [cls._normalize_token(alias) for alias in aliases if cls._normalize_token(alias)]
        best_match = None
        best_score = -1

        for column_name in columns:
            variants = cls._column_variants(column_name)
            score = 0
            for alias_token in alias_tokens:
                if alias_token in variants:
                    score += 100
                elif any(
                    variant.endswith(f"_{alias_token}")
                    or variant.startswith(f"{alias_token}_")
                    or (len(alias_token) >= 4 and alias_token in variant)
                    for variant in variants
                ):
                    score += 60
            if score > best_score:
                best_match = column_name
                best_score = score

        return best_match if best_score > 0 else None

    @classmethod
    def normalize_dataframe(cls, raw_df):
        if raw_df is None:
            return pd.DataFrame(columns=list(CANONICAL_INTERNAL_COLUMNS))

        dataframe = raw_df.copy()
        dataframe.columns = [str(column).strip() for column in dataframe.columns]

        rename_map = {}
        for canonical_name, aliases in COLUMN_ALIASES.items():
            if canonical_name in dataframe.columns:
                continue
            matched_column = cls._find_matching_column(dataframe.columns, aliases)
            if matched_column:
                rename_map[matched_column] = canonical_name

        dataframe = dataframe.rename(columns=rename_map)

        feedback_dates = pd.Series(dtype="datetime64[ns]")
        for alias in DATE_COLUMN_ALIASES:
            matched_column = cls._find_matching_column(dataframe.columns, (alias,))
            if matched_column:
                feedback_dates = pd.to_datetime(
                    dataframe[matched_column],
                    errors="coerce",
                )
                if feedback_dates.notna().any():
                    dataframe["Tanggal Feedback"] = feedback_dates.dt.strftime("%Y-%m-%d")
                    break

        if "Rentang Waktu" not in dataframe.columns:
            for alias in DATE_COLUMN_ALIASES:
                matched_column = cls._find_matching_column(dataframe.columns, (alias,))
                if not matched_column:
                    continue
                parsed_dates = pd.to_datetime(
                    dataframe[matched_column],
                    errors="coerce",
                )
                if parsed_dates.notna().any():
                    dataframe["Rentang Waktu"] = parsed_dates.dt.to_period("M").astype(str)
                    break

        for column_name in CANONICAL_INTERNAL_COLUMNS:
            if column_name not in dataframe.columns:
                dataframe[column_name] = pd.NA

        if dataframe["Record ID"].isna().all():
            dataframe["Record ID"] = [f"FB-{index + 1:05d}" for index in range(len(dataframe))]

        dataframe["Rating"] = pd.to_numeric(dataframe["Rating"], errors="coerce")
        for column_name in CANONICAL_INTERNAL_COLUMNS:
            dataframe[column_name] = dataframe[column_name].fillna("").astype(str).str.strip()

        return dataframe

    def load_feedback_data(self):
        raise NotImplementedError

class DemoCsvProvider(InternalDataProvider):
    source_name = "demo_csv"

    def load_feedback_data(self):
        if not os.path.exists(CSV_PATH):
            raise FileNotFoundError(f"CSV file not found: {CSV_PATH}")
        raw_df = pd.read_csv(CSV_PATH)
        return self.normalize_dataframe(raw_df)

class InternalApiProvider(InternalDataProvider):
    source_name = "company_api"

    def __init__(self):
        self.client = InternalApiClient()
        self.dataset_name = "feedback"

    def load_dataset(self, dataset_name, extra_params=None):
        raw_df = pd.DataFrame(
            self.client.fetch_records(dataset_name, extra_params=extra_params)
        )
        if raw_df.empty:
            raise ValueError(
                f"Internal API returned no records for endpoint '{dataset_name}'."
            )
        return self.normalize_dataframe(raw_df)

    def load_feedback_data(self):
        return self.load_dataset(self.dataset_name)

class KnowledgeBase:
    def __init__(self, db_uri):
        os.makedirs(DATA_DIR, exist_ok=True)
        self.engine = create_engine(db_uri)
        self.app_mode = APP_MODE
        self.internal_data_mode = INTERNAL_DATA_MODE
        self.external_data_mode = EXTERNAL_DATA_MODE
        self.enable_vector_index = ENABLE_VECTOR_INDEX
        self.provider = self._build_provider()
        self.chroma = None
        self.embed_fn = None
        self.collection = None
        if self.enable_vector_index:
            self.chroma = chromadb.Client(Settings(anonymized_telemetry=False))
            self.embed_fn = embedding_functions.OllamaEmbeddingFunction(
                url=f"{OLLAMA_HOST}/api/embeddings",
                model_name=EMBED_MODEL,
            )
            self.collection = self.chroma.get_or_create_collection(
                name="cx_holistic_db",
                embedding_function=self.embed_fn,
            )
        self.df = None
        self.refresh_data()

    def _build_provider(self):
        if self.internal_data_mode == "api":
            return InternalApiProvider()
        return DemoCsvProvider()

    def _load_cached_dataframe(self):
        try:
            cached_df = pd.read_sql("SELECT * FROM feedback", self.engine)
            if cached_df is not None and not cached_df.empty:
                logger.warning("Using cached internal data from SQLite.")
                return cached_df
        except Exception as exc:
            logger.warning("No cached internal dataset available: %s", exc)
        return None

    def _rebuild_vector_store(self):
        if not self.enable_vector_index or self.collection is None:
            return True

        if self.df is None or self.df.empty:
            return False

        existing_ids = self.collection.get().get("ids", [])
        if existing_ids:
            self.collection.delete(existing_ids)

        ids, documents, metadata = [], [], []
        for index, row in self.df.iterrows():
            text_representation = " | ".join(
                f"{column}: {value}" for column, value in row.items()
            )
            ids.append(str(index))
            documents.append(text_representation)
            metadata.append(row.astype(str).to_dict())

        if not ids:
            return False

        logger.info(
            "Sending %s feedback records to Ollama embeddings (%s).",
            len(ids),
            OLLAMA_HOST,
        )
        self.collection.add(documents=documents, metadatas=metadata, ids=ids)
        return True

    def refresh_data(self):
        try:
            latest_df = self.provider.load_feedback_data()
            latest_df.to_sql("feedback", self.engine, index=False, if_exists="replace")
            self.df = latest_df
        except Exception as exc:
            logger.error(
                "Failed to load internal data from %s: %s",
                self.provider.source_name,
                exc,
            )
            self.df = self._load_cached_dataframe()
            if self.df is None or self.df.empty:
                return False

        try:
            return self._rebuild_vector_store()
        except Exception as exc:
            logger.error("Failed to rebuild vector store: %s", exc)
            return False

    def query(self, timeframe, context_keywords=""):
        if not self.enable_vector_index or self.collection is None:
            filtered_df = self.df if self.df is not None else pd.DataFrame()
            if timeframe and not filtered_df.empty:
                filtered_df = filtered_df[filtered_df["Rentang Waktu"] == timeframe]
            if filtered_df.empty:
                return "Tidak ada data feedback internal untuk periode ini."
            limited_rows = filtered_df.head(25)
            documents = []
            for _, row in limited_rows.iterrows():
                documents.append(
                    " | ".join(f"{column}: {value}" for column, value in row.items())
                )
            return "\n---\n".join(documents)

        query_text = (
            f"General feedback, complaints, praise, and operational issues. "
            f"{context_keywords}"
        ).strip()
        query_payload = {"query_texts": [query_text], "n_results": 25}
        if timeframe:
            query_payload["where"] = {"Rentang Waktu": timeframe}

        try:
            result = self.collection.query(**query_payload)
            documents = result.get("documents", [[]])
            if documents and documents[0]:
                return "\n---\n".join(documents[0])
        except Exception as exc:
            logger.error("Query error: %s", exc)

        return "Tidak ada data feedback internal untuk periode ini."


class Researcher:
    SERPER_ENDPOINT = "https://google.serper.dev/search"
    INVALID_API_KEYS = {"", "YOUR_SERPER_API_KEY", "masukkan_api_key_serper_anda_disini"}

    @staticmethod
    def _is_enabled():
        return (SERPER_API_KEY or "").strip() not in Researcher.INVALID_API_KEYS

    @staticmethod
    def fetch_full_markdown(url):
        """Fetches the clean markdown text of any URL using Jina Reader."""
        if not url: return ""
        try:
            jina_url = f"https://r.jina.ai/{url}"
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(jina_url, headers=headers, timeout=12)
            if response.status_code == 200:
                return response.text[:6000] 
            return ""
        except Exception as e:
            logger.warning("Failed to fetch full markdown for %s: %s", url, e)
            return ""

    @classmethod
    def extract_insight_with_llm(cls, url, extraction_goal):
        """Universal Deep Scraper: Reads a URL and extracts a specific qualitative insight via Pydantic/LLM."""
        markdown_text = cls.fetch_full_markdown(url)
        if not markdown_text:
            return ""
            
        prompt = f"""
        You are an expert business researcher. Read the following source text.
        Your goal is to extract: {extraction_goal}
        
        SOURCE TEXT:
        {markdown_text}
        
        Respond ONLY with a valid JSON object using this schema. If the information is not present, use "NOT_FOUND".
        {{
            "insight": "<concise professional summary in Indonesian>"
        }}
        """
        try:
            llm_model = os.getenv("LLM_MODEL", "gpt-oss:120b-cloud")
            client = Client(host=OLLAMA_HOST)
            res = client.chat(
                model=llm_model,
                messages=[{'role': 'user', 'content': prompt}],
                options={'temperature': 0.0}
            )
            raw_text = res['message']['content']
            match = re.search(r'\{.*\}', raw_text, re.DOTALL)
            parsed_dict = json.loads(match.group(0)) if match else json.loads(raw_text)
            
            # Use Pydantic to strictly validate
            data = InsightSchema.model_validate(parsed_dict)
            
            if "NOT_FOUND" in data.insight.upper() or not data.insight:
                return ""
            return data.insight
        except Exception as e:
            logger.warning("Insight extraction failed for %s: %s", url, e)
            return ""

    @staticmethod
    def _search_serper(query, max_results=OSINT_RESULTS_PER_QUERY):
        payload = {
            "q": query,
            "num": max_results,
            "gl": OSINT_SEARCH_REGION,
            "hl": OSINT_SEARCH_LANGUAGE,
        }
        if OSINT_RECENCY:
            payload["tbs"] = OSINT_RECENCY

        headers = {
            "X-API-KEY": SERPER_API_KEY,
            "Content-Type": "application/json",
        }

        response = requests.post(
            Researcher.SERPER_ENDPOINT,
            headers=headers,
            json=payload,
            timeout=10,
        )
        response.raise_for_status()
        return response.json()

    @staticmethod
    def _extract_items(query, payload):
        items = []
        for position, entry in enumerate(payload.get("organic", []), start=1):
            items.append({
                "query": query, "title": entry.get("title", "Tanpa Judul"), "snippet": entry.get("snippet", "").strip(),
                "url": entry.get("link", "").strip(), "date": entry.get("date", "").strip(), "source_type": "organic", "position": position
            })
        for position, entry in enumerate(payload.get("news", []), start=1):
            items.append({
                "query": query, "title": entry.get("title", "Tanpa Judul"), "snippet": entry.get("snippet", "").strip(),
                "url": entry.get("link", "").strip(), "date": entry.get("date", "").strip(), "source_type": "news", "position": position
            })
        return [item for item in items if item["url"]]

    @staticmethod
    def _deduplicate_items(items):
        seen_keys = set()
        unique_items = []
        for item in items:
            key = item["url"] or f"{item['title']}::{item['snippet']}"
            if key in seen_keys: continue
            seen_keys.add(key)
            unique_items.append(item)
        return unique_items

    @staticmethod
    def _score_items(items, context_text):
        keywords = {
            token.lower() for token in re.findall(r"[A-Za-z]{4,}", context_text)
            if token.lower() not in {"yang", "dengan", "untuk", "pada", "from", "this"}
        }
        for item in items:
            corpus = f"{item['title']} {item['snippet']}".lower()
            coverage_score = sum(1 for keyword in keywords if keyword in corpus)
            freshness_bonus = 1 if item["source_type"] == "news" else 0
            ranking_penalty = (item["position"] - 1) * 0.05
            item["score"] = coverage_score + freshness_bonus - ranking_penalty
        return sorted(items, key=lambda value: value["score"], reverse=True)

    @staticmethod
    def _source_domain(url):
        try:
            netloc = urlparse(url).netloc.lower()
            return netloc.replace("www.", "") if netloc else "unknown"
        except Exception:
            return "unknown"

    @staticmethod
    def _format_osint_brief(items, title):
        if not items:
            return "Tidak ada sinyal OSINT eksternal yang dapat digunakan untuk benchmark periode ini."
        lines = [f"{title}:"]
        for index, item in enumerate(items, start=1):
            date_part = f" | tanggal={item['date']}" if item["date"] else ""
            lines.append((f"{index}. {item['title']} | {item['snippet']} | sumber={Researcher._source_domain(item['url'])}{date_part} | url={item['url']}").strip())
        return "\n".join(lines)

    @staticmethod
    def _run_query_batch(queries, max_signals=OSINT_MAX_SIGNALS):
        collected = []
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(4, max(1, len(queries)))) as pool:
            future_map = {pool.submit(Researcher._search_serper, query): query for query in queries}
            for future in concurrent.futures.as_completed(future_map):
                query = future_map[future]
                try:
                    payload = future.result()
                    collected.extend(Researcher._extract_items(query, payload))
                except Exception as exc:
                    logger.warning("OSINT query gagal (%s): %s", query, exc)

        deduplicated = Researcher._deduplicate_items(collected)
        ranked = Researcher._score_items(deduplicated, " ".join(queries))
        return ranked[:max_signals]

    @staticmethod
    @osint_cache.memoize(expire=OSINT_CACHE_TTL_SECONDS)
    def get_macro_trends(timeframe, notes="", score_engine_label="Experience Index"):
        scope = timeframe or "periode terbaru"
        compact_notes = re.sub(r"\s+", " ", notes).strip()
        contextual_query = (
            f"benchmark sentimen pelanggan pelatihan dan konsultasi IT Indonesia {scope} {score_engine_label}"
        )
        if compact_notes:
            contextual_query += f" {compact_notes[:140]}"

        queries = [f"{query} {scope}" for query in OSINT_BASE_QUERIES] + [contextual_query]

        if not Researcher._is_enabled():
            return "Data OSINT eksternal tidak tersedia (SERPER_API_KEY belum diatur)."

        try:
            findings = Researcher._run_query_batch(queries)

            # --- DEEP SCRAPE THE #1 RESULT ---
            deep_insight_text = ""
            if findings and findings[0].get("url"):
                top_link = findings[0]["url"]
                logger.info("Deep scraping OSINT for macro trends: %s", top_link)
                goal = "What are the latest macro trends, challenges, or benchmarks regarding IT training, consulting, and customer expectations in Indonesia?"
                insight = Researcher.extract_insight_with_llm(top_link, goal)
                if insight:
                    source = Researcher._source_domain(top_link)
                    deep_insight_text = f"**Insight Mendalam (via {source}):** {insight}\n\n"

            brief = Researcher._format_osint_brief(findings, "Sinyal OSINT Makro (Indonesia)")
            return deep_insight_text + brief
        except Exception as exc:
            logger.warning("OSINT macro trends failed: %s", exc)
            return "Tidak ada tren eksternal yang berhasil dimuat."

class FeedbackAnalyticsEngine:
    THEME_LIBRARY = (
        {
            "id": "responsiveness", "label": "Respons dan SLA",
            "keywords": ("lambat", "respon", "response", "sla", "timeline", "delay", "mundur", "follow up"),
            "prescription": "Tetapkan SLA respon, dashboard aging, dan owner follow-up per tiket/permintaan.",
        },
        {
            "id": "schedule", "label": "Jadwal dan beban sesi",
            "keywords": ("jadwal", "padat", "jeda", "durasi", "sesi", "waktu"),
            "prescription": "Kalibrasi durasi sesi, sediakan jeda terstruktur, dan review desain agenda per layanan.",
        },
        {
            "id": "facility", "label": "Fasilitas dan infrastruktur",
            "keywords": ("fasilitas", "lab", "ruang", "wifi", "jaringan", "network", "kelas"),
            "prescription": "Audit kesiapan fasilitas sebelum delivery dan tetapkan checklist operasional harian.",
        },
        {
            "id": "instructor", "label": "Kualitas instruktur atau konsultan",
            "keywords": ("instruktur", "trainer", "konsultan", "mentor", "pengajar", "narasumber"),
            "prescription": "Perkuat coaching instruktur, review kompetensi domain, dan standardisasi evaluasi fasilitator.",
        },
        {
            "id": "material", "label": "Materi dan relevansi konten",
            "keywords": ("materi", "kurikulum", "modul", "silabus", "relevan", "contoh"),
            "prescription": "Review kurikulum per segmen, tambahkan contoh kontekstual, dan perbarui modul prioritas.",
        },
        {
            "id": "communication", "label": "Komunikasi dan koordinasi",
            "keywords": ("komunikasi", "informasi", "koordinasi", "brief", "update"),
            "prescription": "Rapikan alur komunikasi pra-delivery dan pastikan semua stakeholder menerima update status yang sama.",
        },
        {
            "id": "outcome", "label": "Dampak hasil layanan",
            "keywords": ("actionable", "implementasi", "hasil", "manfaat", "membantu", "sertifikasi"),
            "prescription": "Pertahankan praktik outcome review dan ubah testimoni hasil menjadi playbook layanan.",
        },
    )

    def __init__(self, dataframe):
        self.full_df = dataframe.copy() if dataframe is not None else pd.DataFrame()
        for column_name in CANONICAL_INTERNAL_COLUMNS:
            if column_name not in self.full_df.columns:
                self.full_df[column_name] = ""
        self.full_df = self.full_df.fillna("")
        if not self.full_df.empty:
            self.full_df["Rating Numeric"] = pd.to_numeric(
                self.full_df["Rating"],
                errors="coerce",
            )
            self.full_df["Sentiment Label"] = self.full_df["Rating Numeric"].apply(
                self._sentiment_label
            )
            self.full_df["Komentar Lower"] = self.full_df["Komentar"].astype(str).str.lower()

    @staticmethod
    def _sentiment_label(value):
        if pd.isna(value): return "unknown"
        if value >= 4: return "positive"
        if value <= 2: return "negative"
        return "neutral"

    @staticmethod
    def _safe_percentage(numerator, denominator):
        if not denominator: return 0.0
        return round((numerator / denominator) * 100, 1)

    @staticmethod
    def _truncate_text(text, max_length=180):
        clean_text = re.sub(r"\s+", " ", str(text)).strip()
        if len(clean_text) <= max_length: return clean_text
        return f"{clean_text[:max_length - 3]}..."

    @staticmethod
    def _series_counts(series, limit=5):
        filtered = series.fillna("").astype(str).str.strip()
        filtered = filtered[filtered != ""]
        return filtered.value_counts().head(limit)

    @staticmethod
    def _column_series(dataframe, column_name):
        if dataframe is None or column_name not in dataframe.columns:
            return pd.Series(dtype="object")
        return dataframe[column_name]

    def _series_counts_for_column(self, dataframe, column_name, limit=5):
        return self._series_counts(self._column_series(dataframe, column_name), limit=limit)

    @staticmethod
    def _label_from_options(options, option_id, fallback):
        for item in options:
            if item["id"] == option_id: return item["label"]
        return fallback

    @staticmethod
    def _clamp(value, minimum=0.0, maximum=100.0):
        return max(minimum, min(maximum, value))

    def _normalize_sentiment_filter(self, sentiment):
        valid_ids = {item["id"] for item in SENTIMENT_OPTIONS}
        return sentiment if sentiment in valid_ids else "all"

    def _normalize_score_engine(self, score_engine):
        return score_engine if score_engine in SCORE_ENGINE_PROFILES else DEFAULT_SCORE_ENGINE

    def _normalize_segment_filter(self, segment):
        cleaned = str(segment or "").strip()
        if not cleaned or cleaned.lower() == "all": return "all"
        available_segments = set(self.full_df["Tipe Stakeholder"].fillna("").astype(str).str.strip().tolist())
        return cleaned if cleaned in available_segments else "all"

    def _score_engine_profile(self, score_engine):
        normalized_engine = self._normalize_score_engine(score_engine)
        return SCORE_ENGINE_PROFILES.get(normalized_engine, SCORE_ENGINE_PROFILES[DEFAULT_SCORE_ENGINE])

    def _analysis_scope_text(self, timeframe, sentiment, segment, score_engine):
        sentiment_label = self._label_from_options(SENTIMENT_OPTIONS, sentiment, "Semua Sentimen")
        profile = self._score_engine_profile(score_engine)
        scope_parts = [f"periode {timeframe}", f"perspektif {profile['label']}"]
        if sentiment != "all": scope_parts.append(f"filter sentimen {sentiment_label.lower()}")
        if segment != "all": scope_parts.append(f"segmen {segment}")
        return ", ".join(scope_parts)

    def _forecast_horizon(self, timeframe):
        normalized = str(timeframe or "").lower()
        if "minggu" in normalized or "weekly" in normalized: return "1-2 minggu ke depan"
        if "semester" in normalized or "6 bulan" in normalized: return "semester berikutnya"
        if "tahun" in normalized or "year" in normalized: return "periode tahun berikutnya"
        if "bulan" in normalized or "monthly" in normalized: return "1-2 bulan ke depan"
        return "1-2 periode evaluasi berikutnya"

    @staticmethod
    def _format_month_year(value):
        month_names = ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
        return f"{month_names[value.month - 1]} {value.year}"

    def _forecast_calendar_reference(self, timeframe):
        now = datetime.now()
        normalized = str(timeframe or "").lower()
        if "minggu" in normalized or "weekly" in normalized:
            start, end = now + timedelta(days=7), now + timedelta(days=14)
            return f"sekitar {self._format_month_year(start)} sampai {self._format_month_year(end)}"
        if "bulan" in normalized or "monthly" in normalized:
            start, end = now + timedelta(days=30), now + timedelta(days=60)
            if start.year == end.year and start.month == end.month: return f"sekitar {self._format_month_year(start)}"
            return f"sekitar {self._format_month_year(start)} sampai {self._format_month_year(end)}"
        if "semester" in normalized or "6 bulan" in normalized or "biannual" in normalized:
            future = now + timedelta(days=180)
            return f"sekitar semester berikutnya pada {future.year}"
        if "tahun" in normalized or "yearly" in normalized:
            return f"pada tahun {now.year + 1}"
        future = now + timedelta(days=60)
        return f"sekitar {self._format_month_year(future)}"

    def _chart_pairs(self, series_counts, total_rows=None, limit=5, use_percentage=False):
        pairs = []
        for label, count in series_counts.head(limit).items():
            value = self._safe_percentage(count, total_rows) if use_percentage and total_rows else count
            pairs.append(f"{label},{value}")
        return "; ".join(pairs)

    def _filter_timeframe(self, timeframe):
        return self._filter_view(timeframe)

    def _filter_view(self, timeframe, sentiment="all", segment="all"):
        if self.full_df.empty: return self.full_df.copy()
        filtered = self.full_df[self.full_df["Rentang Waktu"] == timeframe].copy()
        normalized_sentiment = self._normalize_sentiment_filter(sentiment)
        normalized_segment = self._normalize_segment_filter(segment)
        if normalized_sentiment != "all": filtered = filtered[filtered["Sentiment Label"] == normalized_sentiment]
        if normalized_segment != "all": filtered = filtered[filtered["Tipe Stakeholder"].astype(str).str.strip() == normalized_segment]
        return filtered

    def _customer_journey_keywords(self):
        theme_lookup = {theme["id"]: theme for theme in self.THEME_LIBRARY}
        keyword_map = {}
        for stage in CUSTOMER_JOURNEY_STAGES:
            stage_keywords = []
            for theme_id in stage["theme_ids"]:
                stage_keywords.extend(theme_lookup.get(theme_id, {}).get("keywords", ()))
            keyword_map[stage["label"]] = tuple(dict.fromkeys(stage_keywords))
        return keyword_map

    def _attach_customer_journey(self, dataframe):
        if dataframe.empty:
            enriched = dataframe.copy()
            enriched["Customer Journey Stage"] = pd.Series(dtype="object")
            return enriched

        keyword_map = self._customer_journey_keywords()
        valid_stage_labels = {stage["label"] for stage in CUSTOMER_JOURNEY_STAGES}
        default_stage = next((stage["label"] for stage in CUSTOMER_JOURNEY_STAGES if stage["id"] == "delivery_experience"), CUSTOMER_JOURNEY_STAGES[0]["label"])

        def classify_stage(text):
            lowered = str(text or "").lower()
            best_stage, best_score = default_stage, 0
            for stage_label, keywords in keyword_map.items():
                score = sum(1 for keyword in keywords if keyword in lowered)
                if score > best_score: best_score, best_stage = score, stage_label
            return best_stage

        enriched = dataframe.copy()
        fallback_stages = enriched["Komentar Lower"].apply(classify_stage)
        if "Customer Journey Hint" in enriched.columns:
            hinted_stage = enriched["Customer Journey Hint"].fillna("").astype(str).str.strip()
            enriched["Customer Journey Stage"] = hinted_stage.where(hinted_stage.isin(valid_stage_labels), fallback_stages)
        else:
            enriched["Customer Journey Stage"] = fallback_stages
        return enriched

    def _customer_journey_rows(self, dataframe):
        if dataframe.empty: return []
        enriched = self._attach_customer_journey(dataframe)
        rows = []
        for stage in CUSTOMER_JOURNEY_STAGES:
            label = stage["label"]
            stage_df = enriched[enriched["Customer Journey Stage"] == label]
            if stage_df.empty: continue

            total = len(stage_df)
            positive_count = int((stage_df["Sentiment Label"] == "positive").sum())
            neutral_count = int((stage_df["Sentiment Label"] == "neutral").sum())
            negative_count = int((stage_df["Sentiment Label"] == "negative").sum())
            stage_theme_hits = self._theme_hits(stage_df)
            dominant_theme = next((theme["label"] for theme in stage_theme_hits if theme["negative_hits"] > 0 or theme["positive_hits"] > 0), "Sinyal umum customer journey")

            rows.append({
                "stage_id": stage["id"], "stage_label": label, "description": stage["description"], "volume": total,
                "average_rating": round(stage_df["Rating Numeric"].mean(), 2) if stage_df["Rating Numeric"].notna().any() else 0.0,
                "positive_share": self._safe_percentage(positive_count, total), "neutral_share": self._safe_percentage(neutral_count, total),
                "negative_share": self._safe_percentage(negative_count, total), "dominant_theme": dominant_theme,
            })
        rows.sort(key=lambda item: (item["negative_share"], item["volume"]), reverse=True)
        return rows

    def _score_engine_metrics(self, dataframe, score_engine):
        profile = self._score_engine_profile(score_engine)
        if dataframe.empty:
            return {"label": profile["label"], "current_score": 0.0, "projected_score": 0.0, "delta": 0.0, "direction": "stabil", "theme_rows": []}

        avg_rating = dataframe["Rating Numeric"].mean()
        base_score = ((avg_rating / 5) * 100) if pd.notna(avg_rating) else 0.0
        total_rows = len(dataframe)
        positive_share = self._safe_percentage(int((dataframe["Sentiment Label"] == "positive").sum()), total_rows)
        negative_share = self._safe_percentage(int((dataframe["Sentiment Label"] == "negative").sum()), total_rows)

        weighted_balance, weighted_positive_ratio, weighted_negative_ratio, total_weight = 0.0, 0.0, 0.0, 0.0
        theme_rows = []

        for theme in self._theme_hits(dataframe):
            weight = profile["theme_weights"].get(theme["id"], 0.35)
            total_hits = max(theme["total_hits"], 1)
            positive_ratio = theme["positive_hits"] / total_hits
            negative_ratio = theme["negative_hits"] / total_hits
            balance = positive_ratio - negative_ratio
            priority_score = round(((theme["negative_hits"] * 1.8) + theme["total_hits"]) * weight, 1)

            weighted_balance += balance * weight
            weighted_positive_ratio += positive_ratio * weight
            weighted_negative_ratio += negative_ratio * weight
            total_weight += weight

            theme_rows.append({
                "theme_id": theme["id"], "label": theme["label"], "weight": round(weight, 2), "total_hits": theme["total_hits"],
                "positive_hits": theme["positive_hits"], "negative_hits": theme["negative_hits"], "priority_score": priority_score, "prescription": theme["prescription"],
            })

        if total_weight > 0:
            weighted_balance /= total_weight; weighted_positive_ratio /= total_weight; weighted_negative_ratio /= total_weight

        current_score = self._clamp((base_score * 0.72) + ((50 + (weighted_balance * 50)) * 0.28))
        top_weighted_risk = max(((row["negative_hits"] / max(row["total_hits"], 1)) * row["weight"] for row in theme_rows), default=0.0)
        delta = round((((positive_share - negative_share) / 100) * 6) - (weighted_negative_ratio * 11) + (weighted_positive_ratio * 4) - (top_weighted_risk * 3), 1)
        if abs(delta) < 0.6: delta = 0.0
        projected_score = self._clamp(current_score + delta)
        direction = "naik" if delta > 0 else "turun" if delta < 0 else "stabil"

        theme_rows.sort(key=lambda item: (item["priority_score"], item["negative_hits"]), reverse=True)
        return {"label": profile["label"], "current_score": round(current_score, 1), "projected_score": round(projected_score, 1), "delta": delta, "direction": direction, "theme_rows": theme_rows}

    def _build_analysis_context(self, timeframe_df, timeframe, sentiment, segment, score_engine):
        normalized_sentiment = self._normalize_sentiment_filter(sentiment)
        normalized_segment = self._normalize_segment_filter(segment)
        normalized_score_engine = self._normalize_score_engine(score_engine)
        score_profile = self._score_engine_profile(normalized_score_engine)
        journey_rows = self._customer_journey_rows(timeframe_df)
        score_metrics = self._score_engine_metrics(timeframe_df, normalized_score_engine)
        dominant_journey = journey_rows[0] if journey_rows else None
        dominant_theme = score_metrics["theme_rows"][0] if score_metrics["theme_rows"] else None
        location_counts = self._series_counts_for_column(timeframe_df, "Lokasi", limit=5)
        instructor_type_counts = self._series_counts_for_column(timeframe_df, "Tipe Instruktur", limit=5)

        return {
            "timeframe": timeframe, "sentiment": normalized_sentiment,
            "sentiment_label": self._label_from_options(SENTIMENT_OPTIONS, normalized_sentiment, "Semua Sentimen"),
            "segment": normalized_segment, "segment_label": normalized_segment if normalized_segment != "all" else "Semua Segmen",
            "score_engine": normalized_score_engine, "score_profile": score_profile, "score_metrics": score_metrics,
            "journey_rows": journey_rows, "dominant_journey": dominant_journey, "dominant_theme": dominant_theme,
            "location_counts": location_counts, "instructor_type_counts": instructor_type_counts,
            "scope_text": self._analysis_scope_text(timeframe, normalized_sentiment, normalized_segment, normalized_score_engine),
            "horizon_text": self._forecast_horizon(timeframe),
        }

    def _theme_hits(self, dataframe):
        theme_stats = []
        if dataframe.empty: return theme_stats
        comment_series = dataframe["Komentar Lower"].astype(str)
        for theme in self.THEME_LIBRARY:
            match_mask = comment_series.apply(lambda text: any(keyword in text for keyword in theme["keywords"]))
            matched = dataframe[match_mask]
            if matched.empty: continue
            positive_hits = int((matched["Sentiment Label"] == "positive").sum())
            negative_hits = int((matched["Sentiment Label"] == "negative").sum())
            neutral_hits = int((matched["Sentiment Label"] == "neutral").sum())
            theme_stats.append({
                "id": theme["id"], "label": theme["label"], "prescription": theme["prescription"], "total_hits": int(len(matched)),
                "positive_hits": positive_hits, "negative_hits": negative_hits, "neutral_hits": neutral_hits, "matched_df": matched,
            })
        return sorted(theme_stats, key=lambda item: (item["negative_hits"], item["total_hits"]), reverse=True)

    def _quote_lines(self, dataframe, limit=3):
        if dataframe.empty: return ["- Tidak ada kutipan yang cukup untuk periode ini."]
        lines, seen_comments = [], set()
        for _, row in dataframe.iterrows():
            comment = self._truncate_text(row.get("Komentar", ""))
            if not comment or comment in seen_comments: continue
            seen_comments.add(comment)
            lines.append(f'- "{comment}" ({row.get("Tipe Stakeholder", "Stakeholder")} | {row.get("Layanan", "Layanan")} | rating {row.get("Rating", "-")})')
            if len(lines) >= limit: break
        return lines or ["- Tidak ada kutipan yang cukup untuk periode ini."]

    def _group_risk(self, dataframe, column_name, limit=3):
        if dataframe.empty or column_name not in dataframe.columns: return []
        rows = []
        grouped = dataframe.groupby(column_name, dropna=False)
        for label, group in grouped:
            clean_label = str(label).strip() or "Tidak terklasifikasi"
            rating_avg = group["Rating Numeric"].mean()
            negative_ratio = (group["Sentiment Label"] == "negative").mean()
            volume = len(group)
            safe_avg_rating = round(rating_avg, 2) if pd.notna(rating_avg) else 0.0
            risk_score = round((negative_ratio * 70) + ((5 - safe_avg_rating) * 6) + min(volume, 10), 1)
            rows.append({"label": clean_label, "volume": volume, "average_rating": safe_avg_rating, "negative_ratio": round(negative_ratio * 100, 1), "risk_score": risk_score})
        rows.sort(key=lambda item: item["risk_score"], reverse=True)
        return rows[:limit]

    def _governance_summary(self, timeframe_df):
        total_rows = len(timeframe_df)
        if total_rows == 0: return {"total_rows": 0, "completeness_pct": 0.0, "source_count": 0, "channel_count": 0}
        completeness_scores = [(timeframe_df[field].astype(str).str.strip() != "").mean() for field in ["Tipe Stakeholder", "Layanan", "Rentang Waktu", "Komentar"]]
        source_count = max(self._series_counts(timeframe_df["Sumber Feedback"], limit=20).shape[0], 1)
        channel_count = self._series_counts(timeframe_df["Kanal Feedback"], limit=20).shape[0]
        return {"total_rows": total_rows, "completeness_pct": round(sum(completeness_scores) / len(completeness_scores) * 100, 1), "source_count": source_count, "channel_count": channel_count}

    @staticmethod
    def _rating_assessment(avg_rating):
        if pd.isna(avg_rating): return "belum dapat dinilai secara memadai"
        if avg_rating >= 4.3: return "sangat baik dan relatif konsisten"
        if avg_rating >= 3.75: return "baik, tetapi masih menyisakan beberapa titik perbaikan"
        if avg_rating >= 3.0: return "cukup, namun belum cukup stabil untuk dianggap kuat"
        return "masih lemah dan memerlukan perhatian manajemen segera"

    @staticmethod
    def _negative_share_assessment(negative_share):
        if negative_share >= 35: return "cukup tinggi dan berpotensi mengganggu persepsi layanan jika tidak segera ditangani"
        if negative_share >= 20: return "perlu diawasi karena dapat berkembang menjadi isu yang lebih luas"
        if negative_share > 0: return "masih dalam batas terkendali, namun tetap membutuhkan pemantauan"
        return "belum menunjukkan sinyal keluhan yang berarti"

    @staticmethod
    def _risk_severity(risk_score):
        return "tinggi" if risk_score >= 55 else "menengah" if risk_score >= 40 else "terkendali"

    @staticmethod
    def _primary_label(series_counts, fallback):
        return series_counts.index[0] if not series_counts.empty else fallback

    @staticmethod
    def _format_count_summary(series_counts, unit="feedback", limit=3):
        if series_counts.empty: return "belum terpetakan"
        return ", ".join(f"{label} ({count} {unit})" for label, count in series_counts.head(limit).items())

    @staticmethod
    def _escape_table_cell(value):
        return str(value).replace("|", "\\|").replace("\n", " ").strip()

    @classmethod
    def _markdown_table(cls, headers, rows):
        if not rows: return ""
        header_line = "| " + " | ".join(cls._escape_table_cell(item) for item in headers) + " |"
        separator_line = "| " + " | ".join("---" for _ in headers) + " |"
        row_lines = ["| " + " | ".join(cls._escape_table_cell(cell) for cell in row) + " |" for row in rows]
        return "\n".join([header_line, separator_line, *row_lines])

    def _distribution_rows(self, series_counts, total_rows, limit=5):
        return [[label, count, f"{self._safe_percentage(count, total_rows)}%"] for label, count in series_counts.head(limit).items()]

    def _extract_osint_signals(self, macro_trends, limit=3):
        signals = []
        for line in str(macro_trends).splitlines():
            cleaned = line.strip()
            if not re.match(r"^\d+\.", cleaned): continue
            cleaned = re.sub(r"^\d+\.\s*", "", cleaned)
            parts = [part.strip() for part in cleaned.split(" | ") if part.strip()]
            if not parts: continue
            title = parts[0]
            snippet = parts[1] if len(parts) > 1 else ""
            source, date = "Tidak diketahui", "-"
            for part in parts[2:]:
                if part.startswith("sumber="): source = part.split("=", maxsplit=1)[1] or source
                elif part.startswith("tanggal="): date = part.split("=", maxsplit=1)[1] or date
            signals.append({"title": title, "snippet": snippet, "source": source, "date": date})
            if len(signals) >= limit: break
        return signals

    @staticmethod
    def _extract_deep_insight(macro_trends):
        match = re.search(r"\*\*Insight Mendalam[^*]*\*\*\s*(.*)", str(macro_trends))
        if match:
            return match.group(0)
        return ""

    @staticmethod
    def _theme_owner(theme_id):
        owner_map = {"responsiveness": "Customer Service / Account Management", "schedule": "Operations / Delivery Management", "facility": "Operations / General Affairs", "instructor": "Academic Lead / Service Quality", "material": "Academic Lead / Product Owner", "communication": "Customer Service / Project Coordinator", "outcome": "Service Owner / Quality Assurance"}
        return owner_map.get(theme_id, "Service Owner")

    @staticmethod
    def _theme_outcome(theme_id):
        outcome_map = {"responsiveness": "Waktu respons lebih konsisten dan penutupan isu lebih cepat.", "schedule": "Pengalaman delivery lebih tertata dan beban sesi lebih seimbang.", "facility": "Gangguan operasional di kelas atau sesi layanan dapat ditekan.", "instructor": "Konsistensi kualitas fasilitator meningkat di berbagai layanan.", "material": "Materi lebih relevan dengan kebutuhan peserta dan konteks klien.", "communication": "Ekspektasi stakeholder lebih selaras sejak pra-delivery hingga pasca-delivery.", "outcome": "Nilai manfaat layanan lebih mudah dirasakan dan dibuktikan."}
        return outcome_map.get(theme_id, "Persepsi kualitas layanan membaik secara terukur.")

    @staticmethod
    def _readiness_label(score):
        return "Kuat" if score >= 80 else "Cukup Siap" if score >= 60 else "Perlu Diperkuat" if score >= 40 else "Prioritas Tinggi"

    def _projection_sentence(self, context):
        metrics = context["score_metrics"]
        score_label = context["score_profile"]["forecast_label"]
        horizon_text = context["horizon_text"]
        calendar_reference = self._forecast_calendar_reference(context["timeframe"])

        if metrics["direction"] == "turun":
            return f"{score_label} diproyeksikan turun dari {metrics['current_score']} menjadi sekitar {metrics['projected_score']} dalam {horizon_text}, atau {calendar_reference}, apabila pola saat ini berlanjut."
        if metrics["direction"] == "naik":
            return f"{score_label} diproyeksikan naik dari {metrics['current_score']} menjadi sekitar {metrics['projected_score']} dalam {horizon_text}, atau {calendar_reference}, jika momentum yang ada dapat dipertahankan."
        return f"{score_label} diproyeksikan relatif stabil di kisaran {metrics['projected_score']} dalam {horizon_text}, atau {calendar_reference}, namun tetap perlu dipantau agar tidak bergeser ketika volume feedback bertambah."

    def _descriptive_markdown(self, timeframe_df, timeframe, notes, context):
        governance = self._governance_summary(timeframe_df)
        total_rows = governance["total_rows"]
        if total_rows == 0:
            return "## 1.1 Ringkasan Cakupan Feedback dan Tata Kelola\nTidak ada feedback internal yang sesuai dengan kombinasi filter yang dipilih untuk periode ini.\n"

        avg_rating = timeframe_df["Rating Numeric"].mean()
        positive_count = int((timeframe_df["Sentiment Label"] == "positive").sum())
        neutral_count = int((timeframe_df["Sentiment Label"] == "neutral").sum())
        negative_count = int((timeframe_df["Sentiment Label"] == "negative").sum())

        stakeholder_counts = self._series_counts(timeframe_df["Tipe Stakeholder"])
        service_counts = self._series_counts(timeframe_df["Layanan"])
        source_counts = self._series_counts(timeframe_df["Sumber Feedback"])
        channel_counts = self._series_counts(timeframe_df["Kanal Feedback"])

        top_sources = source_counts.index.tolist() if not source_counts.empty else ["Sumber internal terstandar"]
        top_channels = channel_counts.index.tolist() if not channel_counts.empty else ["Belum terpetakan"]
        positive_share = self._safe_percentage(positive_count, total_rows)
        neutral_share = self._safe_percentage(neutral_count, total_rows)
        negative_share = self._safe_percentage(negative_count, total_rows)
        score_metrics = context["score_metrics"]
        journey_rows = context["journey_rows"]
        scope_text = context["scope_text"]
        location_counts = context["location_counts"]
        instructor_type_counts = context["instructor_type_counts"]

        cleaned_notes = notes.strip().rstrip(".!?")
        focus_line = f"Fokus tambahan dari pengguna pada periode ini adalah: {cleaned_notes}." if notes and notes.strip() else "Tidak ada fokus tambahan dari pengguna, sehingga analisis dilakukan terhadap seluruh sinyal yang tersedia."
        governance_note = "Cakupan sumber sudah mulai terpetakan, tetapi pemetaan kanal masih perlu diperkuat." if governance["channel_count"] == 0 else "Pemetaan sumber dan kanal sudah tersedia sehingga jalur asal feedback lebih mudah diaudit."
        
        descriptive_intro = (
            f"Bagian ini menjelaskan kualitas dasar portofolio feedback yang menjadi fondasi laporan. Analisis dibaca pada {scope_text}. "
            f"Fokus pembacaannya menekankan {context['score_profile']['narrative_focus']}. "
            f"Pada periode {timeframe}, sistem memproses {total_rows} feedback tervalidasi dengan "
            f"rata-rata rating {round(avg_rating, 2) if pd.notna(avg_rating) else 0.0} dari 5, yang menunjukkan kinerja "
            f"layanan berada pada kategori {self._rating_assessment(avg_rating)}. "
            f"Komposisi sentimen memperlihatkan {positive_share}% sinyal positif, {neutral_share}% sinyal netral, "
            f"dan {negative_share}% sinyal negatif."
        )
        governance_intro = f"Dari sisi tata kelola, kelengkapan field inti mencapai {governance['completeness_pct']}%. Data berasal dari {governance['source_count']} sumber feedback dan {governance['channel_count']} kanal yang terpetakan. {governance_note} {focus_line}"
        indicator_table = self._markdown_table(
            ["Indikator", "Nilai"],
            [
                ["Periode analisis", timeframe], ["Cakupan analisis", scope_text], ["Total feedback tervalidasi", f"{total_rows} record"],
                ["Rata-rata rating", f"{round(avg_rating, 2) if pd.notna(avg_rating) else 0.0} dari 5"], [context["score_profile"]["label"], f"{score_metrics['current_score']} / 100"],
                ["Kelengkapan field inti", f"{governance['completeness_pct']}%"], ["Jumlah sumber feedback", governance["source_count"]], ["Jumlah kanal feedback", governance["channel_count"]],
            ],
        )
        score_table = self._markdown_table(
            ["Score Engine", "Nilai Saat Ini", "Arah Bacaan", "Tema Paling Berpengaruh"],
            [[context["score_profile"]["label"], f"{score_metrics['current_score']}", score_metrics["direction"].title(), context["dominant_theme"]["label"] if context["dominant_theme"] else "Belum terpetakan"]]
        )

        sentiment_chart_line = f"[[PIE: Komposisi Sentimen Feedback | Positif,{positive_share}; Netral,{neutral_share}; Negatif,{negative_share}]]"
        journey_chart_line = "[[CHART: Titik Customer Journey dengan Sinyal Negatif | Persentase Negatif | " + self._chart_pairs(pd.Series({item["stage_label"]: item["negative_share"] for item in journey_rows}), use_percentage=False, limit=4) + "]]" if journey_rows else ""
        sentiment_table = self._markdown_table(["Kategori Sentimen", "Jumlah", "Persentase"], [["Positif", positive_count, f"{positive_share}%"], ["Netral", neutral_count, f"{neutral_share}%"], ["Negatif", negative_count, f"{negative_share}%"]])
        stakeholder_table = self._markdown_table(["Segmen Stakeholder", "Jumlah Feedback", "Persentase"], self._distribution_rows(stakeholder_counts, total_rows, limit=5))
        service_table = self._markdown_table(["Layanan", "Jumlah Feedback", "Persentase"], self._distribution_rows(service_counts, total_rows, limit=5))
        location_table = self._markdown_table(["Lokasi Pelatihan", "Jumlah Feedback", "Persentase"], self._distribution_rows(location_counts, total_rows, limit=5))
        instructor_type_table = self._markdown_table(["Tipe Instruktur", "Jumlah Feedback", "Persentase"], self._distribution_rows(instructor_type_counts, total_rows, limit=5))
        
        location_pie_line = "[[PIE: Sebaran Lokasi Pelatihan | " + self._chart_pairs(location_counts, total_rows=total_rows, limit=5, use_percentage=True) + "]]" if not location_counts.empty else ""
        instructor_pie_line = "[[PIE: Komposisi Instruktur Internal vs OL | " + self._chart_pairs(instructor_type_counts, total_rows=total_rows, limit=5, use_percentage=True) + "]]" if not instructor_type_counts.empty else ""
        source_lines = [f"- Sumber utama: {', '.join(top_sources[:3])}", f"- Kanal utama: {', '.join(top_channels[:3])}"]
        
        distribution_paragraph = f"Sebaran volume feedback menunjukkan bahwa konsentrasi terbesar berasal dari segmen {self._format_count_summary(stakeholder_counts, limit=3)}. Dari sisi layanan, perhatian pengguna paling banyak tercurah pada {self._format_count_summary(service_counts, limit=3)}. Pola ini penting untuk dibaca secara hati-hati, karena volume tinggi belum otomatis berarti performa buruk, tetapi menandakan area yang paling banyak terekspos kepada pelanggan."
        source_paragraph = f"Dari sisi asal data, sumber yang paling dominan saat ini adalah {', '.join(top_sources[:3])}. Pada saat yang sama, kanal yang tercatat masih didominasi oleh {', '.join(top_channels[:3])}. Informasi ini perlu dibaca sebagai indikator awal representativitas data: semakin luas sumber dan kanal, semakin kuat dasar analisis untuk pengambilan keputusan lintas fungsi."
        delivery_context_paragraph = f"Lokasi pelatihan pada cakupan terpilih paling banyak berlangsung di {self._format_count_summary(location_counts, limit=3)}. Dari sisi tipe instruktur, komposisi saat ini didominasi oleh {self._format_count_summary(instructor_type_counts, limit=3)}. Informasi ini penting karena performa layanan sering kali dipengaruhi oleh kesiapan lokasi, format delivery, dan model pengajar yang dipakai."
        journey_table = self._markdown_table(["Tahap Customer Journey", "Volume", "Rating Rata-rata", "Positif", "Netral", "Negatif", "Tema Dominan"], [[item["stage_label"], item["volume"], item["average_rating"], f"{item['positive_share']}%", f"{item['neutral_share']}%", f"{item['negative_share']}%", item["dominant_theme"]] for item in journey_rows])
        dominant_journey_text = f"Sentimen paling menantang pada filter yang dipilih saat ini muncul pada tahap {context['dominant_journey']['stage_label']} dengan porsi sinyal negatif {context['dominant_journey']['negative_share']}%." if context["dominant_journey"] else "Belum ada tahap customer journey yang dapat dipetakan secara cukup kuat."

        return "\n".join([
            "## 1.1 Ringkasan Cakupan Feedback dan Tata Kelola", descriptive_intro, "", governance_intro, "", indicator_table, "",
            "## 1.2 Distribusi Sentimen, Rating, dan Volume", f"Distribusi sentimen menunjukkan bahwa proporsi sentimen negatif sebesar {negative_share}% {self._negative_share_assessment(negative_share)}. Sentimen positif tetap menjadi penopang utama pengalaman pelanggan, tetapi keberadaan sentimen netral yang cukup material mengindikasikan masih ada ruang untuk memperkuat pengalaman agar tidak berhenti pada persepsi 'cukup'.", "",
            score_table, "", sentiment_table, "", "Visual berikut memperlihatkan distribusi sentimen untuk kombinasi input yang dipilih, sehingga pembaca dapat segera melihat apakah pengalaman pelanggan lebih banyak berada di area positif, netral, atau negatif.", "", sentiment_chart_line, "",
            "## 1.3 Distribusi Stakeholder, Layanan, dan Kanal/Sumber", distribution_paragraph, "", "### Stakeholder dengan volume feedback terbesar", stakeholder_table, "", "### Layanan dengan volume feedback terbesar", service_table, "",
            "### Pemetaan sentimen pada customer journey", dominant_journey_text, "", journey_table, "", "Visual berikut membantu melihat tahapan customer journey mana yang paling banyak menampung sinyal negatif pada input yang dipilih.", "", journey_chart_line, "",
            "### Lokasi pelatihan dan tipe instruktur", delivery_context_paragraph, "", location_table, "", location_pie_line, "", instructor_type_table, "", instructor_pie_line, "",
            "### Cakupan sumber dan kanal", source_paragraph, "", *source_lines,
        ])

    def _diagnostic_markdown(self, timeframe_df, context):
        if timeframe_df.empty: return "## 2.1 Akar Masalah Utama dan Pain Point Dominan\nTidak ada feedback internal yang sesuai dengan kombinasi filter yang dipilih untuk periode ini.\n"

        theme_hits = self._theme_hits(timeframe_df)
        theme_lookup = {theme["id"]: theme for theme in theme_hits}
        prioritized_theme_rows = context["score_metrics"]["theme_rows"]
        prioritized_negative_ids = [item["theme_id"] for item in prioritized_theme_rows if theme_lookup.get(item["theme_id"], {}).get("negative_hits", 0) > 0][:3]
        negative_themes = [theme_lookup[theme_id] for theme_id in prioritized_negative_ids]
        if not negative_themes: negative_themes = [theme for theme in theme_hits if theme["negative_hits"] > 0][:3]
        positive_themes = sorted(theme_hits, key=lambda item: (item["positive_hits"], item["total_hits"]), reverse=True)[:3]

        if not negative_themes:
            negative_lines = ["- Belum ada pola keluhan dominan yang menonjol; mayoritas feedback berada pada area stabil."]
        else:
            negative_lines = []
            for theme in negative_themes:
                impacted_services = self._series_counts(theme["matched_df"]["Layanan"], limit=2)
                impacted_segments = self._series_counts(theme["matched_df"]["Tipe Stakeholder"], limit=2)
                negative_lines.append(f"- {theme['label']}: {theme['negative_hits']} sinyal negatif. Layanan terdampak: {', '.join(impacted_services.index.tolist()) or 'belum terpetakan'}. Segmen terdampak: {', '.join(impacted_segments.index.tolist()) or 'belum terpetakan'}.")

        positive_lines = []
        for theme in positive_themes:
            if theme["positive_hits"] <= 0: continue
            strongest_services = self._series_counts(theme["matched_df"]["Layanan"], limit=2)
            positive_lines.append(f"- {theme['label']}: {theme['positive_hits']} sinyal positif. Paling banyak muncul pada layanan {', '.join(strongest_services.index.tolist()) or 'belum terpetakan'}.")
        if not positive_lines: positive_lines = ["- Belum ada kekuatan yang cukup konsisten untuk dikonfirmasi pada periode ini."]

        negative_quotes = self._quote_lines(timeframe_df[timeframe_df["Sentiment Label"] == "negative"], limit=3)
        positive_quotes = self._quote_lines(timeframe_df[timeframe_df["Sentiment Label"] == "positive"], limit=2)

        service_risks = self._group_risk(timeframe_df, "Layanan", limit=5)
        location_risks = self._group_risk(timeframe_df, "Lokasi", limit=3)
        instructor_risks = self._group_risk(timeframe_df, "Tipe Instruktur", limit=3)
        process_gap_lines = [f"- {item['label']}: rata-rata rating {item['average_rating']}, proporsi negatif {item['negative_ratio']}%, volume {item['volume']}." for item in service_risks] or ["- Belum ada gap proses yang dapat dipetakan."]
        top_issue = negative_themes[0] if negative_themes else None
        top_strength = next((theme for theme in positive_themes if theme["positive_hits"] > 0 and (not top_issue or theme["id"] != top_issue["id"])), None) or next((theme for theme in positive_themes if theme["positive_hits"] > 0), None)
        dominant_journey = context["dominant_journey"]

        if top_issue and top_strength and top_issue["id"] == top_strength["id"]: strength_context = f"Menariknya, tema {top_strength['label']} muncul sebagai area yang terpolarisasi: sebagian pelanggan menilai sangat baik, sementara sebagian lain masih mengalami hambatan."
        elif top_strength: strength_context = f"Di sisi lain, kekuatan yang paling konsisten terlihat pada {top_strength['label']}."
        else: strength_context = "Kekuatan layanan belum muncul secara cukup konsisten untuk dijadikan diferensiasi yang kuat."

        diagnostic_intro = f"Analisis diagnostik bertujuan menjawab mengapa pola feedback pada periode ini muncul. Karena laporan dibaca dari sudut pandang {context['score_profile']['label']}, perhatian diagnosis terutama diarahkan ke {context['score_profile']['narrative_focus']}. {'Tema keluhan paling dominan saat ini adalah ' + top_issue['label'] + ', yang berulang pada beberapa komentar pelanggan.' if top_issue else 'Belum ada tema keluhan yang sangat dominan, sehingga pola masalah masih relatif tersebar.'} {strength_context}"
        journey_diagnostic = f"Jika dibaca menurut customer journey, titik gesekan yang paling terasa saat ini berada pada tahap {dominant_journey['stage_label']} dengan rating rata-rata {dominant_journey['average_rating']} dan porsi sentimen negatif {dominant_journey['negative_share']}%." if dominant_journey else "Pemetaan customer journey belum menunjukkan titik gesekan yang dominan."
        
        root_cause_table_rows = [[theme["label"], theme["negative_hits"], ", ".join(self._series_counts(theme["matched_df"]["Layanan"], limit=2).index.tolist()) or "Belum terpetakan", ", ".join(self._series_counts(theme["matched_df"]["Tipe Stakeholder"], limit=2).index.tolist()) or "Belum terpetakan"] for theme in negative_themes]
        root_cause_table = self._markdown_table(["Tema Prioritas", "Sinyal Negatif", "Layanan Dominan", "Segmen Dominan"], root_cause_table_rows)
        strength_table_rows = [[theme["label"], theme["positive_hits"], ", ".join(self._series_counts(theme["matched_df"]["Layanan"], limit=2).index.tolist()) or "Belum terpetakan"] for theme in positive_themes if theme["positive_hits"] > 0]
        strength_table = self._markdown_table(["Kekuatan", "Sinyal Positif", "Layanan Dominan"], strength_table_rows)
        service_risk_table = self._markdown_table(["Layanan", "Rata-rata Rating", "Proporsi Negatif", "Volume", "Skor Risiko"], [[item["label"], item["average_rating"], f"{item['negative_ratio']}%", item["volume"], item["risk_score"]] for item in service_risks])
        location_instructor_table = self._markdown_table(["Area Analisis", "Label", "Rata-rata Rating", "Proporsi Negatif", "Volume"], [[ "Lokasi", item["label"], item["average_rating"], f"{item['negative_ratio']}%", item["volume"]] for item in location_risks] + [["Tipe Instruktur", item["label"], item["average_rating"], f"{item['negative_ratio']}%", item["volume"]] for item in instructor_risks])
        operational_context = f"Dari sisi lokasi dan model instruktur, area yang perlu dicermati lebih dekat adalah {location_risks[0]['label'] if location_risks else 'lokasi yang belum terpetakan'} serta komposisi instruktur {instructor_risks[0]['label'] if instructor_risks else 'yang belum terpetakan'}. Pembacaan ini membantu membedakan apakah masalah lebih banyak terkait kesiapan tempat, model pengajar, atau memang tema layanan itu sendiri."

        return "\n".join([
            "## 2.1 Akar Masalah Utama dan Pain Point Dominan", diagnostic_intro, "", "Pembacaan akar masalah dilakukan dengan melihat pengulangan tema, dampaknya pada layanan, dan segmen pelanggan yang paling sering menyinggung isu serupa. Dengan pendekatan ini, tim manajemen dapat membedakan antara keluhan yang bersifat insidental dan keluhan yang sudah layak dibaca sebagai pola struktural.", "",
            root_cause_table, "", *negative_lines, "", "## 2.2 Kekuatan yang Konsisten dan Area yang Perlu Dijaga", "Selain keluhan, periode ini juga memperlihatkan area yang secara berulang diapresiasi oleh pelanggan. Bagian ini penting karena kekuatan yang konsisten dapat dijadikan acuan untuk standardisasi layanan, replikasi praktik baik, dan bahan komunikasi nilai kepada klien.", "",
            strength_table, "", *positive_lines, "", "## 2.3 Bukti Verbatim, Kesenjangan Proses, dan Segmentasi Masalah", "Bukti verbatim di bawah ini digunakan untuk menjaga agar interpretasi manajerial tetap berpijak pada suara pelanggan. Ringkasan kesenjangan proses membantu menerjemahkan komentar individual ke dalam area operasional yang dapat ditindaklanjuti.", "",
            journey_diagnostic, "", "### Kutipan keluhan representatif", *negative_quotes, "### Kutipan apresiasi representatif", *positive_quotes, "### Kesenjangan proses yang paling terlihat", service_risk_table, "", *process_gap_lines, "",
            "### Konteks lokasi pelatihan dan tipe instruktur", operational_context, "", location_instructor_table,
        ])

    def _predictive_markdown(self, timeframe_df, macro_trends, context):
        if timeframe_df.empty: return "## 3.1 Risiko Jangka Pendek Jika Pola Saat Ini Berlanjut\nTidak ada feedback internal yang sesuai dengan kombinasi filter yang dipilih untuk periode ini.\n"

        service_risks = self._group_risk(timeframe_df, "Layanan", limit=5)
        stakeholder_risks = self._group_risk(timeframe_df, "Tipe Stakeholder", limit=5)
        location_risks = self._group_risk(timeframe_df, "Lokasi", limit=3)
        instructor_risks = self._group_risk(timeframe_df, "Tipe Instruktur", limit=3)
        journey_rows = context["journey_rows"]
        score_metrics = context["score_metrics"]

        risk_lines = [f"- {item['label']} diperkirakan tetap menjadi area {self._risk_severity(item['risk_score'])} karena proporsi sinyal negatif {item['negative_ratio']}% dengan rata-rata rating {item['average_rating']}. Jika tidak ada intervensi, skor pengalaman untuk layanan ini cenderung berada di bawah rata-rata periode berjalan." for item in service_risks] or ["- Tidak ada risiko layanan yang cukup kuat untuk diproyeksikan pada periode ini."]
        segment_lines = [f"- Segmen {item['label']} perlu dipantau karena volume {item['volume']} feedback dengan proporsi negatif {item['negative_ratio']}%. Tanpa penanganan, persepsi mereka berpotensi lebih rendah pada periode evaluasi berikutnya." for item in stakeholder_risks] or ["- Tidak ada segmen pelanggan yang cukup dominan untuk diproyeksikan."]
        operational_lines = [f"- Lokasi {item['label']} perlu dipantau karena proporsi sinyal negatifnya {item['negative_ratio']}% dengan rating rata-rata {item['average_rating']}." for item in location_risks] + [f"- Komposisi instruktur {item['label']} juga perlu dibaca karena saat ini mencatat proporsi sinyal negatif {item['negative_ratio']}%." for item in instructor_risks] or ["- Belum ada sinyal lokasi atau tipe instruktur yang cukup kuat untuk diproyeksikan."]
        
        journey_lines = []
        for item in journey_rows[:3]:
            if item["negative_share"] >= 25: journey_lines.append(f"- Tahap {item['stage_label']} diperkirakan tetap menjadi titik gesekan utama karena porsi sentimen negatif masih {item['negative_share']}%.")
            elif item["positive_share"] >= 60: journey_lines.append(f"- Tahap {item['stage_label']} cenderung tetap menjadi area yang lebih kuat karena porsi sentimen positif mencapai {item['positive_share']}%.")
            else: journey_lines.append(f"- Tahap {item['stage_label']} diperkirakan relatif stabil, tetapi perlu dipantau karena sentimennya masih bercampur.")
        if not journey_lines: journey_lines = ["- Belum ada pembacaan customer journey yang cukup kuat untuk dijadikan proyeksi."]

        osint_signals = self._extract_osint_signals(macro_trends, limit=4)
        deep_insight = self._extract_deep_insight(macro_trends)
        osint_lines = []
        if deep_insight: osint_lines.append(f"- {deep_insight}")
        osint_lines.extend([f"- {signal['title']} ({signal['source']}, {signal['date']}): {signal['snippet']}" for signal in osint_signals])
        if not osint_lines: osint_lines = ["- Tren eksternal belum tersedia; prediksi saat ini sepenuhnya didasarkan pada data internal."]

        top_service_risk = service_risks[0] if service_risks else None
        top_segment_risk = stakeholder_risks[0] if stakeholder_risks else None
        predictive_intro = f"Analisis prediktif membaca risiko yang kemungkinan berkembang apabila pola feedback saat ini berlanjut dalam jangka pendek. {self._projection_sentence(context)} {'Layanan yang paling layak diprioritaskan untuk pengawasan adalah ' + top_service_risk['label'] + '.' if top_service_risk else 'Belum ada layanan dengan pola risiko yang cukup kuat untuk diprioritaskan.'} {'Segmen yang paling perlu dipantau adalah ' + top_segment_risk['label'] + '.' if top_segment_risk else 'Belum ada segmen dengan paparan risiko yang dominan.'}"
        score_projection_table = self._markdown_table(["Score Engine", "Nilai Saat Ini", "Arah Proyeksi", "Nilai Proyeksi", "Horizon", "Estimasi Waktu"], [[context["score_profile"]["label"], score_metrics["current_score"], score_metrics["direction"].title(), score_metrics["projected_score"], context["horizon_text"], self._forecast_calendar_reference(context["timeframe"])]])
        service_risk_table = self._markdown_table(["Layanan", "Level Risiko", "Rata-rata Rating", "Proporsi Negatif", "Volume"], [[item["label"], self._risk_severity(item["risk_score"]).title(), item["average_rating"], f"{item['negative_ratio']}%", item["volume"]] for item in service_risks])
        stakeholder_risk_table = self._markdown_table(["Segmen", "Level Risiko", "Rata-rata Rating", "Proporsi Negatif", "Volume"], [[item["label"], self._risk_severity(item["risk_score"]).title(), item["average_rating"], f"{item['negative_ratio']}%", item["volume"]] for item in stakeholder_risks])
        journey_projection_table = self._markdown_table(["Tahap Customer Journey", "Rating Rata-rata", "Negatif", "Positif", "Tema Dominan"], [[item["stage_label"], item["average_rating"], f"{item['negative_share']}%", f"{item['positive_share']}%", item["dominant_theme"]] for item in journey_rows])
        operational_projection_table = self._markdown_table(["Area Operasional", "Label", "Level Risiko", "Rata-rata Rating", "Proporsi Negatif"], [["Lokasi", item["label"], self._risk_severity(item["risk_score"]).title(), item["average_rating"], f"{item['negative_ratio']}%"] for item in location_risks] + [["Tipe Instruktur", item["label"], self._risk_severity(item["risk_score"]).title(), item["average_rating"], f"{item['negative_ratio']}%"] for item in instructor_risks])
        projection_chart_line = f"[[CHART: Perbandingan Score Saat Ini vs Proyeksi | Skor | Saat Ini,{score_metrics['current_score']}; Proyeksi,{score_metrics['projected_score']}]]"
        osint_table = self._markdown_table(["Sinyal Eksternal", "Sumber", "Tanggal"], [[signal["title"], signal["source"], signal["date"]] for signal in osint_signals])

        return "\n".join([
            "## 3.1 Risiko Jangka Pendek Jika Pola Saat Ini Berlanjut", predictive_intro, "", "Prediksi pada dokumen ini tidak dimaksudkan sebagai forecast statistik jangka panjang, melainkan sebagai early warning berbasis pola rating, proporsi sentimen negatif, dan konsentrasi volume feedback. Dengan pendekatan ini, manajemen dapat lebih cepat memutuskan layanan mana yang perlu ditangani lebih dahulu.", "",
            score_projection_table, "", projection_chart_line, "", service_risk_table, "", *risk_lines, "", "## 3.2 Prediksi Segmen dan Layanan yang Paling Rentan", "Selain layanan, pemantauan juga perlu diarahkan pada segmen pelanggan yang memperlihatkan kombinasi antara volume feedback tinggi dan kualitas pengalaman yang menurun. Segmen seperti ini biasanya lebih cepat mempengaruhi reputasi, retensi, dan peluang repeat engagement.", "",
            stakeholder_risk_table, "", *segment_lines, "", "### Pembacaan customer journey ke depan", journey_projection_table, "", *journey_lines, "", "### Area operasional yang perlu diawasi", operational_projection_table, "", *operational_lines, "",
            "## 3.3 Tren Eksternal yang Berpotensi Memperbesar Risiko", "Sinyal eksternal digunakan sebagai benchmark untuk membaca apakah tantangan yang muncul berasal murni dari kondisi internal atau juga diperkuat oleh perubahan ekspektasi pasar. Bila tren eksternal bergerak ke arah yang sama dengan keluhan pelanggan internal, maka urgensi intervensi meningkat.", "",
            osint_table, "", *osint_lines[:6],
        ])

    def _prescriptive_markdown(self, timeframe_df, context):
        if timeframe_df.empty: return "## 4.1 Intervensi Prioritas 30 Hari\nTidak ada feedback internal yang sesuai dengan kombinasi filter yang dipilih untuk periode ini.\n"

        theme_hits = {theme["id"]: theme for theme in self._theme_hits(timeframe_df)}
        prioritized_actions, prioritized_rows = [], []
        for score_theme in context["score_metrics"]["theme_rows"]:
            theme = theme_hits.get(score_theme["theme_id"])
            if not theme or theme["negative_hits"] <= 0: continue
            action_index = len(prioritized_actions) + 1
            prioritized_actions.append(f"{action_index}. {theme['label']}: {theme['prescription']}")
            prioritized_rows.append([action_index, theme["label"], theme["prescription"], self._theme_owner(theme["id"]), self._theme_outcome(theme["id"])])
            if len(prioritized_actions) >= 4: break

        if not prioritized_actions:
            prioritized_actions = ["1. Pertahankan monitoring mingguan karena belum ada pain point dominan yang membutuhkan intervensi besar."]
            prioritized_rows = [[1, "Monitoring berkala", "Pertahankan pemantauan mingguan dan lakukan review tren secara berkala.", "Quality Assurance / CX", "Risiko laten tetap termonitor meskipun belum ada isu dominan."]]

        governance_actions = ["1. Wajibkan field sumber feedback, kanal, stakeholder, layanan, tanggal, dan rating pada setiap record yang masuk.", "2. Satukan kontrak data antar sistem supaya analisis lintas sumber tetap konsisten dan dapat diaudit.", "3. Tetapkan SLA respon dan eskalasi untuk feedback negatif berprioritas tinggi."]
        roadmap_actions = ["1. Minggu 1: validasi kualitas data, pemetaan owner layanan, dan review pain point dominan.", "2. Minggu 2: jalankan quick wins pada layanan berisiko tertinggi serta aktifkan dashboard monitoring.", "3. Minggu 3-4: evaluasi dampak perbaikan, tutup feedback loop ke stakeholder, dan siapkan iterasi berikutnya.", "[[FLOW: Kumpulkan Feedback Multi-Sumber -> Normalisasi dan Audit Data -> Diagnosa Prioritas -> Jalankan Intervensi -> Evaluasi Dampak]]"]
        action_matrix = self._markdown_table(["Prioritas", "Fokus", "Tindakan", "Owner Utama", "Hasil yang Diharapkan"], prioritized_rows)
        roadmap_table = self._markdown_table(["Tahap", "Fokus Kerja", "Output yang Diharapkan"], [["Minggu 1", "Validasi kualitas data dan pemetaan owner layanan", "Daftar isu prioritas dan penanggung jawab yang disepakati."], ["Minggu 2", "Eksekusi quick wins pada layanan berisiko tertinggi", "Perbaikan cepat berjalan dan dashboard monitoring aktif."], ["Minggu 3-4", "Evaluasi dampak, penutupan feedback loop, dan iterasi", "Status dampak awal terdokumentasi dan rencana lanjutan tersusun."]])
        prescriptive_intro = f"Bagian preskriptif menerjemahkan temuan sebelumnya ke dalam tindakan yang dapat dibahas dan diputuskan dalam forum internal. Urutan prioritas disusun berdasarkan intensitas sinyal negatif, potensi dampak ke pengalaman pelanggan, dan kebutuhan koordinasi lintas fungsi dari sudut pandang {context['score_profile']['label']}."

        return "\n".join([
            "## 4.1 Intervensi Prioritas 30 Hari", prescriptive_intro, "", action_matrix, "", *prioritized_actions, "",
            "## 4.2 Penguatan Tata Kelola Feedback dan Eskalasi", "Selain quick wins layanan, perusahaan juga perlu memperkuat tata kelola feedback agar keputusan perbaikan berikutnya tidak selalu dimulai dari data yang parsial. Penguatan tata kelola akan menentukan kualitas diagnosis, kecepatan eskalasi, dan akuntabilitas tindak lanjut.", "", *governance_actions, "",
            "## 4.3 Rencana Tindak Lanjut Lintas Fungsi", "Rencana tindak lanjut di bawah ini disusun agar forum internal tidak berhenti pada pembacaan laporan, tetapi langsung bergerak ke tahap eksekusi. Timeline dapat disesuaikan, namun disiplin implementasi antar fungsi tetap menjadi faktor penentu keberhasilan.", "", roadmap_table, "", *roadmap_actions,
        ])

    def _implementation_readiness_markdown(self, timeframe_df, timeframe, notes, macro_trends, context):
        if timeframe_df.empty: return "## 5.1 Prioritas Sasaran Bisnis\nTidak ada feedback internal yang sesuai dengan kombinasi filter yang dipilih untuk periode ini.\n"

        total_rows = len(timeframe_df)
        avg_rating = timeframe_df["Rating Numeric"].mean()
        positive_count = int((timeframe_df["Sentiment Label"] == "positive").sum())
        negative_count = int((timeframe_df["Sentiment Label"] == "negative").sum())
        positive_share = self._safe_percentage(positive_count, total_rows)
        negative_share = self._safe_percentage(negative_count, total_rows)
        governance = self._governance_summary(timeframe_df)
        theme_hits = self._theme_hits(timeframe_df)
        service_risks = self._group_risk(timeframe_df, "Layanan", limit=3)
        stakeholder_risks = self._group_risk(timeframe_df, "Tipe Stakeholder", limit=3)
        top_service = service_risks[0] if service_risks else None
        top_segment = stakeholder_risks[0] if stakeholder_risks else None
        top_issue = next((theme for theme in theme_hits if theme["negative_hits"] > 0), None)
        top_strength = next((theme for theme in theme_hits if theme["positive_hits"] > 0 and (not top_issue or theme["id"] != top_issue["id"])), None) or next((theme for theme in theme_hits if theme["positive_hits"] > 0), None)
        
        osint_signals = self._extract_osint_signals(macro_trends, limit=2)
        deep_insight = self._extract_deep_insight(macro_trends)
        
        focus_text = notes.strip().rstrip(".!?") if notes and notes.strip() else "Tidak ada fokus tambahan dari pengguna"
        top_service_name = top_service["label"] if top_service else self._primary_label(self._series_counts(timeframe_df["Layanan"], limit=1), "layanan prioritas")
        top_segment_name = top_segment["label"] if top_segment else self._primary_label(self._series_counts(timeframe_df["Tipe Stakeholder"], limit=1), "segmen utama")

        business_score = min(100, 50 + min(total_rows, 10) * 3 + (10 if top_service else 0) + (10 if top_issue else 0))
        data_score = min(100, int(governance["completeness_pct"] * 0.45) + min(governance["source_count"], 3) * 10 + min(governance["channel_count"], 3) * 10 + (10 if governance["source_count"] >= 2 else 0) + (10 if governance["channel_count"] >= 1 else 0))
        architecture_score = min(100, 35 + min(governance["source_count"], 3) * 8 + min(governance["channel_count"], 3) * 10 + (10 if governance["source_count"] >= 2 else 0) + (10 if total_rows >= 10 else 5 if total_rows >= 5 else 0))
        people_score = max(35, min(100, 60 + (10 if top_strength else 0) - (10 if top_issue and top_issue["id"] in {"instructor", "communication"} else 0) - (5 if top_issue and top_issue["id"] in {"responsiveness", "schedule"} else 0)))
        governance_score = min(100, int(governance["completeness_pct"] * 0.35) + min(governance["source_count"], 3) * 8 + min(governance["channel_count"], 3) * 12 + (10 if governance["source_count"] >= 2 else 0) + (10 if governance["channel_count"] >= 1 else 0))
        culture_score = max(35, min(100, 45 + int(positive_share * 0.4) - int(negative_share * 0.3) + (10 if total_rows >= 5 else 0)))

        pillar_map = {
            "business_use_case": {
                "score": business_score,
                "reading": f"Prioritas implementasi yang paling konkret saat ini adalah memperkuat tata kelola feedback untuk mendeteksi lebih dini risiko pada layanan {top_service_name} dan memantau pengalaman segmen {top_segment_name}. Dengan rata-rata rating {round(avg_rating, 2) if pd.notna(avg_rating) else 0.0} dan sinyal negatif {negative_share}%, laporan ini sudah memiliki dasar yang cukup jelas untuk diterjemahkan ke keputusan bisnis. Fokus use case saat ini dibaca pada {context['scope_text']}.",
                "implication": "Inisiatif ini sebaiknya tidak diposisikan sebagai eksperimen AI yang abstrak, melainkan sebagai use case terukur untuk mempercepat diagnosis masalah, prioritisasi perbaikan, dan evaluasi dampak manajerial.",
                "actions": [f"Tentukan satu sasaran bisnis utama, misalnya menurunkan sinyal negatif pada {top_service_name} atau mempercepat penutupan isu pelanggan.", "Sepakati KPI pilot 30 hari yang mudah diukur, misalnya waktu respons, penurunan keluhan berulang, atau peningkatan kepuasan.", f"Gunakan fokus rapat pada area berikut: {focus_text}."],
            },
            "data_model_foundation": {
                "score": data_score,
                "reading": f"Fondasi data saat ini cukup untuk memulai pilot karena kelengkapan field inti mencapai {governance['completeness_pct']}%. Namun cakupan masih berasal dari {governance['source_count']} sumber dan {governance['channel_count']} kanal yang terpetakan, sehingga representativitas lintas kanal belum sepenuhnya kuat.",
                "implication": "Pilot dapat berjalan sekarang, tetapi scale-up akan sulit bila kontrak data, owner data, dan standar mandatory field belum disepakati bersama.",
                "actions": ["Tetapkan owner data untuk setiap sumber feedback dan definisikan field yang wajib terisi.", "Pastikan kanal, tanggal, stakeholder, layanan, rating, dan komentar selalu tercatat secara konsisten.", "Gunakan model analitik saat ini sebagai baseline, lalu perluas sumber data secara bertahap setelah kualitas data stabil."],
            },
            "infrastructure_architecture": {
                "score": architecture_score,
                "reading": "Implementasi untuk use case ini tidak harus dimulai dari arsitektur yang mahal. Kebutuhan dekatnya adalah arsitektur yang aman, dapat dibagikan secara internal, mendukung read-only integration, dan cukup mudah diperluas ketika sumber data bertambah.",
                "implication": "Keputusan cloud, on-prem, atau hybrid sebaiknya mengikuti kebutuhan compliance perusahaan. Untuk tahap saat ini, prioritasnya adalah kestabilan deployment internal, logging, health check, dan jalur ingest data yang dapat diaudit.",
                "actions": ["Mulai dari shared internal deployment yang stabil sebelum memikirkan scale penuh.", "Gunakan akses API read-only untuk data internal dan pisahkan konfigurasi pilot dari produksi.", "Siapkan jalur scale bertahap ke cloud, on-prem, atau hybrid sesuai regulasi dan kebijakan keamanan data."],
            },
            "people_capability": {
                "score": people_score,
                "reading": f"Temuan periode ini menunjukkan bahwa inisiatif ini tidak dapat menjadi urusan IT saja. {'Tema utama yang muncul adalah ' + top_issue['label'] + ', sehingga interpretasi bisnis perlu melibatkan owner layanan.' if top_issue else 'Interpretasi insight tetap membutuhkan owner layanan dan pihak operasional.'} {'Kekuatan yang layak dijaga terlihat pada ' + top_strength['label'] + '.' if top_strength else 'Belum ada kekuatan dominan yang cukup untuk dijadikan patokan lintas tim.'}",
                "implication": "Agar pilot menghasilkan keputusan nyata, perusahaan perlu membedakan peran tim teknis, QA/CX, business translator, owner layanan, dan eksekutor perbaikan di lapangan.",
                "actions": ["Tunjuk business owner yang bertanggung jawab atas use case dan outcome pilot.", "Pastikan QA/CX, owner layanan, dan operasional ikut mereview laporan, bukan hanya tim teknis.", "Siapkan satu business translator yang menerjemahkan insight teknis menjadi keputusan manajerial."],
            },
            "governance": {
                "score": governance_score,
                "reading": "Kontrol risiko dan tata kelola menjadi area penting karena laporan ini memadukan data internal dan konteks eksternal. Internal data harus tetap menjadi sumber kebenaran untuk fakta operasional, sedangkan OSINT dipakai hanya sebagai benchmark dan konteks pasar.",
                "implication": "Organisasi perlu mendefinisikan dengan tegas data apa yang boleh dipakai, siapa yang berhak mengaksesnya, kapan rekomendasi boleh dijalankan, dan kapan hasil AI harus dihentikan atau dikoreksi secara manual.",
                "actions": ["Tetapkan SOP penggunaan data internal vs OSINT agar tidak terjadi pencampuran fakta dan konteks publik.", "Buat review cadence dan approval gate untuk insight berisiko tinggi sebelum menjadi keputusan resmi.", "Dokumentasikan risk control, quality check, dan batas penggunaan AI pada forum evaluasi internal."],
            },
            "culture": {
                "score": culture_score,
                "reading": f"Budaya organisasi untuk inisiatif ini sebaiknya dibangun dengan semangat mencoba secara terstruktur. Komposisi sentimen positif {positive_share}% menunjukkan ada modal kepercayaan yang cukup untuk memulai, sementara keberadaan sinyal negatif tetap penting sebagai bahan belajar dan perbaikan.",
                "implication": "Keberhasilan tahap awal tidak harus berarti sistem langsung sempurna. Yang lebih penting adalah organisasi punya kebiasaan mencoba, mengevaluasi, mengambil pelajaran, dan memutuskan langkah berikutnya secara disiplin.",
                "actions": ["Posisikan pilot sebagai ruang belajar terstruktur, bukan proyek yang harus sempurna sejak hari pertama.", "Dokumentasikan apa yang berhasil, apa yang belum, dan keputusan apa yang diambil setelah setiap periode evaluasi.", "Gunakan laporan ini sebagai alat diskusi lintas fungsi agar AI adoption menjadi perubahan cara kerja, bukan hanya tambahan tools."],
            },
        }

        summary_rows, pillar_sections = [], []
        for pillar in ADOPTION_READINESS_PILLARS:
            pillar_data = pillar_map[pillar["id"]]
            status = self._readiness_label(pillar_data["score"])
            summary_rows.append([" ".join(pillar["title"].split(" ")[1:]), status, pillar_data["reading"], pillar["guiding_question"]])
            pillar_sections.extend([f"## {pillar['title']}", pillar_data["reading"], "", f"Status kesiapan saat ini: **{status}**.", "", f"Pertanyaan pemandu: {pillar['guiding_question']}", "", "### Implikasi untuk Pengambilan Keputusan", pillar_data["implication"], "", "### Aksi Prioritas", *[f"- {action}" for action in pillar_data["actions"]], ""])

        if deep_insight and osint_signals: osint_note = f"{deep_insight} Sinyal eksternal lainnya yang relevan antara lain {osint_signals[0]['title']} dari {osint_signals[0]['source']}."
        elif osint_signals: osint_note = f"Sinyal eksternal yang paling relevan saat ini antara lain {osint_signals[0]['title']} dari {osint_signals[0]['source']}."
        else: osint_note = "Sinyal eksternal belum tersedia, sehingga pembacaan kesiapan implementasi terutama bersandar pada data internal."
        summary_table = self._markdown_table(["Area", "Status", "Pembacaan Saat Ini", "Pertanyaan Diskusi"], summary_rows)

        return "\n".join([
            "Bagian ini menerjemahkan hasil feedback intelligence ke dalam pertimbangan implementasi dan penguatan organisasi agar perusahaan tidak berhenti pada insight, tetapi bergerak menuju eksekusi yang terstruktur. Prinsipnya adalah memulai dari use case yang nyata, membangun fondasi secara bertahap, lalu belajar secara disiplin dari pilot yang dijalankan.", "",
            f"Untuk periode {timeframe}, pertimbangan implementasi perlu dilihat bersama konteks berikut: {osint_note} Analisis saat ini dibaca menggunakan {context['score_profile']['label']} dengan fokus pada {context['score_profile']['narrative_focus']}. Dengan demikian, forum internal dapat menilai bukan hanya apa yang harus diperbaiki, tetapi juga seberapa siap organisasi untuk menjalankan inisiatif ini secara lebih sistematis.", "",
            summary_table, "", *pillar_sections,
        ])

    def build_report_sections(self, timeframe, notes, macro_trends, sentiment="all", segment="all", score_engine=DEFAULT_SCORE_ENGINE):
        timeframe_df = self._filter_view(timeframe, sentiment=sentiment, segment=segment)
        context = self._build_analysis_context(timeframe_df, timeframe, sentiment, segment, score_engine)
        section_map = {
            "cx_chap_1": self._descriptive_markdown(timeframe_df, timeframe, notes, context),
            "cx_chap_2": self._diagnostic_markdown(timeframe_df, context),
            "cx_chap_3": self._predictive_markdown(timeframe_df, macro_trends, context),
            "cx_chap_4": self._prescriptive_markdown(timeframe_df, context),
            "cx_chap_5": self._implementation_readiness_markdown(timeframe_df, timeframe, notes, macro_trends, context),
        }
        return [{"id": chapter["id"], "title": chapter["title"], "content": section_map.get(chapter["id"], "")} for chapter in CX_SENTIMENT_STRUCTURE]

    def build_executive_snapshot(self, timeframe, notes="", sentiment="all", segment="all", score_engine=DEFAULT_SCORE_ENGINE):
        timeframe_df = self._filter_view(timeframe, sentiment=sentiment, segment=segment)
        if timeframe_df.empty: return "## Ringkasan Eksekutif\n- Tidak ada data internal yang cukup untuk menyusun snapshot eksekutif pada kombinasi filter yang dipilih.\n"

        context = self._build_analysis_context(timeframe_df, timeframe, sentiment, segment, score_engine)
        total_rows = len(timeframe_df)
        avg_rating = timeframe_df["Rating Numeric"].mean()
        negative_count = int((timeframe_df["Sentiment Label"] == "negative").sum())
        negative_share = self._safe_percentage(negative_count, total_rows)
        top_service = self._series_counts(timeframe_df["Layanan"], limit=1)
        top_stakeholder = self._series_counts(timeframe_df["Tipe Stakeholder"], limit=1)
        top_risk = self._group_risk(timeframe_df, "Layanan", limit=1)
        governance = self._governance_summary(timeframe_df)
        top_issue = next((theme for theme in self._theme_hits(timeframe_df) if theme["negative_hits"] > 0), None)
        focus_text = notes.strip() if notes and notes.strip() else "Tidak ada fokus tambahan dari pengguna."
        dominant_journey, score_metrics = context["dominant_journey"], context["score_metrics"]
        top_location = self._series_counts_for_column(timeframe_df, "Lokasi", limit=1)
        top_instructor_type = self._series_counts_for_column(timeframe_df, "Tipe Instruktur", limit=1)

        risk_statement = f"- Risiko teratas saat ini ada pada layanan {top_risk[0]['label']} dengan proporsi sinyal negatif {top_risk[0]['negative_ratio']}%." if top_risk else "- Belum ada layanan dengan risiko dominan yang teridentifikasi."
        executive_intro = f"Laporan ini merangkum kondisi pengalaman pelanggan untuk periode {timeframe} berdasarkan {total_rows} feedback tervalidasi. Analisis saat ini dibaca pada {context['scope_text']} dengan fokus pada {context['score_profile']['narrative_focus']}. Secara umum, rata-rata rating berada pada level {round(avg_rating, 2) if pd.notna(avg_rating) else 0.0} dari 5, yang menunjukkan kualitas layanan {self._rating_assessment(avg_rating)}. Proporsi sentimen negatif tercatat sebesar {negative_share}% ({negative_count} feedback), sehingga kondisi ini {self._negative_share_assessment(negative_share)}."
        meeting_context = f"Untuk kebutuhan rapat internal, perhatian utama sebaiknya diarahkan pada layanan {top_risk[0]['label'] if top_risk else self._primary_label(top_service, 'yang memiliki volume feedback terbesar')} serta pada isu {top_issue['label'] if top_issue else 'konsistensi kualitas layanan'}. {self._projection_sentence(context)} Fokus tambahan yang diminta pengguna: {focus_text}"
        snapshot_table = self._markdown_table(["Indikator Kunci", "Nilai"], [["Total feedback dianalisis", f"{total_rows} record"], ["Cakupan analisis", context["scope_text"]], [context["score_profile"]["label"], f"{score_metrics['current_score']} / 100"], ["Rata-rata rating", f"{round(avg_rating, 2) if pd.notna(avg_rating) else 0.0} dari 5"], ["Proporsi sentimen negatif", f"{negative_share}%"], ["Layanan dengan volume terbesar", self._primary_label(top_service, "Belum terpetakan")], ["Segmen dengan volume terbesar", self._primary_label(top_stakeholder, "Belum terpetakan")], ["Lokasi pelatihan dominan", self._primary_label(top_location, "Belum terpetakan")], ["Tipe instruktur dominan", self._primary_label(top_instructor_type, "Belum terpetakan")], ["Kelengkapan field inti", f"{governance['completeness_pct']}%"]])
        meeting_agenda = [f"- Apakah layanan {top_risk[0]['label']} memerlukan intervensi prioritas lintas fungsi pada 30 hari ke depan?" if top_risk else "- Apakah perusahaan perlu memperluas pengumpulan feedback agar risiko layanan lebih mudah dibaca?", f"- Bagaimana tindak lanjut yang paling tepat untuk tema {top_issue['label']} agar tidak berkembang menjadi keluhan berulang?" if top_issue else "- Kekuatan layanan mana yang paling layak distandardisasi dan direplikasi?", f"- Tahap customer journey mana yang paling perlu dikoreksi lebih dulu, mengingat titik gesekan terbesar saat ini berada pada {dominant_journey['stage_label']}?" if dominant_journey else "- Tahap customer journey mana yang paling perlu dipetakan lebih rinci pada periode berikutnya?", "- Apakah tata kelola sumber, kanal, dan owner tindak lanjut sudah cukup jelas untuk mendukung evaluasi periodik berikutnya?"]

        return "\n".join([
            "## Ringkasan Eksekutif", executive_intro, "", meeting_context, "", snapshot_table, "", "### Agenda Diskusi Prioritas", *meeting_agenda, "",
            "### Poin Utama untuk Pembacaan Cepat", f"- Total feedback yang dianalisis: {total_rows} record.", f"- Rata-rata rating periode ini: {round(avg_rating, 2) if pd.notna(avg_rating) else 0.0} dari 5.", f"- Volume layanan terbesar: {top_service.index[0] if not top_service.empty else 'Belum terpetakan'}.", f"- Segmen dengan volume terbesar: {top_stakeholder.index[0] if not top_stakeholder.empty else 'Belum terpetakan'}.", f"- Lokasi pelatihan dominan: {top_location.index[0] if not top_location.empty else 'Belum terpetakan'}.", f"- Tipe instruktur dominan: {top_instructor_type.index[0] if not top_instructor_type.empty else 'Belum terpetakan'}.", f"- Proporsi sentimen negatif: {negative_share}%.", f"- {context['score_profile']['label']} saat ini: {score_metrics['current_score']} dengan proyeksi {score_metrics['direction']} ke {score_metrics['projected_score']}.", f"- Tahap customer journey yang paling perlu diperhatikan: {dominant_journey['stage_label']}." if dominant_journey else "- Pemetaan customer journey belum menunjukkan titik perhatian yang dominan.", risk_statement, "- Struktur laporan ini disusun untuk mendukung analisis Descriptive, Diagnostic, Predictive, Prescriptive, serta kesiapan implementasi dan penguatan organisasi secara konsisten.",
        ])


class StyleEngine:
    @staticmethod
    def _configure_text_style(style, font_name, font_size, color, bold=False):
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor(*color)
        style.font.bold = bold

    @staticmethod
    def apply_document_styles(doc, theme_color):
        for section in doc.sections:
            section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = Cm(2.54), Cm(2.54), Cm(2.54), Cm(2.54)
            footer_paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
            footer_paragraph.clear()
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            footer_run = footer_paragraph.add_run("STRICTLY CONFIDENTIAL | Inixindo Jogja Executive Report | Page ")
            footer_run.font.name, footer_run.font.size, footer_run.font.color.rgb = "Calibri", Pt(9), RGBColor(128, 128, 128)
            append_field(footer_paragraph, "PAGE")

        normal_style = doc.styles["Normal"]
        StyleEngine._configure_text_style(normal_style, "Calibri", 11, (33, 37, 41))
        normal_format = normal_style.paragraph_format
        normal_format.alignment, normal_format.line_spacing_rule, normal_format.line_spacing, normal_format.space_after = WD_ALIGN_PARAGRAPH.JUSTIFY, WD_LINE_SPACING.MULTIPLE, 1.15, Pt(8)

        heading_1 = doc.styles["Heading 1"]
        StyleEngine._configure_text_style(heading_1, "Calibri", 16, theme_color, True)
        heading_1.paragraph_format.space_before, heading_1.paragraph_format.space_after = Pt(18), Pt(8)

        heading_2 = doc.styles["Heading 2"]
        StyleEngine._configure_text_style(heading_2, "Calibri", 13, (0, 0, 0), True)
        heading_2.paragraph_format.space_before, heading_2.paragraph_format.space_after = Pt(14), Pt(6)

        heading_3 = doc.styles["Heading 3"]
        StyleEngine._configure_text_style(heading_3, "Calibri", 12, (54, 54, 54), True)
        heading_3.paragraph_format.space_before, heading_3.paragraph_format.space_after = Pt(10), Pt(4)

        for list_style_name in ("List Bullet", "List Bullet 2", "List Bullet 3", "List Number", "List Number 2", "List Number 3"):
            if list_style_name in doc.styles:
                list_style = doc.styles[list_style_name]
                StyleEngine._configure_text_style(list_style, "Calibri", 11, (33, 37, 41))
                list_style.paragraph_format.space_after = Pt(4)

class ChartEngine:
    @staticmethod
    def _get_plt_color(theme_color): return tuple(channel / 255 for channel in theme_color)

    @staticmethod
    def _parse_chart_points(raw_data):
        labels, values = [], []
        for pair in raw_data.split(";"):
            if "," not in pair: continue
            label, value = pair.split(",", maxsplit=1)
            cleaned_value = re.sub(r"[^\d.]", "", value)
            if cleaned_value:
                labels.append(label.strip())
                values.append(float(cleaned_value))
        return labels, values

    @staticmethod
    def create_bar_chart(data_str, theme_color):
        try:
            parts = data_str.split("|")
            title_str, ylabel_str, raw_data = [p.strip() for p in parts] if len(parts) == 3 else ("Sentimen Makro", "Persentase", data_str)
            labels, values = ChartEngine._parse_chart_points(raw_data)
            if not labels: return None

            fig, axis = plt.subplots(figsize=(7, 4.5))
            axis.bar(labels, values, color=ChartEngine._get_plt_color(theme_color), alpha=0.9, width=0.5)
            axis.set_title(title_str, fontsize=12, fontweight="bold", pad=20)
            axis.set_ylabel(ylabel_str, fontsize=10)
            axis.spines["top"].set_visible(False)
            axis.spines["right"].set_visible(False)

            image_stream = io.BytesIO()
            plt.savefig(image_stream, format="png", bbox_inches="tight", dpi=150)
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception as exc:
            logger.warning("Gagal membuat bar chart: %s", exc)
            return None

    @staticmethod
    def create_pie_chart(data_str, theme_color):
        try:
            parts = data_str.split("|")
            title_str, raw_data = [p.strip() for p in parts] if len(parts) == 2 else ("Distribusi", data_str)
            labels, values = ChartEngine._parse_chart_points(raw_data)
            if not labels: return None

            palette = [ChartEngine._get_plt_color(theme_color), (0.85, 0.35, 0.25), (0.20, 0.45, 0.70), (0.35, 0.65, 0.45), (0.75, 0.60, 0.25)]
            fig, axis = plt.subplots(figsize=(6.2, 4.8))
            axis.pie(values, labels=labels, autopct="%1.1f%%", startangle=90, colors=palette[: len(labels)], textprops={"fontsize": 9})
            axis.set_title(title_str, fontsize=12, fontweight="bold", pad=16)
            axis.axis("equal")

            image_stream = io.BytesIO()
            plt.savefig(image_stream, format="png", bbox_inches="tight", dpi=150)
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception as exc:
            logger.warning("Gagal membuat pie chart: %s", exc)
            return None

    @staticmethod
    def create_flowchart(data_str, theme_color):
        try:
            steps = ["\n".join(textwrap.wrap(step.strip(), width=18)) for step in data_str.split("->") if step.strip()]
            if len(steps) < 2: return None

            fig, axis = plt.subplots(figsize=(8, 3))
            axis.axis("off")
            x_positions = [index * 2.5 for index in range(len(steps))]

            for index in range(len(steps) - 1):
                axis.annotate("", xy=(x_positions[index + 1] - 1.0, 0.5), xytext=(x_positions[index] + 1.0, 0.5), arrowprops={"arrowstyle": "-|>", "lw": 1.5})

            for index, step in enumerate(steps):
                box = patches.FancyBboxPatch((x_positions[index] - 1.0, 0.1), 2.0, 0.8, boxstyle="round,pad=0.1", fc=ChartEngine._get_plt_color(theme_color), alpha=0.9)
                axis.add_patch(box)
                axis.text(x_positions[index], 0.5, step, ha="center", va="center", size=9, color="white", fontweight="bold")

            axis.set_xlim(-1.2, (len(steps) - 1) * 2.5 + 1.2)
            axis.set_ylim(0, 1)

            image_stream = io.BytesIO()
            plt.savefig(image_stream, format="png", bbox_inches="tight", dpi=200, transparent=True)
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception as exc:
            logger.warning("Gagal membuat flowchart: %s", exc)
            return None

class DocumentBuilder:
    LIST_STYLES = {"ul": ["List Bullet", "List Bullet 2", "List Bullet 3"], "ol": ["List Number", "List Number 2", "List Number 3"]}

    @staticmethod
    def _resolve_list_style(doc, list_tag, level):
        style_candidates = DocumentBuilder.LIST_STYLES.get(list_tag, ["List Bullet"])
        return style_candidates[min(level, len(style_candidates) - 1)] if style_candidates[min(level, len(style_candidates) - 1)] in doc.styles else style_candidates[0]

    @staticmethod
    def _append_inline_runs(paragraph, node, bold=False, italic=False):
        if isinstance(node, NavigableString):
            if str(node):
                run = paragraph.add_run(str(node))
                run.bold, run.italic = bold, italic
            return
        if not isinstance(node, Tag): return
        if node.name == "br":
            paragraph.add_run("\n")
            return
        next_bold = bold or node.name in {"strong", "b"}
        next_italic = italic or node.name in {"em", "i"}
        for child in node.children:
            DocumentBuilder._append_inline_runs(paragraph, child, bold=next_bold, italic=next_italic)

    @staticmethod
    def _render_list(doc, list_node, level=0):
        style_name = DocumentBuilder._resolve_list_style(doc, list_node.name, level)
        for list_item in list_node.find_all("li", recursive=False):
            paragraph = doc.add_paragraph(style=style_name)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for child in list_item.contents:
                if isinstance(child, Tag) and child.name in {"ul", "ol"}: continue
                DocumentBuilder._append_inline_runs(paragraph, child)
            for nested in [child for child in list_item.contents if isinstance(child, Tag) and child.name in {"ul", "ol"}]:
                DocumentBuilder._render_list(doc, nested, level=level + 1)

    @staticmethod
    def _render_table(doc, table_node):
        rows = table_node.find_all("tr")
        if not rows: return
        column_count = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
        table = doc.add_table(rows=0, cols=column_count)
        table.style = "Table Grid"

        for row_index, row_node in enumerate(rows):
            cells = row_node.find_all(["th", "td"], recursive=False)
            table_row = table.add_row().cells
            for col_index in range(column_count):
                if col_index < len(cells): table_row[col_index].text = cells[col_index].get_text(" ", strip=True)
                if row_index == 0 and col_index < len(cells) and cells[col_index].name == "th":
                    for run in table_row[col_index].paragraphs[0].runs: run.bold = True

    @staticmethod
    def parse_html_to_docx(doc, html_content):
        soup = BeautifulSoup(html_content, "html.parser")
        for element in soup.contents:
            if not isinstance(element, Tag): continue
            if element.name in {"h1", "h2", "h3"}:
                doc.add_heading(element.get_text(" ", strip=True), level=int(element.name[1]))
                continue
            if element.name == "p":
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for child in element.children: DocumentBuilder._append_inline_runs(paragraph, child)
                continue
            if element.name in {"ul", "ol"}:
                DocumentBuilder._render_list(doc, element, level=0)
                continue
            if element.name == "table":
                DocumentBuilder._render_table(doc, element)

    @staticmethod
    def process_content(doc, raw_text, theme_color=DEFAULT_COLOR):
        clean_lines = []
        for line in raw_text.split("\n"):
            stripped_line = line.strip()
            if stripped_line.startswith("[[CHART:") and stripped_line.endswith("]]"):
                image = ChartEngine.create_bar_chart(stripped_line.replace("[[CHART:", "").replace("]]", "").strip(), theme_color)
                if image: doc.add_paragraph().add_run().add_picture(image, width=Inches(5.5))
                continue
            if stripped_line.startswith("[[PIE:") and stripped_line.endswith("]]"):
                image = ChartEngine.create_pie_chart(stripped_line.replace("[[PIE:", "").replace("]]", "").strip(), theme_color)
                if image: doc.add_paragraph().add_run().add_picture(image, width=Inches(5.4))
                continue
            if stripped_line.startswith("[[FLOW:") and stripped_line.endswith("]]"):
                image = ChartEngine.create_flowchart(stripped_line.replace("[[FLOW:", "").replace("]]", "").strip(), theme_color)
                if image: doc.add_paragraph().add_run().add_picture(image, width=Inches(6.5))
                continue
            clean_lines.append(line)

        html = markdown.markdown("\n".join(clean_lines), extensions=["tables"])
        DocumentBuilder.parse_html_to_docx(doc, html)

    @staticmethod
    def add_table_of_contents(doc):
        doc.add_heading("DAFTAR ISI", level=1)
        toc_paragraph = doc.add_paragraph()
        append_field(toc_paragraph, 'TOC \\o "1-3" \\h \\z \\u')
        note = doc.add_paragraph("Perbarui field di Microsoft Word agar daftar isi otomatis terisi.")
        note.runs[0].italic, note.runs[0].font.size, note.alignment = True, Pt(9), WD_ALIGN_PARAGRAPH.LEFT
        doc.add_page_break()

    @staticmethod
    def create_cover(doc, timeframe, theme_color=DEFAULT_COLOR):
        StyleEngine.apply_document_styles(doc, theme_color)
        for _ in range(5): doc.add_paragraph()

        confidentiality = doc.add_paragraph("S T R I C T L Y   C O N F I D E N T I A L")
        confidentiality.alignment = WD_ALIGN_PARAGRAPH.CENTER
        confidentiality.runs[0].font.size, confidentiality.runs[0].font.color.rgb, confidentiality.runs[0].font.bold = Pt(10), RGBColor(128, 128, 128), True

        doc.add_paragraph()
        report_title = doc.add_paragraph("HOLISTIC CUSTOMER EXPERIENCE REPORT")
        report_title.alignment, report_title.runs[0].font.name, report_title.runs[0].font.size = WD_ALIGN_PARAGRAPH.CENTER, "Calibri", Pt(20)

        company_name = doc.add_paragraph("INIXINDO JOGJA")
        company_name.alignment, company_name.runs[0].font.name, company_name.runs[0].font.bold, company_name.runs[0].font.size, company_name.runs[0].font.color.rgb = WD_ALIGN_PARAGRAPH.CENTER, "Calibri", True, Pt(32), RGBColor(*theme_color)

        doc.add_paragraph()
        period_text = doc.add_paragraph(f"Periode Evaluasi Laporan: {timeframe}")
        period_text.alignment, period_text.runs[0].font.size = WD_ALIGN_PARAGRAPH.CENTER, Pt(13)

        generated_text = doc.add_paragraph(f"Tanggal Generasi: {datetime.now().strftime('%d %B %Y')}")
        generated_text.alignment, generated_text.runs[0].font.size, generated_text.runs[0].font.color.rgb = WD_ALIGN_PARAGRAPH.CENTER, Pt(11), RGBColor(128, 128, 128)

        for _ in range(8): doc.add_paragraph()
        prepared_for = doc.add_paragraph(f"Prepared for Executive Board by:\n{WRITER_FIRM_NAME}")
        prepared_for.alignment, prepared_for.runs[0].font.bold = WD_ALIGN_PARAGRAPH.CENTER, True

        doc.add_page_break()
        DocumentBuilder.add_table_of_contents(doc)

class ReportQualityValidator:
    REQUIRED_CHAPTER_IDS = {"cx_chap_1": "Descriptive chapter tersedia", "cx_chap_2": "Diagnostic chapter tersedia", "cx_chap_3": "Predictive chapter tersedia", "cx_chap_4": "Prescriptive chapter tersedia", "cx_chap_5": "Implementation readiness chapter tersedia"}

    @staticmethod
    def _plain_text(value):
        text = str(value or "")
        text = re.sub(r"\[\[(?:CHART|PIE|FLOW):.*?\]\]", " ", text)
        text = re.sub(r"[#*`>|_]", " ", text)
        return re.sub(r"\s+", " ", text).strip()

    @staticmethod
    def _section_map(report_sections):
        return {section.get("id"): section.get("content", "") for section in report_sections}

    @staticmethod
    def _check(checks, label, passed):
        checks.append({"label": label, "passed": bool(passed)})

    @classmethod
    def evaluate(cls, document, executive_snapshot, report_sections, score_label):
        checks = []
        section_map = cls._section_map(report_sections)
        plain_combined = cls._plain_text("\n".join([executive_snapshot or "", "\n".join(section_map.values())])).lower()

        cls._check(checks, "Executive snapshot substantif", len(cls._plain_text(executive_snapshot)) >= 400)
        for section_id, label in cls.REQUIRED_CHAPTER_IDS.items(): cls._check(checks, label, len(cls._plain_text(section_map.get(section_id, ""))) >= 250)

        cls._check(checks, "Score engine POV tercermin", score_label.lower() in plain_combined)
        cls._check(checks, "Customer journey teridentifikasi", "customer journey" in plain_combined or "tahap customer journey" in plain_combined)
        cls._check(checks, "Lokasi pelatihan tercantum", "lokasi pelatihan" in plain_combined or " lokasi " in f" {plain_combined} ")
        cls._check(checks, "Tipe instruktur tercantum", "tipe instruktur" in plain_combined or "instruktur" in plain_combined)
        cls._check(checks, "Prediksi menggunakan bahasa manusia", bool(re.search(r"diproyeksikan (turun|naik|relatif stabil)", plain_combined)))
        cls._check(checks, "Prediksi menyebut horizon waktu", bool(re.search(r"(januari|februari|maret|april|mei|juni|juli|agustus|september|oktober|november|desember)\s+\d{4}|pada tahun \d{4}|1-2 bulan ke depan|1-2 minggu ke depan|semester berikutnya", plain_combined)))

        nonempty_paragraphs = sum(1 for paragraph in document.paragraphs if paragraph.text.strip())
        table_count, visual_count = len(document.tables), len(document.inline_shapes)

        cls._check(checks, "Dokumen memiliki paragraf yang memadai", nonempty_paragraphs >= 80)
        cls._check(checks, "Dokumen memiliki tabel pendukung", table_count >= 8)
        cls._check(checks, "Dokumen memiliki visual pendukung", visual_count >= 3)

        passed_checks = sum(1 for check in checks if check["passed"])
        total_checks = len(checks)
        completeness_score = round((passed_checks / total_checks) * 100, 1) if total_checks else 0.0
        verified_complete = completeness_score >= 80.0

        return {
            "verification_status": "verified" if verified_complete else "needs_review", "verified_complete": verified_complete,
            "completeness_score": completeness_score, "passed_checks": passed_checks, "total_checks": total_checks,
            "missing_checks": [check["label"] for check in checks if not check["passed"]],
            "document_stats": {"paragraph_count": nonempty_paragraphs, "table_count": table_count, "visual_count": visual_count},
            "summary": f"{passed_checks}/{total_checks} checks passed. Completeness score {completeness_score}%.",
        }

class ReportGenerator:
    def __init__(self, kb_instance):
        self.kb = kb_instance
        self.research_pool = concurrent.futures.ThreadPoolExecutor(max_workers=4)

    def run(self, timeframe, notes="", sentiment="all", segment="all", score_engine=DEFAULT_SCORE_ENGINE):
        logger.info("Starting feedback intelligence report generation for timeframe=%s, sentiment=%s, segment=%s, score_engine=%s", timeframe, sentiment, segment, score_engine)
        score_profile = SCORE_ENGINE_PROFILES.get(score_engine, SCORE_ENGINE_PROFILES[DEFAULT_SCORE_ENGINE])

        macro_future = self.research_pool.submit(Researcher.get_macro_trends, timeframe, notes, score_profile["label"])
        try: macro_trends = macro_future.result(timeout=45) # Increased timeout to accommodate deep scraping
        except Exception: macro_trends = "Tidak ada tren eksternal yang berhasil dimuat."

        analytics = FeedbackAnalyticsEngine(self.kb.df)
        executive_snapshot = analytics.build_executive_snapshot(timeframe, notes, sentiment=sentiment, segment=segment, score_engine=score_engine)
        report_sections = analytics.build_report_sections(timeframe, notes, macro_trends, sentiment=sentiment, segment=segment, score_engine=score_engine)

        document = Document()
        DocumentBuilder.create_cover(document, timeframe, DEFAULT_COLOR)
        document.add_heading("EXECUTIVE SNAPSHOT", level=1)
        DocumentBuilder.process_content(document, executive_snapshot, DEFAULT_COLOR)
        document.add_page_break()

        for index, section in enumerate(report_sections):
            document.add_heading(section["title"], level=1)
            DocumentBuilder.process_content(document, section["content"], DEFAULT_COLOR)
            if index < len(report_sections) - 1: document.add_page_break()

        filename = f"Inixindo_Feedback_Intelligence_Report_{score_profile['label']}_{timeframe}".replace(" ", "_")
        quality = ReportQualityValidator.evaluate(document, executive_snapshot, report_sections, score_profile["label"])
        if not quality["verified_complete"]: logger.warning("Generated report is below completeness target: %s", quality["summary"])
        return document, filename, quality
