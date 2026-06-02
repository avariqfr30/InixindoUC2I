import sys
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

PROJECT_DIR = Path(__file__).resolve().parents[1]
if str(PROJECT_DIR) not in sys.path:
    sys.path.insert(0, str(PROJECT_DIR))

from document_builder import DocumentBuilder
from report_analytics import FeedbackAnalyticsEngine


class ReaderFacingDocumentContractTests(unittest.TestCase):
    def test_reader_facing_text_suppresses_internal_source_disclosures(self):
        raw = (
            "Semua Data APIDog (tanggal tidak tersedia) memakai Internal API / feedback cache. "
            "Jejak sumber yang dipakai: internal_api, komentar, osint, rating. "
            "internal facts remain separated from OSINT benchmark context."
        )

        clean = DocumentBuilder.reader_facing_text(raw)

        forbidden = ["APIDog", "Internal API", "internal_api", "internal facts", "feedback cache"]
        for token in forbidden:
            self.assertNotIn(token, clean)
        self.assertIn("Seluruh Periode Evaluasi", clean)
        self.assertIn("Jejak bukti yang dipakai", clean)

    def test_reader_facing_text_synthesizes_raw_ui_and_api_helpers(self):
        raw = (
            "Nama Perusahaan Klien: ReferenceAccount mencatat source=/api/Resource/dataset "
            "dataset_code=ConsultantProjectExpertHistory. Dirangkum dari sumber APIDog: "
            "Problem, Opportunity, Directive untuk Pain Points dan endpoint sync status."
        )

        clean = DocumentBuilder.reader_facing_text(raw)

        forbidden = [
            "Nama Perusahaan Klien",
            "ReferenceAccount",
            "source=",
            "/api/Resource/dataset",
            "dataset_code",
            "ConsultantProjectExpertHistory",
            "Dirangkum dari sumber",
            "APIDog",
            "Problem, Opportunity, Directive",
            "Pain Points",
            "endpoint",
            "sync status",
        ]
        for token in forbidden:
            self.assertNotIn(token, clean)
        self.assertIn("catatan klien", clean)
        self.assertIn("riwayat pengalaman konsultan", clean)
        self.assertIn("kebutuhan prioritas yang perlu dipertegas", clean)

    def test_reader_facing_text_translates_unnecessary_english_report_terms(self):
        raw = (
            "EXECUTIVE SUMMARY\nBLUF\nKey Findings\nRecommendation\n"
            "DESCRIPTIVE ANALYTICS & FEEDBACK GOVERNANCE\n"
            "Customer journey, score engine, rating, dashboard, insight, stakeholder, "
            "benchmark, owner, review, quick wins, approval gate, field, record, delivery, outcome."
        )

        clean = DocumentBuilder.reader_facing_text(raw)

        forbidden = [
            "EXECUTIVE SUMMARY",
            "BLUF",
            "Key Findings",
            "Recommendation",
            "DESCRIPTIVE ANALYTICS",
            "FEEDBACK GOVERNANCE",
            "Customer journey",
            "score engine",
            "rating",
            "dashboard",
            "insight",
            "stakeholder",
            "benchmark",
            "owner",
            "review",
            "quick wins",
            "approval gate",
            "field",
            "record",
            "delivery",
            "outcome",
        ]
        for token in forbidden:
            self.assertNotIn(token, clean)
        self.assertIn("Ringkasan Eksekutif", clean)
        self.assertIn("Inti Keputusan", clean)
        self.assertIn("Temuan Utama", clean)
        self.assertIn("Rekomendasi", clean)

    def test_generated_cover_has_real_word_toc_and_neutral_period_label(self):
        document = Document()

        DocumentBuilder.create_cover(document, "Semua Data APIDog (tanggal tidak tersedia)")
        document.add_heading("Ringkasan Eksekutif", level=1)
        document.add_heading("BAB I - Detail", level=1)

        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("DAFTAR ISI", text)
        self.assertIn("Ringkasan Eksekutif", text)
        self.assertIn("BAB I", text)
        self.assertIn("Seluruh Periode Evaluasi", text)
        self.assertNotIn("APIDog", text)
        self.assertNotIn("Perbarui field", text)
        settings_xml = document.settings.element.xml
        self.assertIn("updateFields", settings_xml)
        self.assertEqual(document.settings.element.find(qn("w:updateFields")).get(qn("w:val")), "true")
        document_xml = document.element.xml
        self.assertIn('TOC \\o "1-3" \\h \\z \\u', document_xml)

    def test_executive_summary_is_business_first_and_source_neutral(self):
        data = [
            {
                "Rentang Waktu": "Semua Data APIDog (tanggal tidak tersedia)",
                "Rating Numeric": 4,
                "Sentiment Label": "positive",
                "Layanan": "Kinerja instruktur",
                "Tipe Stakeholder": "Peserta Kelas",
                "Sumber Feedback": "Internal API",
                "Kanal Feedback": "Evaluasi Kelas Internal",
                "Komentar": "Instruktur jelas tetapi waktu praktik masih kurang.",
                "Customer Journey Stage": "Pelaksanaan Layanan",
            }
        ]
        engine = FeedbackAnalyticsEngine.from_records(data)

        summary = engine.build_executive_snapshot("Semua Data APIDog (tanggal tidak tersedia)")

        self.assertNotIn("## Ringkasan Eksekutif", summary)
        self.assertIn("### Kesimpulan Utama", summary)
        self.assertIn("### Alasan Utama", summary)
        self.assertIn("### Bukti Pendukung", summary)
        self.assertIn("### Keputusan yang Diminta", summary)
        self.assertIn("### Agenda Follow-up Meeting", summary)
        for token in ["APIDog", "Internal API", "internal_api", "internal facts", "Dataset Spesialis", "BLUF", "Key Findings", "Recommendation"]:
            self.assertNotIn(token, summary)

    def test_class_report_raw_form_labels_are_not_promoted_as_business_cases(self):
        data = [
            {
                "Record ID": "class-report-001",
                "Sumber Feedback": "class_report",
                "Kanal Feedback": "Evaluasi Kelas Internal",
                "Tanggal Feedback": "Tanggal tidak tersedia",
                "Tipe Stakeholder": "Peserta Kelas",
                "Layanan": "BRAND EQUITY (Mengapa Inixindo Jogja menjadi pilihan?) Pilih 4 Bintang untuk mengisi",
                "Rentang Waktu": "Semua Data APIDog (tanggal tidak tersedia)",
                "Rating": "47.6",
                "Komentar": "Rata-rata rating BRAND EQUITY (Mengapa Inixindo Jogja menjadi pilihan?) Pilih 4 Bintang untuk mengisi: 4.76 dari 5. Mengapa: reputasi dan instruktur menjadi alasan memilih.",
                "Customer Journey Hint": "Pelaksanaan Layanan",
                "Raw Response Count": "100",
                "Rating Response Count": "80",
                "Text Response Count": "20",
            }
        ]
        engine = FeedbackAnalyticsEngine.from_records(data)

        summary = engine.build_executive_snapshot("Semua Data APIDog (tanggal tidak tersedia)")
        sections = engine.build_report_sections("Semua Data APIDog (tanggal tidak tersedia)", "", "")
        combined = summary + "\n" + "\n".join(section["content"] for section in sections)

        self.assertNotIn("4.76/5", summary)
        for forbidden in [
            "BRAND EQUITY",
            "Mengapa Inixindo Jogja menjadi pilihan",
            "Pilih 4 Bintang untuk mengisi",
            "47.6/5",
            "rating 47.6",
            "Reputasi dan alasan memilih Inixindo",
        ]:
            self.assertNotIn(forbidden, combined)
        self.assertIn("Belum ada bukti evaluasi", summary)

    def test_executive_summary_uses_bluf_and_hides_agentic_desk_terms(self):
        data = [
            {
                "Record ID": "class-report-001",
                "Sumber Feedback": "class_report",
                "Kanal Feedback": "Evaluasi Kelas Internal",
                "Tanggal Feedback": "Tanggal tidak tersedia",
                "Tipe Stakeholder": "Peserta Kelas",
                "Layanan": "Kinerja instruktur",
                "Rentang Waktu": "1 Bulan Terakhir (Monthly)",
                "Rating": "4",
                "Komentar": "Instruktur jelas, tetapi waktu praktik masih kurang dan follow up perlu dipercepat.",
                "Customer Journey Hint": "Pelaksanaan Layanan",
            }
        ]
        engine = FeedbackAnalyticsEngine.from_records(data)

        summary = engine.build_executive_snapshot("1 Bulan Terakhir (Monthly)")

        expected_order = [
            "### Kesimpulan Utama",
            "### Keputusan yang Diminta",
            "### Dasbor Keputusan",
            "### Matriks Keputusan",
            "### Interpretasi Manajemen",
            "### Alasan Utama",
            "### Bukti Pendukung",
            "### Catatan Keyakinan dan Batasan",
            "### Agenda Follow-up Meeting",
        ]
        for marker in expected_order:
            self.assertIn(marker, summary)
        for before, after in zip(expected_order, expected_order[1:]):
            self.assertLess(summary.index(before), summary.index(after))
        for forbidden in [
            "Agentic",
            "Desk",
            "Review Tim Analis Internal",
            "Confidence Desk",
            "Evidence Ledger",
            "QA Guardrail",
            "Report Audit Trail",
            "Contradiction Check",
            "Historical Trend Desk",
            "Prediction Boundary",
        ]:
            self.assertNotIn(forbidden, summary)

    def test_positive_only_executive_summary_frames_strength_not_risk_fire_drill(self):
        data = [
            {
                "Record ID": "class-report-001",
                "Sumber Feedback": "Internal API",
                "Kanal Feedback": "Evaluasi Kelas Internal",
                "Tanggal Feedback": "Tanggal tidak tersedia",
                "Tipe Stakeholder": "Peserta Kelas",
                "Layanan": "Kinerja instruktur",
                "Rentang Waktu": "Semua Data APIDog (tanggal tidak tersedia)",
                "Rating": "5",
                "Komentar": "Instruktur jelas dan praktiknya mudah diikuti.",
                "Customer Journey Hint": "Pelaksanaan Layanan",
            }
        ]
        engine = FeedbackAnalyticsEngine.from_records(data)

        summary = engine.build_executive_snapshot("Semua Data APIDog (tanggal tidak tersedia)")

        self.assertIn("mempertahankan dan mereplikasi", summary)
        self.assertIn("kekuatan pengalaman pelanggan", summary.lower())
        self.assertNotIn("sinyal risiko pengalaman pelanggan paling jelas", summary)
        self.assertNotIn("memerlukan intervensi prioritas lintas fungsi", summary)


if __name__ == "__main__":
    unittest.main()
