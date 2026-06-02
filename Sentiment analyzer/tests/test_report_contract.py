import sys
import unittest
from pathlib import Path

PROJECT_DIR = Path(__file__).resolve().parents[1]
if str(PROJECT_DIR) not in sys.path:
    sys.path.insert(0, str(PROJECT_DIR))

import pandas as pd
from docx import Document
from docx.shared import Twips

from config import CX_SENTIMENT_STRUCTURE
from document_builder import DocumentBuilder
from report_evidence import ContextIntelligenceDesk, ReportEvidenceBuilder
from report_analytics import FeedbackAnalyticsEngine
from report_agents import FeedbackProposalTeam, HiddenAgentDesk
from report_quality import ReportQualityValidator


class ReportAnalyticsContractTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.dataframe = pd.read_csv(PROJECT_DIR / "data" / "db.csv")
        cls.engine = FeedbackAnalyticsEngine(cls.dataframe)
        cls.timeframe = "1 Bulan Terakhir (Monthly)"
        cls.notes = "Periksa risiko jadwal dan tindak lanjut layanan."
        cls.macro_trends = (
            "**Insight Mendalam (via example.com):** Benchmark pelatihan IT menekankan follow-up pasca kelas.\n\n"
            "1. Tren pelatihan korporat Indonesia | Kebutuhan evaluasi dampak meningkat | sumber=example.com | tanggal=2026"
        )

    def test_report_sections_match_configured_structure(self):
        sections = self.engine.build_report_sections(
            self.timeframe,
            self.notes,
            self.macro_trends,
            score_engine="experience_index",
        )

        self.assertEqual([item["id"] for item in sections], [item["id"] for item in CX_SENTIMENT_STRUCTURE])
        self.assertEqual([item["title"] for item in sections], [item["title"] for item in CX_SENTIMENT_STRUCTURE])
        self.assertTrue(all(section["content"].strip() for section in sections))
        self.assertTrue(all(not section["content"].lstrip().startswith("### Bukti yang Dipakai") for section in sections))

        combined = "\n".join(section["content"] for section in sections)
        self.assertNotIn("Bukti yang Dipakai", combined)
        self.assertNotIn("### Cakupan sumber dan kanal", combined)
        required_markers = [
            "## 1.1 Ringkasan Cakupan Umpan Balik dan Tata Kelola",
            "## 2.1 Akar Masalah Utama dan Titik Keluhan Dominan",
            "## 3.1 Risiko Jangka Pendek Jika Pola Saat Ini Berlanjut",
            "## 4.1 Intervensi Prioritas 30 Hari",
            "## 5.1 Prioritas Sasaran Bisnis",
            "Experience Index",
        ]
        for marker in required_markers:
            self.assertIn(marker, combined)

    def test_placeholder_brand_reason_dimension_is_not_promoted_to_report_priority(self):
        dataframe = pd.DataFrame(
            [
                {
                    "Record ID": "brand-placeholder",
                    "Tanggal Feedback": "2026-05-20",
                    "Rentang Waktu": "2026-05",
                    "Tipe Stakeholder": "Peserta Kelas",
                    "Layanan": "BRAND EQUITY (Mengapa Inixindo Jogja menjadi pilihan?) Pilih 4 Bintang untuk mengisi",
                    "Rating": "1",
                    "Komentar": "Rata-rata rating BRAND EQUITY (Mengapa Inixindo Jogja menjadi pilihan?): 1.0 dari 5. Mengapa: TIDAK",
                    "Customer Journey Hint": "Tindak Lanjut dan Outcome",
                },
                {
                    "Record ID": "material-critique",
                    "Tanggal Feedback": "2026-05-20",
                    "Rentang Waktu": "2026-05",
                    "Tipe Stakeholder": "Peserta Kelas",
                    "Layanan": "Materi dan kurikulum",
                    "Rating": "4",
                    "Komentar": "Materi menggunakan versi lama dan perlu diupdate agar sesuai kebutuhan peserta.",
                    "Customer Journey Hint": "Pelaksanaan Layanan",
                },
                {
                    "Record ID": "instructor-positive",
                    "Tanggal Feedback": "2026-05-20",
                    "Rentang Waktu": "2026-05",
                    "Tipe Stakeholder": "Peserta Kelas",
                    "Layanan": "Kinerja instruktur",
                    "Rating": "5",
                    "Komentar": "Instruktur jelas dan praktiknya mudah diikuti.",
                    "Customer Journey Hint": "Pelaksanaan Layanan",
                },
            ]
        )
        engine = FeedbackAnalyticsEngine(dataframe)

        snapshot = engine.build_executive_snapshot("2026-05", score_engine="experience_index")
        sections = engine.build_report_sections("2026-05", "", "Tidak ada tren eksternal yang berhasil dimuat.", score_engine="experience_index")
        combined = snapshot + "\n" + "\n".join(section["content"] for section in sections)

        self.assertNotIn("Reputasi dan alasan memilih Inixindo", combined)
        self.assertNotIn("BRAND EQUITY", combined)
        self.assertNotIn("Mengapa Inixindo Jogja menjadi pilihan", combined)
        self.assertIn("Materi dan kurikulum", combined)

    def test_text_aware_sentiment_reclassifies_four_star_critique_and_weak_low_rating(self):
        dataframe = pd.DataFrame(
            [
                {
                    "Record ID": "four-critique",
                    "Tanggal Feedback": "2026-05-20",
                    "Rentang Waktu": "2026-05",
                    "Tipe Stakeholder": "Peserta Kelas",
                    "Layanan": "Materi dan kurikulum",
                    "Rating": "4",
                    "Komentar": "Materi masih menggunakan versi lama dan perlu diupdate.",
                },
                {
                    "Record ID": "weak-negative",
                    "Tanggal Feedback": "2026-05-20",
                    "Rentang Waktu": "2026-05",
                    "Tipe Stakeholder": "Peserta Kelas",
                    "Layanan": "Logistik",
                    "Rating": "1",
                    "Komentar": "TIDAK",
                },
                {
                    "Record ID": "clean-positive",
                    "Tanggal Feedback": "2026-05-20",
                    "Rentang Waktu": "2026-05",
                    "Tipe Stakeholder": "Peserta Kelas",
                    "Layanan": "Kinerja instruktur",
                    "Rating": "5",
                    "Komentar": "Instruktur jelas dan contoh praktiknya relevan.",
                },
            ]
        )
        engine = FeedbackAnalyticsEngine(dataframe)
        labels = dict(zip(engine.full_df["Record ID"], engine.full_df["Sentiment Label"]))

        self.assertEqual(labels["four-critique"], "mixed")
        self.assertEqual(labels["weak-negative"], "weak_negative")
        self.assertEqual(labels["clean-positive"], "positive")

        sections = engine.build_report_sections("2026-05", "", "Tidak ada tren eksternal yang berhasil dimuat.")
        combined = "\n".join(section["content"] for section in sections)

        self.assertIn("Kritik konstruktif", combined)
        self.assertIn("Negatif bukti lemah", combined)
        self.assertNotIn('"TIDAK"', combined)

    def test_rolling_month_timeframe_filters_by_feedback_date(self):
        dataframe = pd.DataFrame(
            [
                {
                    "Record ID": "A",
                    "Tanggal Feedback": "2026-03-29",
                    "Rentang Waktu": "2026-03-01 sampai 2026-03-31",
                    "Tipe Stakeholder": "Peserta",
                    "Layanan": "Materi",
                    "Rating": "5",
                    "Komentar": "Materi relevan.",
                },
                {
                    "Record ID": "B",
                    "Tanggal Feedback": "2026-02-15",
                    "Rentang Waktu": "2026-02-01 sampai 2026-02-28",
                    "Tipe Stakeholder": "Peserta",
                    "Layanan": "Instruktur",
                    "Rating": "4",
                    "Komentar": "Instruktur jelas.",
                },
                {
                    "Record ID": "C",
                    "Tanggal Feedback": "2025-10-01",
                    "Rentang Waktu": "2025-10-01 sampai 2025-10-31",
                    "Tipe Stakeholder": "Peserta",
                    "Layanan": "Fasilitas",
                    "Rating": "2",
                    "Komentar": "Fasilitas perlu diperbaiki.",
                },
            ]
        )
        engine = FeedbackAnalyticsEngine(dataframe)

        one_month = engine._filter_view("1 Bulan Terakhir")
        six_months = engine._filter_view("6 Bulan Terakhir")

        self.assertEqual(set(one_month["Record ID"]), {"A", "B"})
        self.assertEqual(set(six_months["Record ID"]), {"A", "B", "C"})

    def test_timeframe_options_include_only_rolling_month_choices_when_dates_exist(self):
        from timeframe_filters import build_available_date_options, build_timeframe_options

        dataframe = pd.DataFrame(
            [
                {"Tanggal Feedback": "2026-04-03", "Rentang Waktu": "2026-04-01 sampai 2026-04-03"},
                {"Tanggal Feedback": "2026-03-01", "Rentang Waktu": "2026-03"},
                {"Tanggal Feedback": "2026-05-12", "Rentang Waktu": "2026-05-10 sampai 2026-05-12"},
                {"Tanggal Feedback": "Tanggal tidak tersedia", "Rentang Waktu": "Semua Data APIDog (tanggal tidak tersedia)"},
            ]
        )

        options = build_timeframe_options(dataframe)
        available_dates = build_available_date_options(dataframe)

        self.assertEqual(options, [
            "1 Bulan Terakhir",
            "3 Bulan Terakhir",
            "6 Bulan Terakhir",
            "12 Bulan Terakhir / 1 Tahun",
        ])
        self.assertNotIn("2026-04-01 sampai 2026-04-03", options)
        self.assertNotIn("Semua Data APIDog (tanggal tidak tersedia)", options)
        self.assertEqual(available_dates["min"], "2026-03-01")
        self.assertEqual(available_dates["max"], "2026-05-12")
        self.assertEqual(available_dates["dates"], ["2026-03-01", "2026-04-03", "2026-05-12"])

    def test_custom_date_range_timeframe_filters_by_start_and_end_date(self):
        from timeframe_filters import custom_timeframe_label

        dataframe = pd.DataFrame(
            [
                {
                    "Record ID": "A",
                    "Tanggal Feedback": "2026-05-12",
                    "Rentang Waktu": "2026-05",
                    "Tipe Stakeholder": "Peserta",
                    "Layanan": "Materi",
                    "Rating": "5",
                    "Komentar": "Materi relevan.",
                },
                {
                    "Record ID": "B",
                    "Tanggal Feedback": "2026-04-03",
                    "Rentang Waktu": "2026-04",
                    "Tipe Stakeholder": "Peserta",
                    "Layanan": "Instruktur",
                    "Rating": "4",
                    "Komentar": "Instruktur jelas.",
                },
                {
                    "Record ID": "C",
                    "Tanggal Feedback": "2026-03-01",
                    "Rentang Waktu": "2026-03",
                    "Tipe Stakeholder": "Peserta",
                    "Layanan": "Fasilitas",
                    "Rating": "2",
                    "Komentar": "Fasilitas perlu diperbaiki.",
                },
            ]
        )
        engine = FeedbackAnalyticsEngine(dataframe)

        scoped = engine._filter_view(custom_timeframe_label("2026-04-01", "2026-05-31"))

        self.assertEqual(set(scoped["Record ID"]), {"A", "B"})

        summary = engine.build_executive_snapshot(custom_timeframe_label("2026-04-01", "2026-05-31"))
        self.assertIn("2026-04-01 sampai 2026-05-31", summary)
        self.assertNotIn("custom_range:", summary)

    def test_report_evidence_cards_are_source_safe(self):
        sections = [
            {
                "id": "cx_chap_1",
                "title": "Analisis Deskriptif",
                "content": "APIDog source=/api/Resource/dataset menunjukkan 120 respons dan risiko onboarding.",
            }
        ]

        enriched = ReportEvidenceBuilder.attach_to_sections(sections)
        content = enriched[0]["content"]

        self.assertIn("### Bukti yang Dipakai", content)
        self.assertIn("120 respons", content)
        for forbidden in ["APIDog", "/api/Resource/dataset", "source=", "Evidence Ledger"]:
            self.assertNotIn(forbidden, content)

    def test_external_osint_brief_is_summarized_not_raw_links(self):
        from osint_research import Researcher

        brief = Researcher._format_osint_brief(
            [
                {
                    "title": "Raw search result title",
                    "snippet": "Survei pelanggan menunjukkan ekspektasi follow-up layanan meningkat 28% setelah program pelatihan digital.",
                    "url": "https://example.com/report/customer-expectation",
                    "date": "2026",
                    "source_quality": 2,
                }
            ],
            "Sinyal OSINT Makro (Indonesia)",
        )

        self.assertIn("ekspektasi follow-up layanan meningkat", brief)
        self.assertIn("(Sumber: example.com, 2026)", brief)
        self.assertNotIn("url=", brief)
        self.assertNotIn("source=", brief)
        self.assertNotIn("https://", brief)
        self.assertNotIn("Raw search result title |", brief)

    def test_executive_section_synthesis_summarizes_chapter_callouts(self):
        synthesis = self.engine._executive_section_synthesis(
            [
                {
                    "title": "Analisis Prediktif",
                    "content": "Rating turun 12% pada layanan prioritas sehingga tindak lanjut perlu dipercepat. url=https://example.com source=example.",
                }
            ]
        )

        self.assertIn("Analisis Prediktif", synthesis)
        self.assertIn("Rating turun 12%", synthesis)
        self.assertNotIn("url=", synthesis)
        self.assertNotIn("source=", synthesis)

    def test_context_intelligence_desk_turns_raw_notes_into_reader_safe_focus(self):
        packet = ContextIntelligenceDesk.build(
            dataframe=self.dataframe,
            notes="APIDog source=/api/Resource/dataset meminta fokus Problem, Opportunity, Directive pada onboarding.",
            timeframe=self.timeframe,
            sentiment="all",
            segment="all",
            score_engine="experience_index",
        )

        self.assertIn("onboarding", packet["focus_note"].lower())
        self.assertIn("periode", packet["coverage_note"].lower())
        for forbidden in ["APIDog", "/api/Resource/dataset", "source=", "Problem, Opportunity, Directive", "Context Intelligence Desk"]:
            self.assertNotIn(forbidden, "\n".join(str(value) for value in packet.values()))

    def test_context_intelligence_desk_marks_weak_external_research_conservatively(self):
        packet = ContextIntelligenceDesk.build(
            dataframe=self.dataframe,
            notes="Perlu cek tindak lanjut kelas.",
            timeframe=self.timeframe,
            macro_trends="Tidak ada tren eksternal yang berhasil dimuat.",
        )

        self.assertFalse(packet["external_context_ready"])
        self.assertIn("belum cukup kuat", packet["external_context_note"].lower())
        self.assertNotIn("OSINT", packet["external_context_note"])

    def test_narrative_preflight_rejects_empty_and_raw_source_content(self):
        result = ReportQualityValidator.evaluate_narrative(
            "## Ringkasan Eksekutif\nRingkas.",
            [{"id": "cx_chap_1", "title": "Bab 1", "content": "APIDog endpoint."}],
        )

        self.assertFalse(result["passes"])
        self.assertIn("raw_source_label", result["categories"])

    def test_executive_snapshot_keeps_decision_ready_contract(self):
        snapshot = self.engine.build_executive_snapshot(
            self.timeframe,
            self.notes,
            score_engine="experience_index",
        )

        required_markers = [
            "### Kesimpulan Utama",
            "### Keputusan yang Diminta",
            "### Dasbor Keputusan",
            "### Matriks Keputusan",
            "### Interpretasi Manajemen",
            "### Alasan Utama",
            "### Bukti Pendukung",
            "### Catatan Keyakinan dan Batasan",
            "### Agenda Follow-up Meeting",
            "| Pertanyaan Eksekutif | Jawaban Singkat |",
        ]
        for marker in required_markers:
            self.assertIn(marker, snapshot)
        self.assertNotIn("Formula Experience Index", snapshot)
        self.assertLess(snapshot.index("### Kesimpulan Utama"), snapshot.index("### Keputusan yang Diminta"))
        self.assertLess(snapshot.index("### Keputusan yang Diminta"), snapshot.index("### Dasbor Keputusan"))
        self.assertLess(snapshot.index("### Dasbor Keputusan"), snapshot.index("### Matriks Keputusan"))
        self.assertLess(snapshot.index("### Matriks Keputusan"), snapshot.index("### Interpretasi Manajemen"))
        self.assertLess(snapshot.index("### Interpretasi Manajemen"), snapshot.index("### Alasan Utama"))
        self.assertLess(snapshot.index("### Alasan Utama"), snapshot.index("### Bukti Pendukung"))
        self.assertLess(snapshot.index("### Bukti Pendukung"), snapshot.index("### Catatan Keyakinan dan Batasan"))

        sections = self.engine.build_report_sections(
            self.timeframe,
            self.notes,
            self.macro_trends,
            score_engine="experience_index",
        )
        combined_sections = "\n".join(section["content"] for section in sections)
        self.assertIn("Penjelasan Perhitungan Experience Index", combined_sections)
        self.assertIn("titik sentuh", combined_sections.lower())
        self.assertIn("dirasakan", combined_sections.lower())
        self.assertIn("agenda", combined_sections.lower())
        self.assertNotIn("Indeks Pengalaman", snapshot + "\n" + combined_sections)

    def test_predictive_section_explains_challengeable_early_warning_model(self):
        sections = self.engine.build_report_sections(
            self.timeframe,
            self.notes,
            self.macro_trends,
            score_engine="experience_index",
        )
        predictive = next(section["content"] for section in sections if section["id"] == "cx_chap_3")

        self.assertIn("peringatan dini", predictive.lower())
        self.assertIn("faktor pembentuk proyeksi", predictive.lower())
        self.assertIn("penggerak keyakinan", predictive.lower())
        self.assertIn("uji kewajaran", predictive.lower())
        self.assertIn("sensitivitas", predictive.lower())
        self.assertIn("belum diklaim sebagai backtesting statistik", predictive.lower())

    def test_executive_snapshot_follows_minto_decision_order(self):
        snapshot = self.engine.build_executive_snapshot(
            self.timeframe,
            self.notes,
            score_engine="experience_index",
        )

        expected_order = [
            "### Kesimpulan Utama",
            "### Keputusan yang Diminta",
            "### Dasbor Keputusan",
            "### Matriks Keputusan",
            "### Interpretasi Manajemen",
            "### Alasan Utama",
            "### Bukti Pendukung",
            "### Catatan Keyakinan dan Batasan",
            "### Agenda Follow-up Meeting",
        ]
        for marker in expected_order:
            self.assertIn(marker, snapshot)
        for before, after in zip(expected_order, expected_order[1:]):
            self.assertLess(snapshot.index(before), snapshot.index(after))

    def test_executive_snapshot_uses_inixindo_management_interpretation_layer(self):
        snapshot = self.engine.build_executive_snapshot(
            self.timeframe,
            self.notes,
            score_engine="experience_index",
        )

        self.assertIn("### Interpretasi Manajemen", snapshot)
        self.assertIn("| Sinyal | Makna | Keputusan | Aksi | Keyakinan |", snapshot)
        self.assertIn("penilaian", snapshot.lower())
        self.assertIn("pengalaman pelanggan", snapshot.lower())
        self.assertIn("30 hari", snapshot)
        self.assertLess(snapshot.index("### Dasbor Keputusan"), snapshot.index("### Interpretasi Manajemen"))
        self.assertLess(snapshot.index("### Interpretasi Manajemen"), snapshot.index("### Alasan Utama"))

    def test_executive_snapshot_uses_finished_sections_without_generic_report_meta(self):
        sections = [
            {"title": "Analisis Deskriptif", "content": "Evaluasi menunjukkan keluhan onboarding terkonsentrasi pada respons awal."},
            {"title": "Kesiapan Implementasi", "content": "Owner layanan perlu menetapkan tindak lanjut 30 hari."},
        ]

        snapshot = self.engine.build_executive_snapshot(
            self.timeframe,
            self.notes,
            score_engine="experience_index",
            report_sections=sections,
            section_context={"focus_note": "onboarding perlu dipercepat", "coverage_note": "Cakupan pembacaan memakai periode bulanan."},
        )

        self.assertIn("Sintesis Lintas Bab", snapshot)
        self.assertIn("onboarding", snapshot.lower())
        self.assertNotIn("Ringkasan ini dirancang", snapshot)
        self.assertNotIn("Fokus tambahan:", snapshot)

    def test_executive_snapshot_reports_raw_response_volume_for_aggregated_class_report(self):
        dataframe = pd.DataFrame(
            [
                {
                    "Record ID": "class_report-00001",
                    "Sumber Feedback": "class_report",
                    "Kanal Feedback": "Evaluasi Kelas Internal",
                    "Tanggal Feedback": "Tanggal tidak tersedia",
                    "Tipe Stakeholder": "Peserta Kelas",
                    "Layanan": "Materi dan kurikulum",
                    "Lokasi": "",
                    "Tipe Instruktur": "",
                    "Rentang Waktu": "Semua Data APIDog (tanggal tidak tersedia)",
                    "Rating": "3",
                    "Komentar": "Rata-rata rating Kesesuaian materi bahan ajar: 3.0 dari 5. Mengapa: Materi perlu contoh praktik tambahan.",
                    "Customer Journey Hint": "Pelaksanaan Layanan",
                    "Raw Response Count": "3",
                    "Rating Response Count": "2",
                    "Text Response Count": "1",
                    "Rating Distribution": "2: 1; 4: 1",
                    "Representative Why": "Materi perlu contoh praktik tambahan.",
                }
            ]
        )
        engine = FeedbackAnalyticsEngine(dataframe)

        snapshot = engine.build_executive_snapshot(
            "Semua Data APIDog (tanggal tidak tersedia)",
            score_engine="experience_index",
        )

        self.assertIn("3 respons", snapshot)
        self.assertIn("1 dimensi", snapshot)
        self.assertIn("2 penilaian", snapshot)
        self.assertIn("1 komentar teks", snapshot)

    def test_docx_quality_accepts_generated_contract(self):
        snapshot = self.engine.build_executive_snapshot(
            self.timeframe,
            self.notes,
            score_engine="experience_index",
        )
        sections = self.engine.build_report_sections(
            self.timeframe,
            self.notes,
            self.macro_trends,
            score_engine="experience_index",
        )

        document = Document()
        DocumentBuilder.process_content(document, snapshot)
        for section in sections:
            document.add_heading(section["title"], level=1)
            DocumentBuilder.process_content(document, section["content"])

        quality = ReportQualityValidator.evaluate(document, snapshot, sections, "Experience Index")
        self.assertTrue(quality["verified_complete"], quality)
        self.assertEqual(quality["passed_checks"], quality["total_checks"])
        table_headers = [
            [cell.text for cell in table.rows[0].cells]
            for table in document.tables
        ]
        self.assertIn(["Lensa Experience Index", "Skor", "Pembacaan", "Bukti Analitis"], table_headers)
        self.assertFalse(any("| Lensa Experience Index |" in paragraph.text for paragraph in document.paragraphs))

    def test_docx_ordered_lists_restart_and_use_hanging_indents(self):
        document = Document()
        DocumentBuilder.process_content(
            document,
            "1. Tahap pertama\n2. Tahap kedua\n\nParagraf pemisah.\n\n1. Siklus baru\n2. Nomor kedua",
        )

        numbered_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None]
        self.assertEqual(len(numbered_paragraphs), 4)

        first_list_num_id = numbered_paragraphs[0]._p.pPr.numPr.numId.val
        second_list_num_id = numbered_paragraphs[2]._p.pPr.numPr.numId.val
        self.assertNotEqual(first_list_num_id, second_list_num_id)

        for paragraph in numbered_paragraphs:
            paragraph_format = paragraph.paragraph_format
            self.assertEqual(paragraph_format.left_indent, Twips(720))
            self.assertEqual(paragraph_format.first_line_indent, Twips(-360))

    def test_experience_index_is_explained_as_touchpoint_journey_not_bare_metric(self):
        sections = self.engine.build_report_sections(
            self.timeframe,
            self.notes,
            self.macro_trends,
            score_engine="experience_index",
        )
        predictive = next(section["content"] for section in sections if section["id"] == "cx_chap_3")
        lower_predictive = predictive.lower()

        self.assertNotRegex(predictive, r"(?m)^\|\s*Indeks Pengalaman\s*\|")
        self.assertIn("titik sentuh", lower_predictive)
        self.assertIn("dirasakan", lower_predictive)
        self.assertIn("perjalanan", lower_predictive)
        self.assertIn("agenda", lower_predictive)

    def test_predictive_section_links_osint_to_company_current_and_future_state(self):
        sections = self.engine.build_report_sections(
            self.timeframe,
            self.notes,
            self.macro_trends,
            score_engine="experience_index",
        )
        predictive = next(section["content"] for section in sections if section["id"] == "cx_chap_3")

        self.assertIn("## 3.4 Keterkaitan Faktor Eksternal dengan Kondisi Perusahaan", predictive)
        self.assertIn("kondisi perusahaan saat ini", predictive)
        self.assertIn("kondisi perusahaan ke depan", predictive)
        self.assertIn("Ringkasan tren eksternal", predictive)
        self.assertNotIn("**Insight Mendalam", predictive)
        self.assertNotIn("| Sinyal Eksternal | Sumber | Tanggal |", predictive)

    def test_weak_osint_keeps_prediction_internal_data_first(self):
        sections = self.engine.build_report_sections(
            self.timeframe,
            self.notes,
            "Tidak ada tren eksternal yang berhasil dimuat.",
            score_engine="experience_index",
            section_context={"external_context_ready": False},
        )
        predictive = next(section["content"] for section in sections if section["id"] == "cx_chap_3")

        self.assertIn("bukti evaluasi internal", predictive.lower())
        self.assertNotIn("| Sinyal Eksternal | Sumber | Tanggal |", predictive)
        self.assertNotIn("urgensi intervensi meningkat", predictive)

    def test_predictive_section_includes_challenge_and_confidence_defensibility_layer(self):
        sections = self.engine.build_report_sections(
            self.timeframe,
            self.notes,
            self.macro_trends,
            score_engine="experience_index",
        )
        predictive = next(section["content"] for section in sections if section["id"] == "cx_chap_3")
        lower_predictive = predictive.lower()

        self.assertIn("peringatan dini", lower_predictive)
        self.assertRegex(lower_predictive, r"penggerak keyakinan|pemicu keyakinan|confidence driver")
        self.assertRegex(lower_predictive, r"uji kewajaran|uji tantangan")
        self.assertNotRegex(
            lower_predictive,
            r"hasil backtesting|uji balik statistik|akurasi historis|validated forecast|statistical backtest",
        )

    def test_specialist_agents_generate_bounded_briefings_from_internal_datasets(self):
        briefing = FeedbackProposalTeam().run(
            self.engine,
            self.dataframe,
            self.timeframe,
            self.macro_trends,
        )

        self.assertEqual(
            [item["role"] for item in briefing["specialists"]],
            [
                "Data Steward",
                "Rating Analyst",
                "Voice-of-Customer Analyst",
                "External Context Analyst",
                "Action Planner",
            ],
        )
        self.assertTrue(all(item["finding"].strip() for item in briefing["specialists"]))
        self.assertIn("rating", briefing["sources_used"])
        self.assertIn("komentar", briefing["sources_used"])
        self.assertIn("osint", briefing["sources_used"])
        self.assertLessEqual(len(briefing["manager_summary"]), 420)

    def test_specialist_briefing_exposes_evidence_confidence_and_qa_guardrails(self):
        briefing = FeedbackProposalTeam().run(
            self.engine,
            self.dataframe,
            self.timeframe,
            self.macro_trends,
        )

        self.assertIn("evidence_ledger", briefing)
        self.assertIn("qa_review", briefing)
        self.assertTrue(briefing["evidence_ledger"])
        self.assertIn(briefing["confidence"], {"Tinggi", "Sedang", "Rendah"})
        self.assertTrue(all(item["evidence_type"] for item in briefing["evidence_ledger"]))
        self.assertTrue(all(item["source"] for item in briefing["evidence_ledger"]))
        self.assertTrue(all(item["confidence"] in {"Tinggi", "Sedang", "Rendah"} for item in briefing["specialists"]))
        self.assertTrue(all("temuan evaluasi" in item for item in briefing["qa_review"]))

    def test_hidden_agent_desk_uses_one_model_across_role_separated_passes(self):
        class FakeModelClient:
            def __init__(self):
                self.calls = []

            def chat(self, **kwargs):
                self.calls.append(kwargs)
                return {"message": {"content": '{"finding":"cukup","implication":"lanjut","confidence":"Sedang"}'}}

        client = FakeModelClient()
        desk = HiddenAgentDesk(model_name="same-model:latest", model_client=client, mode="ollama")
        briefing = FeedbackProposalTeam(agent_desk=desk).run(
            self.engine,
            self.dataframe,
            self.timeframe,
            self.macro_trends,
            score_engine="experience_index",
        )

        passes = briefing["agent_desk"]["passes"]
        self.assertEqual(len(client.calls), len(passes))
        self.assertTrue(all(call["model"] == "same-model:latest" for call in client.calls))
        self.assertEqual(len({item["model"] for item in passes}), 1)
        self.assertEqual(
            [item["role"] for item in passes],
            [item["role"] for item in briefing["specialists"]],
        )
        self.assertTrue(all("Return JSON with keys" in item["prompt_contract"] for item in passes))
        self.assertTrue(all(item["model_status"] == "completed" for item in passes))

    def test_hidden_agent_desk_produces_deterministic_ledger_and_final_quality_gate(self):
        briefing_one = FeedbackProposalTeam().run(
            self.engine,
            self.dataframe,
            self.timeframe,
            self.macro_trends,
            score_engine="experience_index",
        )
        briefing_two = FeedbackProposalTeam().run(
            self.engine,
            self.dataframe,
            self.timeframe,
            self.macro_trends,
            score_engine="experience_index",
        )

        ledger_one = briefing_one["evidence_ledger"]
        ledger_two = briefing_two["evidence_ledger"]
        self.assertEqual(ledger_one, ledger_two)
        self.assertTrue(all(item["evidence_id"] for item in ledger_one))
        self.assertEqual(
            [(item["evidence_type"], item["evidence_id"]) for item in ledger_one],
            sorted((item["evidence_type"], item["evidence_id"]) for item in ledger_one),
        )
        gate = briefing_one["agent_desk"]["final_quality_gate"]
        editor = briefing_one["agent_desk"]["editor_review"]
        self.assertTrue(gate["passes"], briefing_one["agent_desk"])
        self.assertTrue(editor["reader_safe"], editor)
        self.assertTrue(editor["ledger_complete"], editor)

    def test_specialist_briefing_adds_audit_trend_contradiction_and_prediction_controls(self):
        briefing = FeedbackProposalTeam().run(
            self.engine,
            self.dataframe,
            self.timeframe,
            self.macro_trends,
            score_engine="experience_index",
        )

        self.assertIn("audit_trail", briefing)
        self.assertIn("contradiction_review", briefing)
        self.assertIn("trend_review", briefing)
        self.assertIn("prediction_review", briefing)
        self.assertEqual(briefing["audit_trail"]["timeframe"], self.timeframe)
        self.assertEqual(briefing["audit_trail"]["score_engine"], "experience_index")
        self.assertGreater(briefing["audit_trail"]["raw_response_count"], 0)
        self.assertIn("rating_text_alignment", briefing["contradiction_review"])
        self.assertIn(briefing["contradiction_review"]["severity"], {"Rendah", "Sedang", "Tinggi"})
        self.assertIn("comparison_period", briefing["trend_review"])
        self.assertIn("rating_delta", briefing["trend_review"])
        self.assertFalse(briefing["prediction_review"]["statistical_forecast"])
        self.assertIn("early warning", briefing["prediction_review"]["method"].lower())

    def test_executive_summary_hides_specialist_workflow_labels(self):
        snapshot = self.engine.build_executive_snapshot(
            self.timeframe,
            self.notes,
            score_engine="experience_index",
        )

        self.assertIn("### Kesimpulan Utama", snapshot)
        self.assertIn("### Alasan Utama", snapshot)
        self.assertIn("### Keputusan yang Diminta", snapshot)
        for forbidden in [
            "Review Tim Analis Internal",
            "Data Steward",
            "Rating Analyst",
            "Voice-of-Customer Analyst",
            "Confidence Desk",
            "Evidence Ledger",
            "QA Guardrail",
            "Report Audit Trail",
            "Contradiction Check",
            "Historical Trend Desk",
            "Prediction Boundary",
        ]:
            self.assertNotIn(forbidden, snapshot)


if __name__ == "__main__":
    unittest.main()
