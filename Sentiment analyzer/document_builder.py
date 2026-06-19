import io
import logging
import re
import textwrap
from datetime import datetime

from bs4 import BeautifulSoup, NavigableString, Tag
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips
import markdown
import matplotlib
import matplotlib.patches as patches
import matplotlib.pyplot as plt

from config import DEFAULT_COLOR, WRITER_FIRM_NAME
from editorial_intelligence import compact_sentence

matplotlib.use("Agg")

logger = logging.getLogger(__name__)

def append_field(paragraph, instruction):
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    field_instruction = OxmlElement("w:instrText")
    field_instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    field_instruction.text = instruction
    field_separator = OxmlElement("w:fldChar")
    field_separator.set(qn("w:fldCharType"), "separate")
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, field_instruction, field_separator, field_end])

def enable_update_fields_on_open(doc):
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

class StyleEngine:
    @staticmethod
    def _configure_text_style(style, font_name, font_size, color, bold=False):
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor(*color)
        style.font.bold = bold

    @staticmethod
    def _configure_paragraph_format(paragraph_format, before=0, after=6, line_spacing=1.08):
        paragraph_format.space_before = Pt(before)
        paragraph_format.space_after = Pt(after)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = line_spacing

    @staticmethod
    def apply_document_styles(doc, theme_color):
        for section in doc.sections:
            section.top_margin = Cm(2.2)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.35)
            section.right_margin = Cm(2.35)
            footer_paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
            footer_paragraph.clear()
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            footer_run = footer_paragraph.add_run("SANGAT RAHASIA | Laporan Eksekutif Inixindo Jogja | Halaman ")
            footer_run.font.name, footer_run.font.size, footer_run.font.color.rgb = "Calibri", Pt(9), RGBColor(128, 128, 128)
            append_field(footer_paragraph, "PAGE")

        normal_style = doc.styles["Normal"]
        StyleEngine._configure_text_style(normal_style, "Calibri", 11, (33, 37, 41))
        normal_format = normal_style.paragraph_format
        normal_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        StyleEngine._configure_paragraph_format(normal_format, after=6, line_spacing=1.08)

        heading_1 = doc.styles["Heading 1"]
        StyleEngine._configure_text_style(heading_1, "Calibri", 16, theme_color, True)
        StyleEngine._configure_paragraph_format(heading_1.paragraph_format, before=16, after=6, line_spacing=1.0)
        heading_1.paragraph_format.keep_with_next = True

        heading_2 = doc.styles["Heading 2"]
        StyleEngine._configure_text_style(heading_2, "Calibri", 13, (0, 0, 0), True)
        StyleEngine._configure_paragraph_format(heading_2.paragraph_format, before=10, after=4, line_spacing=1.0)
        heading_2.paragraph_format.keep_with_next = True

        heading_3 = doc.styles["Heading 3"]
        StyleEngine._configure_text_style(heading_3, "Calibri", 12, (54, 54, 54), True)
        StyleEngine._configure_paragraph_format(heading_3.paragraph_format, before=8, after=3, line_spacing=1.0)
        heading_3.paragraph_format.keep_with_next = True

        for list_style_name in ("List Bullet", "List Bullet 2", "List Bullet 3", "List Number", "List Number 2", "List Number 3"):
            if list_style_name in doc.styles:
                list_style = doc.styles[list_style_name]
                StyleEngine._configure_text_style(list_style, "Calibri", 11, (33, 37, 41))
                StyleEngine._configure_paragraph_format(list_style.paragraph_format, after=3, line_spacing=1.0)

class ChartEngine:
    @staticmethod
    def _get_plt_color(theme_color): return tuple(channel / 255 for channel in theme_color)

    @staticmethod
    def _parse_chart_points(raw_data):
        labels, values = [], []
        for pair in raw_data.split(";"):
            if "," not in pair: continue
            label, value = pair.split(",", maxsplit=1)
            cleaned_value = re.sub(r"[^\d.\-]", "", value)
            if cleaned_value:
                labels.append(label.strip())
                values.append(float(cleaned_value))
        return labels, values

    @staticmethod
    def create_bar_chart(data_str, theme_color):
        try:
            parts = data_str.split("|")
            title_str, ylabel_str, raw_data = [p.strip() for p in parts] if len(parts) == 3 else ("Sentimen Makro", "Persentase", data_str)
            labels, values = ChartEngine._parse_chart_points(raw_data)
            if not labels: return None

            fig, axis = plt.subplots(figsize=(7, 4.5))
            axis.bar(labels, values, color=ChartEngine._get_plt_color(theme_color), alpha=0.9, width=0.5)
            axis.set_title(title_str, fontsize=12, fontweight="bold", pad=20)
            axis.set_ylabel(ylabel_str, fontsize=10)
            axis.spines["top"].set_visible(False)
            axis.spines["right"].set_visible(False)

            image_stream = io.BytesIO()
            plt.savefig(image_stream, format="png", bbox_inches="tight", dpi=150)
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception as exc:
            logger.warning("Gagal membuat bar chart: %s", exc)
            return None

    @staticmethod
    def create_line_chart(data_str, theme_color):
        try:
            parts = data_str.split("|")
            title_str, ylabel_str, raw_data = [p.strip() for p in parts] if len(parts) == 3 else ("Tren", "Skor", data_str)
            labels, values = ChartEngine._parse_chart_points(raw_data)
            if not labels: return None

            color = ChartEngine._get_plt_color(theme_color)
            fig, axis = plt.subplots(figsize=(7, 4.2))
            axis.plot(labels, values, color=color, linewidth=2.4, marker="o", markersize=6)
            axis.fill_between(labels, values, min(values), color=color, alpha=0.12)
            axis.set_title(title_str, fontsize=12, fontweight="bold", pad=18)
            axis.set_ylabel(ylabel_str, fontsize=10)
            axis.grid(axis="y", alpha=0.25)
            axis.spines["top"].set_visible(False)
            axis.spines["right"].set_visible(False)
            for index, value in enumerate(values):
                axis.annotate(f"{value:g}", (index, value), textcoords="offset points", xytext=(0, 8), ha="center", fontsize=8)

            image_stream = io.BytesIO()
            plt.savefig(image_stream, format="png", bbox_inches="tight", dpi=150)
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception as exc:
            logger.warning("Gagal membuat line chart: %s", exc)
            return None

    @staticmethod
    def create_pie_chart(data_str, theme_color):
        try:
            parts = data_str.split("|")
            title_str, raw_data = [p.strip() for p in parts] if len(parts) == 2 else ("Distribusi", data_str)
            labels, values = ChartEngine._parse_chart_points(raw_data)
            if not labels: return None

            palette = [ChartEngine._get_plt_color(theme_color), (0.85, 0.35, 0.25), (0.20, 0.45, 0.70), (0.35, 0.65, 0.45), (0.75, 0.60, 0.25)]
            fig, axis = plt.subplots(figsize=(6.2, 4.8))
            axis.pie(values, labels=labels, autopct="%1.1f%%", startangle=90, colors=palette[: len(labels)], textprops={"fontsize": 9})
            axis.set_title(title_str, fontsize=12, fontweight="bold", pad=16)
            axis.axis("equal")

            image_stream = io.BytesIO()
            plt.savefig(image_stream, format="png", bbox_inches="tight", dpi=150)
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception as exc:
            logger.warning("Gagal membuat pie chart: %s", exc)
            return None

    @staticmethod
    def create_flowchart(data_str, theme_color):
        try:
            steps = ["\n".join(textwrap.wrap(step.strip(), width=18)) for step in data_str.split("->") if step.strip()]
            if len(steps) < 2: return None

            fig, axis = plt.subplots(figsize=(8, 3))
            axis.axis("off")
            x_positions = [index * 2.5 for index in range(len(steps))]

            for index in range(len(steps) - 1):
                axis.annotate("", xy=(x_positions[index + 1] - 1.0, 0.5), xytext=(x_positions[index] + 1.0, 0.5), arrowprops={"arrowstyle": "-|>", "lw": 1.5})

            for index, step in enumerate(steps):
                box = patches.FancyBboxPatch((x_positions[index] - 1.0, 0.1), 2.0, 0.8, boxstyle="round,pad=0.1", fc=ChartEngine._get_plt_color(theme_color), alpha=0.9)
                axis.add_patch(box)
                axis.text(x_positions[index], 0.5, step, ha="center", va="center", size=9, color="white", fontweight="bold")

            axis.set_xlim(-1.2, (len(steps) - 1) * 2.5 + 1.2)
            axis.set_ylim(0, 1)

            image_stream = io.BytesIO()
            plt.savefig(image_stream, format="png", bbox_inches="tight", dpi=200, transparent=True)
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception as exc:
            logger.warning("Gagal membuat flowchart: %s", exc)
            return None

class DocumentBuilder:
    LIST_STYLES = {"ul": ["List Bullet", "List Bullet 2", "List Bullet 3"], "ol": ["List Number", "List Number 2", "List Number 3"]}
    LIST_INDENTS = (
        (Twips(720), Twips(-360)),
        (Twips(1080), Twips(-360)),
        (Twips(1440), Twips(-360)),
    )
    DEFAULT_TOC_ITEMS = (
        "Ringkasan Eksekutif",
        "BAB I - Analitik Deskriptif dan Tata Kelola Umpan Balik",
        "BAB II - Analitik Diagnostik dan Akar Masalah",
        "BAB III - Analitik Prediktif dan Peringatan Dini",
        "BAB IV - Analitik Preskriptif dan Rencana Aksi",
        "BAB V - Kesiapan Implementasi",
    )

    SOURCE_REPLACEMENTS = (
        (r"\bNama\s+Perusahaan\s+Klien\s*:", "catatan klien:"),
        (r"\bReferenceAccount\s+mencatat\b", "catatan klien menunjukkan"),
        (r"\bsource\s*=\s*(?:https?://\S+|/api/[A-Za-z0-9_./-]+)?", ""),
        (r"\bdataset[_\s-]*code\s*=\s*ConsultantProjectExpertHistory\b", "riwayat pengalaman konsultan"),
        (r"\bConsultantProjectExpertHistory\b", "riwayat pengalaman konsultan"),
        (r"\bDirangkum\s+dari\s+sumber[^:]*:\s*", "Berdasarkan catatan pendukung yang sudah dipadatkan: "),
        (r"\bProblem\s*,\s*Opportunity\s*,\s*Directive\b", "kebutuhan prioritas yang perlu dipertegas"),
        (r"\bReferenceAccount\b", "catatan klien"),
        (r"\bPain\s+Points\b", "titik masalah"),
        (r"\bS\s*T\s*R\s*I\s*C\s*T\s*L\s*Y\s+C\s*O\s*N\s*F\s*I\s*D\s*E\s*N\s*T\s*I\s*A\s*L\b", "SANGAT RAHASIA"),
        (r"HOLISTIC CUSTOMER EXPERIENCE REPORT", "LAPORAN MENYELURUH PENGALAMAN PELANGGAN"),
        (r"Prepared for Executive Board by:", "Disiapkan untuk Dewan Eksekutif oleh:"),
        (r"EXECUTIVE SUMMARY", "Ringkasan Eksekutif"),
        (r"\bBLUF\b", "Inti Keputusan"),
        (r"Key Findings", "Temuan Utama"),
        (r"Recommendation", "Rekomendasi"),
        (r"DESCRIPTIVE ANALYTICS\s*&\s*FEEDBACK GOVERNANCE", "ANALITIK DESKRIPTIF DAN TATA KELOLA UMPAN BALIK"),
        (r"DIAGNOSTIC ANALYTICS", "ANALITIK DIAGNOSTIK"),
        (r"PREDICTIVE ANALYTICS", "ANALITIK PREDIKTIF"),
        (r"PRESCRIPTIVE ANALYTICS", "ANALITIK PRESKRIPTIF"),
        (r"Semua Data APIDog \(tanggal tidak tersedia\)", "Seluruh Periode Evaluasi"),
        (r"Semua Data sistem data evaluasi \(tanggal tidak tersedia\)", "Seluruh Periode Evaluasi"),
        (r"APIDog", "sistem data evaluasi"),
        (r"Internal API\\s*/\\s*feedback cache", "basis evaluasi layanan"),
        (r"Internal API", "basis evaluasi layanan"),
        (r"internal_api", "evaluasi layanan"),
        (r"internal facts", "temuan evaluasi"),
        (r"Internal facts", "Temuan evaluasi"),
        (r"feedback cache", "basis evaluasi layanan"),
        (r"Dataset Spesialis", "Fokus Bukti"),
        (r"OSINT benchmark", "pembanding eksternal"),
        (r"OSINT", "konteks eksternal"),
        (r"Jejak sumber yang dipakai", "Jejak bukti yang dipakai"),
        (r"BRAND EQUITY\s*\(Mengapa Inixindo Jogja menjadi pilihan\?\)\s*Pilih\s+\d+\s+Bintang\s+untuk\s+mengisi", "Reputasi dan alasan memilih Inixindo"),
        (r"\bPilih\s+\d+\s+Bintang\s+untuk\s+mengisi\b", ""),
        (r"\bCustomer\s+journey\b", "perjalanan pelanggan"),
        (r"\bcustomer\s+journey\b", "perjalanan pelanggan"),
        (r"\bScore\s+Engine\b", "Perspektif Skor"),
        (r"\bscore\s+engine\b", "perspektif skor"),
        (r"\bExperience\s+Index\b", "Experience Index"),
        (r"\bLearning\s+Score\b", "Skor Pembelajaran"),
        (r"\bService\s+Score\b", "Skor Layanan"),
        (r"\bFacility\s+Score\b", "Skor Fasilitas"),
        (r"\bfeedback\b", "umpan balik"),
        (r"\brating\b", "penilaian"),
        (r"\bdashboard\b", "dasbor"),
        (r"\binsight\b", "wawasan"),
        (r"\bstakeholder\b", "pemangku kepentingan"),
        (r"\bbenchmark\b", "pembanding"),
        (r"\bowner\b", "penanggung jawab"),
        (r"\breview\b", "tinjauan"),
        (r"\bquick wins?\b", "perbaikan cepat"),
        (r"\bapproval gate\b", "gerbang persetujuan"),
        (r"\bpain point\b", "titik keluhan"),
        (r"\bfield\b", "kolom"),
        (r"\brecord\b", "catatan"),
        (r"\bdelivery\b", "pelaksanaan"),
        (r"\boutcome\b", "dampak hasil"),
        (r"\bclass_report\b", "laporan kelas"),
    )

    @staticmethod
    def reader_facing_text(raw_text):
        clean_text = str(raw_text or "")
        for pattern, replacement in DocumentBuilder.SOURCE_REPLACEMENTS:
            clean_text = re.sub(pattern, replacement, clean_text, flags=re.IGNORECASE)
        presentation_replacements = (
            (r"Ringkasan Cakupan umpan balik", "Ringkasan Cakupan Umpan Balik"),
            (r"Distribusi Sentimen, penilaian", "Distribusi Sentimen, Penilaian"),
            (r"penilaian Rata-rata", "Rata-rata Penilaian"),
            (r"Rating Rata-rata", "Rata-rata Penilaian"),
            (r"penanggung jawab Utama", "Penanggung Jawab Utama"),
            (r"Distribusi pemangku kepentingan", "Distribusi Pemangku Kepentingan"),
            (r"pemangku kepentingan dengan volume", "Pemangku kepentingan dengan volume"),
            (r"Akar Masalah Utama dan titik keluhan", "Akar Masalah Utama dan Titik Keluhan"),
            (r"Tahap perjalanan pelanggan", "Tahap Perjalanan Pelanggan"),
            (r"tahapan perjalanan pelanggan", "tahapan perjalanan pelanggan"),
            (r"Penjelasan Perhitungan Indeks Pengalaman", "Penjelasan Perhitungan Indeks Pengalaman"),
            (r"umpan balik Score\\.xlsx", "Parameter Skor Evaluasi.xlsx"),
            (r"\.kondisi", ". Kondisi"),
            (r"\bfalse comfort\b", "rasa aman semu"),
            (r"\bChallenge Check\b", "Uji Kewajaran"),
            (r"\bchallenge check\b", "uji kewajaran"),
            (r"\bProspective Validation\b", "Validasi Periode Berikutnya"),
            (r"\bprospective validation\b", "validasi periode berikutnya"),
            (r"\bearly-warning\b", "peringatan dini"),
            (r"\bearly warning\b", "peringatan dini"),
            (r"\btouchpoint\b", "titik sentuh"),
            (r"\bjourney\b", "perjalanan"),
            (r"\bBelum terpetakan\b", "Perlu pemetaan data"),
            (r"\bTidak terklasifikasi\b", "Perlu klasifikasi"),
        )
        for pattern, replacement in presentation_replacements:
            clean_text = re.sub(pattern, replacement, clean_text)
        clean_text = re.sub(r"\b(endpoint|schema|sync status|source-of-truth|Waitress|thread|runtime)\b", "kesiapan operasional", clean_text, flags=re.IGNORECASE)
        clean_text = re.sub(r"\s+", " ", clean_text) if "\n" not in clean_text else "\n".join(re.sub(r"[ \t]+", " ", line).rstrip() for line in clean_text.splitlines())
        return clean_text.strip()

    @staticmethod
    def _format_paragraph(paragraph, alignment=None, before=0, after=6, line_spacing=1.08):
        if alignment is not None:
            paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = line_spacing
        return paragraph

    @staticmethod
    def _add_spacer(doc, height=6):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(height)
        paragraph.paragraph_format.line_spacing = 1.0
        return paragraph

    @staticmethod
    def _add_picture(doc, image, width):
        paragraph = doc.add_paragraph()
        DocumentBuilder._format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=8)
        paragraph.add_run().add_picture(image, width=width)

    @staticmethod
    def _set_cell_shading(cell, fill):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def _format_table(table):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        column_count = len(table.columns)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for paragraph in cell.paragraphs:
                    DocumentBuilder._format_paragraph(
                        paragraph,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        after=4,
                        line_spacing=1.08,
                    )
                    for run in paragraph.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(8.8 if column_count >= 5 else 9.2)
                if row_index == 0:
                    DocumentBuilder._set_cell_shading(cell, "F3F4F6")
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    @staticmethod
    def _resolve_list_style(doc, list_tag, level):
        style_candidates = DocumentBuilder.LIST_STYLES.get(list_tag, ["List Bullet"])
        return style_candidates[min(level, len(style_candidates) - 1)] if style_candidates[min(level, len(style_candidates) - 1)] in doc.styles else style_candidates[0]

    @staticmethod
    def _next_numbering_id(numbering, tag_name, attr_name):
        values = []
        for element in numbering.findall(qn(tag_name)):
            raw_value = element.get(qn(attr_name))
            if str(raw_value).isdigit():
                values.append(int(raw_value))
        return max(values, default=0) + 1

    @staticmethod
    def _append_level(parent, level, list_tag):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if list_tag == "ol" else "bullet")
        lvl.append(num_fmt)

        lvl_text = OxmlElement("w:lvlText")
        if list_tag == "ol":
            lvl_text.set(qn("w:val"), f"%{level + 1}.")
        else:
            lvl_text.set(qn("w:val"), ("•", "o", "▪")[min(level, 2)])
        lvl.append(lvl_text)

        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        left_indent, hanging_indent = DocumentBuilder.LIST_INDENTS[min(level, len(DocumentBuilder.LIST_INDENTS) - 1)]
        ind.set(qn("w:left"), str(left_indent.twips))
        ind.set(qn("w:hanging"), str(abs(hanging_indent.twips)))
        p_pr.append(ind)
        lvl.append(p_pr)
        parent.append(lvl)

    @staticmethod
    def _create_numbering_instance(doc, list_tag):
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        abstract_id = DocumentBuilder._next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
        num_id = DocumentBuilder._next_numbering_id(numbering, "w:num", "w:numId")

        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
        multi_level = OxmlElement("w:multiLevelType")
        multi_level.set(qn("w:val"), "hybridMultilevel")
        abstract_num.append(multi_level)
        for level in range(len(DocumentBuilder.LIST_INDENTS)):
            DocumentBuilder._append_level(abstract_num, level, list_tag)
        numbering.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num_id)
        numbering.append(num)
        return num_id

    @staticmethod
    def _apply_numbering(paragraph, num_id, level):
        p_pr = paragraph._p.get_or_add_pPr()
        existing_num_pr = p_pr.find(qn("w:numPr"))
        if existing_num_pr is not None:
            p_pr.remove(existing_num_pr)

        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num])
        p_pr.append(num_pr)

        left_indent, hanging_indent = DocumentBuilder.LIST_INDENTS[min(level, len(DocumentBuilder.LIST_INDENTS) - 1)]
        paragraph.paragraph_format.left_indent = left_indent
        paragraph.paragraph_format.first_line_indent = hanging_indent

    @staticmethod
    def _append_inline_runs(paragraph, node, bold=False, italic=False):
        if isinstance(node, NavigableString):
            if str(node):
                run = paragraph.add_run(str(node))
                run.bold, run.italic = bold, italic
            return
        if not isinstance(node, Tag): return
        if node.name == "br":
            paragraph.add_run("\n")
            return
        next_bold = bold or node.name in {"strong", "b"}
        next_italic = italic or node.name in {"em", "i"}
        for child in node.children:
            DocumentBuilder._append_inline_runs(paragraph, child, bold=next_bold, italic=next_italic)

    @staticmethod
    def _render_list(doc, list_node, level=0, list_state=None):
        if list_state is None:
            list_state = {}
        if list_node.name not in list_state:
            list_state[list_node.name] = DocumentBuilder._create_numbering_instance(doc, list_node.name)
        style_name = DocumentBuilder._resolve_list_style(doc, list_node.name, level)
        for list_item in list_node.find_all("li", recursive=False):
            paragraph = doc.add_paragraph(style=style_name)
            DocumentBuilder._format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=3, line_spacing=1.0)
            DocumentBuilder._apply_numbering(paragraph, list_state[list_node.name], level)
            for child in list_item.contents:
                if isinstance(child, Tag) and child.name in {"ul", "ol"}: continue
                DocumentBuilder._append_inline_runs(paragraph, child)
            for nested in [child for child in list_item.contents if isinstance(child, Tag) and child.name in {"ul", "ol"}]:
                DocumentBuilder._render_list(doc, nested, level=level + 1, list_state=list_state)

    @staticmethod
    def _render_table(doc, table_node):
        rows = table_node.find_all("tr")
        if not rows: return
        column_count = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
        table = doc.add_table(rows=0, cols=column_count)
        table.style = "Table Grid"

        for row_index, row_node in enumerate(rows):
            cells = row_node.find_all(["th", "td"], recursive=False)
            table_row = table.add_row().cells
            for col_index in range(column_count):
                if col_index < len(cells):
                    table_row[col_index].text = compact_sentence(
                        cells[col_index].get_text(" ", strip=True),
                        max_words=22 if column_count >= 4 else 30,
                    )
                if row_index == 0 and col_index < len(cells) and cells[col_index].name == "th":
                    for run in table_row[col_index].paragraphs[0].runs:
                        run.bold = True
        DocumentBuilder._format_table(table)
        DocumentBuilder._add_spacer(doc, height=4)

    @staticmethod
    def parse_html_to_docx(doc, html_content):
        soup = BeautifulSoup(html_content, "html.parser")
        for element in soup.contents:
            if not isinstance(element, Tag):
                continue
            if element.name in {"h1", "h2", "h3"}:
                doc.add_heading(element.get_text(" ", strip=True), level=int(element.name[1]))
                continue
            if element.name == "p":
                paragraph = doc.add_paragraph()
                DocumentBuilder._format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
                for child in element.children:
                    DocumentBuilder._append_inline_runs(paragraph, child)
                continue
            if element.name in {"ul", "ol"}:
                DocumentBuilder._render_list(doc, element, level=0)
                continue
            if element.name == "table":
                DocumentBuilder._render_table(doc, element)

    @staticmethod
    def process_content(doc, raw_text, theme_color=DEFAULT_COLOR):
        raw_text = DocumentBuilder.reader_facing_text(raw_text)
        clean_lines = []

        def flush_markdown():
            if not any(line.strip() for line in clean_lines):
                clean_lines.clear()
                return
            html = markdown.markdown("\n".join(clean_lines), extensions=["tables"])
            DocumentBuilder.parse_html_to_docx(doc, html)
            clean_lines.clear()

        for line in raw_text.split("\n"):
            stripped_line = line.strip()
            if stripped_line.startswith("[[CHART:") and stripped_line.endswith("]]"):
                flush_markdown()
                image = ChartEngine.create_bar_chart(stripped_line.replace("[[CHART:", "").replace("]]", "").strip(), theme_color)
                if image:
                    DocumentBuilder._add_picture(doc, image, width=Inches(5.5))
                continue
            if stripped_line.startswith("[[LINE:") and stripped_line.endswith("]]"):
                flush_markdown()
                image = ChartEngine.create_line_chart(stripped_line.replace("[[LINE:", "").replace("]]", "").strip(), theme_color)
                if image:
                    DocumentBuilder._add_picture(doc, image, width=Inches(5.5))
                continue
            if stripped_line.startswith("[[PIE:") and stripped_line.endswith("]]"):
                flush_markdown()
                image = ChartEngine.create_pie_chart(stripped_line.replace("[[PIE:", "").replace("]]", "").strip(), theme_color)
                if image:
                    DocumentBuilder._add_picture(doc, image, width=Inches(5.4))
                continue
            if stripped_line.startswith("[[FLOW:") and stripped_line.endswith("]]"):
                flush_markdown()
                image = ChartEngine.create_flowchart(stripped_line.replace("[[FLOW:", "").replace("]]", "").strip(), theme_color)
                if image:
                    DocumentBuilder._add_picture(doc, image, width=Inches(6.5))
                continue
            clean_lines.append(line)

        flush_markdown()

    @staticmethod
    def add_table_of_contents(doc, items=None):
        enable_update_fields_on_open(doc)
        doc.add_heading("DAFTAR ISI", level=1)
        toc_paragraph = doc.add_paragraph()
        DocumentBuilder._format_paragraph(toc_paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=8)
        append_field(toc_paragraph, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_page_break()

    @staticmethod
    def create_cover(doc, timeframe, theme_color=DEFAULT_COLOR):
        StyleEngine.apply_document_styles(doc, theme_color)
        for _ in range(4):
            DocumentBuilder._add_spacer(doc, height=8)

        confidentiality = doc.add_paragraph("S A N G A T   R A H A S I A")
        DocumentBuilder._format_paragraph(confidentiality, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=10)
        confidentiality.runs[0].font.size, confidentiality.runs[0].font.color.rgb, confidentiality.runs[0].font.bold = Pt(10), RGBColor(128, 128, 128), True

        DocumentBuilder._add_spacer(doc, height=8)
        report_title = doc.add_paragraph("LAPORAN MENYELURUH PENGALAMAN PELANGGAN")
        DocumentBuilder._format_paragraph(report_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=4, line_spacing=1.0)
        report_title.runs[0].font.name, report_title.runs[0].font.size = "Calibri", Pt(20)

        company_name = doc.add_paragraph("INIXINDO JOGJA")
        DocumentBuilder._format_paragraph(company_name, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=14, line_spacing=1.0)
        company_name.runs[0].font.name, company_name.runs[0].font.bold, company_name.runs[0].font.size, company_name.runs[0].font.color.rgb = "Calibri", True, Pt(32), RGBColor(*theme_color)

        period_text = doc.add_paragraph(f"Periode Evaluasi Laporan: {DocumentBuilder.reader_facing_text(timeframe)}")
        DocumentBuilder._format_paragraph(period_text, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=4, line_spacing=1.0)
        period_text.runs[0].font.size = Pt(13)

        generated_text = doc.add_paragraph(f"Tanggal Generasi: {datetime.now().strftime('%d %B %Y')}")
        DocumentBuilder._format_paragraph(generated_text, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=18, line_spacing=1.0)
        generated_text.runs[0].font.size, generated_text.runs[0].font.color.rgb = Pt(11), RGBColor(128, 128, 128)

        for _ in range(6):
            DocumentBuilder._add_spacer(doc, height=8)
        prepared_for = doc.add_paragraph(f"Disiapkan untuk Dewan Eksekutif oleh:\n{WRITER_FIRM_NAME}")
        DocumentBuilder._format_paragraph(prepared_for, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=0, line_spacing=1.0)
        prepared_for.runs[0].font.bold = True

        doc.add_page_break()
        DocumentBuilder.add_table_of_contents(doc)
