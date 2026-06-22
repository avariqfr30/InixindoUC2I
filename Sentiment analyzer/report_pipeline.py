import concurrent.futures
import logging
import threading
import time
from dataclasses import dataclass

from docx import Document

from config import DEFAULT_COLOR, DEFAULT_SCORE_ENGINE, SCORE_ENGINE_PROFILES
from document_builder import DocumentBuilder
from osint_research import Researcher
from report_agents import FeedbackProposalTeam
from report_analytics import FeedbackAnalyticsEngine
from report_evidence import ContextIntelligenceDesk
from report_factuality import ReportFactRegistry
from report_planning import FeedbackSectionPlanner
from report_quality import ReportQualityValidator
from editorial_intelligence import repair_feedback_document_spine
from feedback_deliberation import FeedbackDeliberationBuilder
from timeframe_filters import readable_timeframe_label

logger = logging.getLogger(__name__)

_REPORT_ORCHESTRATION_EXECUTOR = concurrent.futures.ThreadPoolExecutor(max_workers=2)
_REPORT_WRITING_EXECUTOR = concurrent.futures.ThreadPoolExecutor(max_workers=2)
_TIMED_PIPELINE_STAGES = (
    "research",
    "analysis",
    "narrative",
    "writing",
    "preflight",
    "render",
    "quality",
)


class _PipelineStageTimings:
    def __init__(self):
        self._durations = {}
        self._lock = threading.Lock()

    def run(self, stage_name, operation, *args):
        started_at = time.perf_counter()
        try:
            return operation(*args)
        finally:
            self.finish(stage_name, started_at)

    def track_future(self, stage_name, future, started_at):
        add_done_callback = getattr(future, "add_done_callback", None)
        if callable(add_done_callback):
            add_done_callback(lambda _completed: self.finish(stage_name, started_at))

    def finish(self, stage_name, started_at):
        duration = max(0.0, time.perf_counter() - started_at)
        with self._lock:
            self._durations.setdefault(stage_name, duration)

    def log_summary(self):
        with self._lock:
            summary = " ".join(
                f"{stage_name}={self._durations[stage_name]:.6f}"
                for stage_name in _TIMED_PIPELINE_STAGES
            )
        logger.info("Report pipeline stage timings (seconds): %s", summary)


@dataclass(frozen=True)
class ReportRequestContext:
    timeframe: str
    notes: str = ""
    sentiment: str = "all"
    segment: str = "all"
    score_engine: str = DEFAULT_SCORE_ENGINE

    @property
    def score_profile(self):
        return SCORE_ENGINE_PROFILES.get(
            self.score_engine,
            SCORE_ENGINE_PROFILES[DEFAULT_SCORE_ENGINE],
        )

    @property
    def timeframe_label(self):
        return readable_timeframe_label(self.timeframe)


class ReportResearchStage:
    def __init__(self, executor=None):
        self.executor = executor

    def _lookup(self, context):
        return Researcher.get_macro_trends(
            context.timeframe_label,
            context.notes,
            context.score_profile["label"],
        )

    def submit(self, context, executor=None):
        selected_executor = executor or self.executor or _REPORT_ORCHESTRATION_EXECUTOR
        return selected_executor.submit(self._lookup, context)

    @staticmethod
    def resolve(future):
        try:
            return future.result(timeout=45)
        except Exception:
            logger.exception("OSINT macro trend lookup failed during report generation.")
            return "Tidak ada tren eksternal yang berhasil dimuat."

    def run(self, context):
        return self.resolve(self.submit(context))


class ReportAnalysisStage:
    def run(self, dataframe, context=None):
        analytics = FeedbackAnalyticsEngine(dataframe)
        if context is not None:
            analytics.prepare_report_analysis(
                context.timeframe,
                sentiment=context.sentiment,
                segment=context.segment,
                score_engine=context.score_engine,
            )
        return analytics


class ReportNarrativeStage:
    def __init__(self, planner=None):
        self.planner = planner or FeedbackSectionPlanner()

    def run(self, analytics, context, macro_trends):
        get_prepared = getattr(analytics, "get_prepared_report_analysis", None)
        prepared_analysis = (
            get_prepared(
                context.timeframe,
                sentiment=context.sentiment,
                segment=context.segment,
                score_engine=context.score_engine,
            )
            if callable(get_prepared)
            else None
        )
        osint_dossier = Researcher.build_osint_dossier(
            macro_trends,
            {
                "timeframe_label": context.timeframe_label,
                "timeframe": context.timeframe,
                "sentiment": context.sentiment,
                "segment": context.segment,
                "score_engine": context.score_engine,
                "score_engine_label": context.score_profile["label"],
            },
        )
        section_context = ContextIntelligenceDesk.build(
            dataframe=getattr(analytics, "full_df", None),
            notes=context.notes,
            timeframe=context.timeframe,
            sentiment=context.sentiment,
            segment=context.segment,
            score_engine=context.score_engine,
            macro_trends=macro_trends,
        )
        scoped_dataframe = (
            prepared_analysis.scoped_dataframe
            if prepared_analysis is not None
            else getattr(analytics, "full_df", None)
        )
        if prepared_analysis is None and callable(getattr(analytics, "_filter_view", None)):
            scoped_dataframe = analytics._filter_view(
                context.timeframe,
                sentiment=context.sentiment,
                segment=context.segment,
            )
        if prepared_analysis is not None:
            analysis_context = prepared_analysis.analysis_context
        elif callable(getattr(analytics, "_build_analysis_context", None)):
            analysis_context = analytics._build_analysis_context(
                scoped_dataframe,
                context.timeframe,
                context.sentiment,
                context.segment,
                context.score_engine,
            )
        else:
            analysis_context = {
                "score_profile": context.score_profile,
                "score_metrics": {},
                "dominant_journey": None,
            }
        if prepared_analysis is not None:
            governance = prepared_analysis.governance_summary
        elif callable(getattr(analytics, "_governance_summary", None)):
            governance = analytics._governance_summary(scoped_dataframe)
        else:
            governance = {
                "total_rows": int(section_context.get("row_count") or 0),
                "dimension_count": int(section_context.get("row_count") or 0),
                "rating_response_count": 0,
                "text_response_count": int(section_context.get("text_response_count") or 0),
                "completeness_pct": 0.0,
                "source_count": 0,
                "channel_count": 0,
            }
        contradiction_review = (
            prepared_analysis.contradiction_review
            if prepared_analysis is not None
            else FeedbackProposalTeam._contradiction_review(scoped_dataframe)
            if scoped_dataframe is not None
            else {
                "rating_text_alignment": "Belum dinilai pada jalur analitik terbatas.",
                "severity": "Rendah",
            }
        )
        trust_packet = ReportFactRegistry.build(
            scoped_dataframe,
            analysis_context,
            governance,
            contradiction_review,
            {
                "timeframe": context.timeframe,
                "timeframe_label": context.timeframe_label,
                "sentiment": context.sentiment,
                "segment": context.segment,
                "score_engine": context.score_engine,
                "external_context_ready": section_context.get("external_context_ready"),
            },
        )
        section_context["trust_packet"] = trust_packet
        section_context["data_version"] = trust_packet["snapshot_fingerprint"]
        section_context["row_count"] = int(governance.get("total_rows", 0) or 0)
        section_context["text_response_count"] = int(governance.get("text_response_count", 0) or 0)
        coverage_parts = [
            f"periode {context.timeframe_label}",
            f"{section_context['row_count']} respons terolah",
        ]
        if context.sentiment != "all":
            coverage_parts.append(f"filter sentimen {context.sentiment}")
        if context.segment != "all":
            coverage_parts.append(f"segmen {context.segment}")
        section_context["coverage_note"] = "Cakupan pembacaan memakai " + ", ".join(coverage_parts) + "."
        report_sections = analytics.build_report_sections(
            context.timeframe,
            section_context["focus_note"],
            macro_trends,
            sentiment=context.sentiment,
            segment=context.segment,
            score_engine=context.score_engine,
            section_context=section_context,
            prepared_analysis=prepared_analysis,
        )
        deliberation_builder = FeedbackDeliberationBuilder()
        document_contract = deliberation_builder.build(
            report_sections or [],
            {
                **section_context,
                "timeframe_label": context.timeframe_label,
                "timeframe": context.timeframe,
                "sentiment": context.sentiment,
                "segment": context.segment,
            },
            data_version=str(section_context.get("data_version") or ""),
        )
        planning_block = self.planner.build_prompt_block(
            sections=["Ringkasan Eksekutif", *[section.get("title", "") for section in report_sections or []]],
            context={
                "timeframe_label": context.timeframe_label,
                "timeframe": context.timeframe,
                "sentiment": context.sentiment,
                "segment": context.segment,
                "osint_dossier": osint_dossier,
                "row_count": section_context.get("row_count"),
                "text_response_count": section_context.get("text_response_count"),
                "external_context_ready": section_context.get("external_context_ready"),
                "insight_cards": section_context.get("insight_cards"),
                "data_version": section_context.get("data_version"),
                "document_contract": document_contract,
            },
        )
        for section in report_sections or []:
            section["_writing_plan"] = "\n".join(
                [planning_block, deliberation_builder.for_section(document_contract, section.get("id") or "")]
            )
            section["_document_contract"] = document_contract
        executive_snapshot = analytics.build_executive_snapshot(
            context.timeframe,
            section_context["focus_note"],
            sentiment=context.sentiment,
            segment=context.segment,
            score_engine=context.score_engine,
            macro_trends=macro_trends,
            report_sections=report_sections,
            section_context=section_context,
            prepared_analysis=prepared_analysis,
        )
        return executive_snapshot, report_sections, planning_block


class ReportWritingQualityStage:
    def __init__(self, editor=None):
        if editor is None:
            from writing_quality import ProtectedIndonesianEditor
            editor = ProtectedIndonesianEditor()
        self.editor = editor

    def _polish_section(self, section, planning_block):
        item = dict(section)
        item["content"] = self.editor.polish(
            item.get("content", ""),
            guidance=item.get("_writing_plan") or planning_block,
        )
        item.pop("_writing_plan", None)
        return item

    def run(self, executive_snapshot, report_sections, planning_block=""):
        snapshot_future = _REPORT_WRITING_EXECUTOR.submit(
            self.editor.polish,
            executive_snapshot,
            guidance=planning_block,
        )
        section_futures = [
            _REPORT_WRITING_EXECUTOR.submit(self._polish_section, section, planning_block)
            for section in report_sections or []
        ]
        futures = [snapshot_future, *section_futures]
        try:
            completed, _pending = concurrent.futures.wait(
                futures,
                return_when=concurrent.futures.FIRST_EXCEPTION,
            )
            failed_future = next(
                (future for future in completed if future.exception() is not None),
                None,
            )
            if failed_future is not None:
                failed_future.result()
            polished_snapshot = snapshot_future.result()
            polished_sections = [future.result() for future in section_futures]
        except Exception:
            for future in futures:
                future.cancel()
            raise
        polished_snapshot, polished_sections = repair_feedback_document_spine(polished_snapshot, polished_sections)
        return polished_snapshot, polished_sections


class ReportPreflightQualityStage:
    def run(self, executive_snapshot, report_sections):
        contract = next(
            (section.get("_document_contract") for section in report_sections or [] if section.get("_document_contract")),
            None,
        )
        appendix = FeedbackDeliberationBuilder.build_appendix_markdown(contract) if contract else ""
        result = ReportQualityValidator.evaluate_narrative(
            executive_snapshot,
            report_sections,
            deliberation_contract=contract,
            appendix_content=appendix,
        )
        if not result["passes"]:
            raise ValueError("Report narrative preflight failed: " + "; ".join(result["findings"]))
        return result


class DocumentRenderStage:
    def run(self, context, executive_snapshot, report_sections):
        document = Document()
        DocumentBuilder.create_cover(document, context.timeframe_label, DEFAULT_COLOR)
        document.add_heading("Ringkasan Eksekutif", level=1)
        DocumentBuilder.process_content(document, executive_snapshot, DEFAULT_COLOR)

        for section in report_sections:
            document.add_page_break()
            document.add_heading(section["title"], level=1)
            DocumentBuilder.process_content(document, section["content"], DEFAULT_COLOR)
        contract = next(
            (section.get("_document_contract") for section in report_sections or [] if section.get("_document_contract")),
            None,
        )
        if contract:
            document.add_page_break()
            appendix = FeedbackDeliberationBuilder.build_appendix_markdown(contract)
            DocumentBuilder.process_content(document, appendix, DEFAULT_COLOR)
        return document


DocumentAssemblyStage = DocumentRenderStage


class ReportQualityStage:
    def run(
        self,
        document,
        executive_snapshot,
        report_sections,
        analytics,
        dataframe,
        context,
        macro_trends,
        prepared_analysis=None,
    ):
        quality = ReportQualityValidator.evaluate(
            document,
            executive_snapshot,
            report_sections,
            context.score_profile["label"],
        )
        briefing = FeedbackProposalTeam().run(
            analytics,
            dataframe,
            context.timeframe,
            macro_trends=macro_trends,
            sentiment=context.sentiment,
            segment=context.segment,
            score_engine=context.score_engine,
            prepared_analysis=prepared_analysis,
        )
        quality["audit_trail"] = briefing.get("audit_trail", {})
        quality["confidence"] = briefing.get("confidence")
        quality["contradiction_review"] = briefing.get("contradiction_review", {})
        quality["trend_review"] = briefing.get("trend_review", {})
        quality["prediction_review"] = briefing.get("prediction_review", {})
        quality["agent_desk"] = briefing.get("agent_desk", {})
        contract = next(
            (section.get("_document_contract") for section in report_sections or [] if section.get("_document_contract")),
            {},
        )
        quality["document_deliberation"] = {
            "cache_key": contract.get("cache_key"),
            "accepted_claim_count": len(contract.get("claim_ledger") or []),
            "data_gap_count": len(contract.get("data_gap_register") or []),
            "appendix_sections": list((contract.get("appendix_manifest") or {}).keys()),
        }
        if not quality["verified_complete"]:
            logger.warning(
                "Generated report is below completeness target: %s",
                quality["summary"],
            )
        return quality


class ReportPipeline:
    def __init__(
        self,
        kb_instance,
        research_stage=None,
        analysis_stage=None,
        narrative_stage=None,
        document_stage=None,
        quality_stage=None,
        preflight_stage=None,
        writing_stage=None,
        orchestration_executor=None,
    ):
        self.kb = kb_instance
        self._uses_default_research_stage = research_stage is None
        self._uses_default_analysis_stage = analysis_stage is None
        self._uses_default_quality_stage = quality_stage is None
        self.research_stage = research_stage or ReportResearchStage()
        self.analysis_stage = analysis_stage or ReportAnalysisStage()
        self.narrative_stage = narrative_stage or ReportNarrativeStage()
        self.document_stage = document_stage or DocumentRenderStage()
        self.preflight_stage = preflight_stage or ReportPreflightQualityStage()
        self.writing_stage = writing_stage or ReportWritingQualityStage()
        self.quality_stage = quality_stage or ReportQualityStage()
        self.orchestration_executor = orchestration_executor or _REPORT_ORCHESTRATION_EXECUTOR

    def run(
        self,
        timeframe,
        notes="",
        sentiment="all",
        segment="all",
        score_engine=DEFAULT_SCORE_ENGINE,
    ):
        context = ReportRequestContext(timeframe, notes, sentiment, segment, score_engine)
        timings = _PipelineStageTimings()
        if self._uses_default_research_stage:
            research_started_at = time.perf_counter()
            research_future = self.research_stage.submit(
                context,
                executor=self.orchestration_executor,
            )
            timings.track_future("research", research_future, research_started_at)
        else:
            research_started_at = None
            research_future = self.orchestration_executor.submit(
                timings.run,
                "research",
                self.research_stage.run,
                context,
            )
        try:
            analysis_args = (
                (self.kb.df, context)
                if self._uses_default_analysis_stage
                else (self.kb.df,)
            )
            analytics = timings.run("analysis", self.analysis_stage.run, *analysis_args)
        except Exception:
            research_future.cancel()
            raise
        if self._uses_default_research_stage:
            macro_trends = self.research_stage.resolve(research_future)
            timings.finish("research", research_started_at)
        else:
            macro_trends = research_future.result()
        executive_snapshot, report_sections, planning_block = timings.run(
            "narrative",
            self.narrative_stage.run,
            analytics,
            context,
            macro_trends,
        )
        executive_snapshot, report_sections = timings.run(
            "writing",
            self.writing_stage.run,
            executive_snapshot,
            report_sections,
            planning_block,
        )
        preflight_quality = timings.run(
            "preflight",
            self.preflight_stage.run,
            executive_snapshot,
            report_sections,
        )
        document = timings.run(
            "render",
            self.document_stage.run,
            context,
            executive_snapshot,
            report_sections,
        )
        quality_args = (
            document,
            executive_snapshot,
            report_sections,
            analytics,
            self.kb.df,
            context,
            macro_trends,
        )
        if self._uses_default_quality_stage:
            get_prepared = getattr(analytics, "get_prepared_report_analysis", None)
            prepared_analysis = (
                get_prepared(
                    context.timeframe,
                    sentiment=context.sentiment,
                    segment=context.segment,
                    score_engine=context.score_engine,
                )
                if callable(get_prepared)
                else None
            )
            quality_args = (*quality_args, prepared_analysis)
        quality = timings.run("quality", self.quality_stage.run, *quality_args)
        quality["preflight"] = preflight_quality
        filename = (
            f"Inixindo_Feedback_Intelligence_Report_{context.score_profile['label']}_{context.timeframe_label}"
        ).replace(" ", "_")
        timings.log_summary()
        return document, filename, quality
