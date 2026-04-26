import concurrent.futures
import logging

from docx import Document

from config import DEFAULT_COLOR, DEFAULT_SCORE_ENGINE, SCORE_ENGINE_PROFILES
from document_builder import DocumentBuilder
from osint_research import Researcher
from report_analytics import FeedbackAnalyticsEngine
from report_quality import ReportQualityValidator

logger = logging.getLogger(__name__)

class ReportGenerator:
    def __init__(self, kb_instance):
        self.kb = kb_instance
        self.research_pool = concurrent.futures.ThreadPoolExecutor(max_workers=4)

    def run(self, timeframe, notes="", sentiment="all", segment="all", score_engine=DEFAULT_SCORE_ENGINE):
        logger.info(
            "Starting feedback intelligence report generation for timeframe=%s, sentiment=%s, segment=%s, score_engine=%s",
            timeframe,
            sentiment,
            segment,
            score_engine,
        )
        score_profile = SCORE_ENGINE_PROFILES.get(score_engine, SCORE_ENGINE_PROFILES[DEFAULT_SCORE_ENGINE])

        macro_future = self.research_pool.submit(Researcher.get_macro_trends, timeframe, notes, score_profile["label"])
        try:
            macro_trends = macro_future.result(timeout=45)
        except Exception:
            logger.exception("OSINT macro trend lookup failed during report generation.")
            macro_trends = "Tidak ada tren eksternal yang berhasil dimuat."

        analytics = FeedbackAnalyticsEngine(self.kb.df)
        executive_snapshot = analytics.build_executive_snapshot(timeframe, notes, sentiment=sentiment, segment=segment, score_engine=score_engine)
        report_sections = analytics.build_report_sections(timeframe, notes, macro_trends, sentiment=sentiment, segment=segment, score_engine=score_engine)

        document = Document()
        DocumentBuilder.create_cover(document, timeframe, DEFAULT_COLOR)
        document.add_heading("EXECUTIVE SNAPSHOT", level=1)
        DocumentBuilder.process_content(document, executive_snapshot, DEFAULT_COLOR)
        document.add_page_break()

        for index, section in enumerate(report_sections):
            document.add_heading(section["title"], level=1)
            DocumentBuilder.process_content(document, section["content"], DEFAULT_COLOR)
            if index < len(report_sections) - 1:
                document.add_page_break()

        filename = f"Inixindo_Feedback_Intelligence_Report_{score_profile['label']}_{timeframe}".replace(" ", "_")
        quality = ReportQualityValidator.evaluate(document, executive_snapshot, report_sections, score_profile["label"])
        if not quality["verified_complete"]:
            logger.warning(
                "Generated report is below completeness target: %s",
                quality["summary"],
            )
        return document, filename, quality
